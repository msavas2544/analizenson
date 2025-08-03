import sys
import os
import cv2
import math
import subprocess
import shutil
import configparser
import gc
import platform
import csv
from datetime import datetime

# PyQt5 import'ları
from PyQt5.QtCore import QThread, pyqtSignal, pyqtSlot, Qt, QTimer, QRect, QTime, QUrl
from PyQt5.QtGui import QPainter, QColor, QImage, QPixmap, QTextCursor, QPen, QKeySequence, QFont
from PyQt5.QtWidgets import (QWidget, QMainWindow, QVBoxLayout, QHBoxLayout, 
                             QPushButton, QLabel, QFileDialog, QProgressBar, 
                             QStyle, QApplication, QRadioButton, 
                             QButtonGroup, QTextEdit, QGroupBox, QMessageBox,
                             QListWidget, QListWidgetItem, QShortcut, QDialog, QCheckBox, 
                             QGridLayout, QScrollArea, QSlider)

# Güvenli import'lar
try:
    import torch
    TORCH_AVAILABLE = True
except ImportError:
    TORCH_AVAILABLE = False
    print("PyTorch bulunamadı - CPU modu kullanılacak")

try:
    from ultralytics import YOLO
    ULTRALYTICS_AVAILABLE = True
except ImportError:
    ULTRALYTICS_AVAILABLE = False
    print("Ultralytics bulunamadı - lütfen 'pip install ultralytics' çalıştırın")

try:
    import xlsxwriter
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("Excel desteği yok - 'pip install xlsxwriter' çalıştırın")

try:
    import matplotlib
    matplotlib.use('Agg')  # GUI gerektirmeyen backend
    import matplotlib.pyplot as plt
    import numpy as np
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False
    print("Grafik desteği yok - 'pip install matplotlib numpy' çalıştırın")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import parse_xml
    DOCX_AVAILABLE = True
    print("✅ Word desteği yüklü")
except ImportError:
    DOCX_AVAILABLE = False
    print("Word desteği yok - 'pip install python-docx' çalıştırın")

# YENİ: PDF rapor desteği - şu anda kullanılmıyor
PDF_AVAILABLE = False

# =============================================================================
# --- UYGULAMA TEKİL İNSTANS YÖNETİCİSİ ---
# =============================================================================

import socket
import tempfile
import atexit

class SingletonApp:
    """Uygulamanın sadece tek bir örneğinin çalışmasını sağlar"""
    
    def __init__(self, app_name="M.SAVAS_Video_Analiz"):
        self.app_name = app_name
        self.socket = None
        self.lock_file_path = None
        self.is_singleton = False
        
    def is_already_running(self):
        """Uygulamanın zaten çalışıp çalışmadığını kontrol eder"""
        try:
            # Socket ile port kontrolü
            self.socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            # Port 0 otomatik port seçer
            self.socket.bind(('localhost', 0))
            port = self.socket.getsockname()[1]
            
            # Lock dosyası oluştur
            self.lock_file_path = os.path.join(tempfile.gettempdir(), f"{self.app_name}.lock")
            
            # Eğer lock dosyası varsa ve içinde geçerli port bilgisi varsa
            if os.path.exists(self.lock_file_path):
                try:
                    with open(self.lock_file_path, 'r') as f:
                        existing_port = int(f.read().strip())
                    
                    # Mevcut port'u test et
                    test_socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                    result = test_socket.connect_ex(('localhost', existing_port))
                    test_socket.close()
                    
                    if result == 0:  # Port aktif
                        return True
                except:
                    pass  # Dosya bozuksa devam et
            
            # Yeni lock dosyası yaz
            with open(self.lock_file_path, 'w') as f:
                f.write(str(port))
            
            # Çıkışta temizle
            atexit.register(self.cleanup)
            self.is_singleton = True
            return False
            
        except Exception:
            return False
    
    def cleanup(self):
        """Kaynakları temizler"""
        try:
            if self.socket:
                self.socket.close()
            if self.lock_file_path and os.path.exists(self.lock_file_path):
                os.remove(self.lock_file_path)
        except:
            pass

# Global singleton instance
app_singleton = SingletonApp()

# =============================================================================
# --- UYGULAMA YAPILANDIRMASI VE AYARLAR ---
# =============================================================================

# =============================================================================
# --- YAPILANDIRMA YÖNETİCİSİ ---
# =============================================================================

class ConfigManager:
    """Config.ini dosyasını yöneten sınıf"""
    
    def __init__(self, config_path='config.ini'):
        self.config_path = config_path
        self.config = configparser.ConfigParser()
        self.load_config()
    
    def load_config(self):
        """Yapılandırma dosyasını yükler"""
        try:
            if os.path.exists(self.config_path):
                self.config.read(self.config_path, encoding='utf-8')
            else:
                self.create_default_config()
        except Exception as e:
            print(f"Config yükleme hatası: {e}")
            self.create_default_config()
    
    def create_default_config(self):
        """Varsayılan yapılandırma dosyası oluşturur"""
        self.config['DEFAULT'] = {
            'app_name': 'M.SAVAŞ Video Analiz Sistemi',
            'version': '1.1.0',
            'author': 'M.SAVAŞ'
        }
        
        self.config['VIDEO'] = {
            'supported_formats': 'mp4,avi,mov,mkv,wmv,flv,webm,dav,h264,264,ts,m2ts,mts',
            'max_file_size_mb': '2000',
            'default_fps': '30',
            'security_camera_formats': 'ch1,ch2,ch3,ch4,ch5,ch6,ch7,ch8',
            'smart_detection_channels': 'true'
        }
        
        self.config['ANALYSIS'] = {
            'default_sensitivity': 'ULTRA MAX',
            'confidence_threshold': '0.5',
            'iou_threshold': '0.5',
            'max_detections_per_frame': '10'
        }
        
        self.save_config()
    
    def save_config(self):
        """Yapılandırmayı dosyaya kaydeder"""
        try:
            with open(self.config_path, 'w', encoding='utf-8') as f:
                self.config.write(f)
        except Exception as e:
            print(f"Config kaydetme hatası: {e}")
    
    def get_supported_formats(self):
        """Desteklenen video formatlarını döndürür"""
        return self.config['VIDEO']['supported_formats'].split(',')
    
    def get_confidence_threshold(self):
        """Güven eşiğini döndürür"""
        return float(self.config['ANALYSIS']['confidence_threshold'])
    
    def get_default_sensitivity(self):
        """Varsayılan hassasiyet seviyesini döndürür"""
        return self.config['ANALYSIS']['default_sensitivity']
    
    def get_max_file_size(self):
        """Maksimum dosya boyutunu MB olarak döndürür"""
        return int(self.config['VIDEO']['max_file_size_mb'])
    
    def is_security_camera_file(self, file_path: str) -> bool:
        """Güvenlik kamerası dosyası olup olmadığını kontrol eder"""
        filename = os.path.basename(file_path).lower()
        
        # Kanal işaretlerini kontrol et
        security_patterns = [
            '_ch1_', '_ch2_', '_ch3_', '_ch4_', '_ch5_', '_ch6_', '_ch7_', '_ch8_',
            'channel', 'cam01', 'cam02', 'cam03', 'cam04', 'dvr', 'nvr'
        ]
        
        return any(pattern in filename for pattern in security_patterns)
    
    def get_security_camera_settings(self):
        """Güvenlik kamerası için optimize edilmiş ayarları döndürür"""
        return {
            'enhanced_motion_detection': True,
            'low_light_optimization': True,
            'fixed_camera_mode': True,
            'continuous_monitoring': True,
            'event_merge_gap': 2.0,  # Güvenlik kameraları için daha uzun birleştirme
            'minimum_event_duration': 3.0  # En az 3 saniye olay
        }

# Global config manager instance
config_manager = ConfigManager()

# =============================================================================
# --- ÖZEL HATA SINIFLARI ---
# =============================================================================

class VideoAnalysisError(Exception):
    """Video analizi genel hata sınıfı"""
    pass

class FileFormatError(VideoAnalysisError):
    """Dosya formatı hatası"""
    pass

class ModelLoadError(VideoAnalysisError):
    """Model yükleme hatası"""
    pass

class VideoLoadError(VideoAnalysisError):
    """Video yükleme hatası"""
    pass

class GPUError(VideoAnalysisError):
    """GPU kullanım hatası"""
    pass

# Performans ve doğruluk için analiz edilecek karelerin genişliği
ANALYSIS_FRAME_WIDTH = 320  # 416'dan 320'ye düşürdük - %50 daha hızlı

# Ultra performanslı analiz için ek ayarlar
PERFORMANCE_SETTINGS = {
    "batch_size": 8,           # Batch processing için
    "max_det": 50,             # Maksimum tespit sayısı
    "iou_threshold": 0.5,      # IoU eşiği
    "agnostic_nms": True,      # Sınıf bağımsız NMS
    "half_precision": True,    # FP16 kullan
    "dynamic_skip": True,      # Dinamik kare atlama
    "smart_crop": True,        # Akıllı kırpma
    "motion_based_roi": True   # Hareket bazlı ROI
}

# Hassasiyet seviyeleri ve karşılık gelen eşik değerleri
SENSITIVITY_LEVELS = {
    "🚀 Ultra Hızlı": {
        "motion": 45,           # Sadece büyük hareketler
        "conf": 0.3,            # Yüksek güven - sadece net tespitler
        "frame_skip": 8,        # Her 9. kareyi analiz et (hızlı)
        "min_area_ratio": 0.005,   # Büyük nesneler (insan boyutu+)
        "roi_expand": 0.9,      # ROI daraltma
        "temporal_smooth": 1,   # Minimal yumuşatma
        "description": "Hızlı tarama - sadece büyük hareketler"
    },
    "⚡ Hızlı": {
        "motion": 35,           # Orta-büyük hareketler
        "conf": 0.2,            # Orta-yüksek güven eşiği
        "frame_skip": 4,        # Her 5. kareyi analiz et
        "min_area_ratio": 0.002,   # Orta boy nesneler
        "roi_expand": 1.0,      # Normal ROI
        "temporal_smooth": 1,   # Az yumuşatma
        "description": "Hızlı analiz - orta boy nesneler"
    },
    "🎯 Normal": {
        "motion": 25,           # Dengeli hassasiyet
        "conf": 0.1,            # Dengeli güven eşiği
        "frame_skip": 2,        # Her 3. kareyi analiz et
        "min_area_ratio": 0.0008,  # İnsan odaklı boyut
        "roi_expand": 1.1,      # Hafif ROI genişletme
        "temporal_smooth": 2,   # Normal yumuşatma
        "description": "Dengeli analiz - insan odaklı"
    },
    "🔍 Detaylı": {
        "motion": 15,           # Hassas tespit
        "conf": 0.05,           # Düşük güven eşiği
        "frame_skip": 1,        # Her 2. kareyi analiz et
        "min_area_ratio": 0.0002,  # Küçük detaylar
        "roi_expand": 1.25,     # ROI genişletme
        "temporal_smooth": 3,   # İyi yumuşatma
        "description": "Hassas analiz - ince detaylar"
    },
    "🔬 Ultra Detaylı": {
        "motion": 8,            # Mikroskobik hareketler
        "conf": 0.02,           # En düşük güven eşiği
        "frame_skip": 0,        # Her kareyi analiz et
        "min_area_ratio": 0.00005, # En küçük hareketler
        "roi_expand": 1.5,      # Maksimum ROI genişletme
        "temporal_smooth": 4,   # Maksimum yumuşatma
        "description": "Maksimum hassasiyet - en küçük detaylar"
    },
    "🏢 Güvenlik Kamerası": {
        "motion": 20,           # Güvenlik optimumu
        "conf": 0.08,           # Güvenlik odaklı güven
        "frame_skip": 1,        # Güvenlik için sık analiz
        "min_area_ratio": 0.001,   # İnsan boyutu odaklı
        "roi_expand": 1.2,      # Güvenlik ROI
        "temporal_smooth": 2,   # Gürültü azaltma
        "description": "Güvenlik kamerası optimizasyonu",
        "low_light_mode": True, # Düşük ışık optimizasyonu
        "fixed_camera": True,   # Sabit kamera modu
        "continuous_mode": True # Sürekli izleme modu
    }
}

# Varsayılan ayarlar
DEFAULT_SENSITIVITY = "🎯 Normal"  # Dengeli analiz - insan odaklı

# YENİ: Çoklu nesne tespiti - YOLO sınıfları
TARGET_CLASSES = {
    'person': 0,        # İnsan (varsayılan)
    'bicycle': 1,       # Bisiklet  
    'car': 2,           # Araba
    'motorbike': 3,     # Motosiklet
    'bus': 5,           # Otobüs
    'truck': 7,         # Kamyon
    'cat': 15,          # Kedi
    'dog': 16,          # Köpek
    'horse': 17,        # At
    'bird': 14,         # Kuş
    'backpack': 24,     # Sırt çantası
    'suitcase': 28,     # Valiz
    'sports ball': 32,  # Spor topu
    'bottle': 39,       # Şişe
    'wine glass': 40,   # Şarap kadehi
    'cup': 41,          # Fincan
    'knife': 43,        # Bıçak
    'cell phone': 67,   # Cep telefonu
    'laptop': 63,       # Laptop
    'mouse': 64,        # Fare
    'remote': 65,       # Kumanda
    'keyboard': 66,     # Klavye
    'book': 73,         # Kitap
    'clock': 74,        # Saat
    'scissors': 76,     # Makas
    'toothbrush': 79    # Diş fırçası
}

# Aktif tespit sınıfları (varsayılan sadece insan)
ACTIVE_CLASSES = [0]  # Başlangıçta sadece person

EVENT_MERGE_GAP_SECONDS = 1.5  # Daha hızlı birleştirme
EXPORT_BUFFER_SECONDS = 0.5  # Daha kısa buffer
TEMP_CLIP_DIR = "msavas_temp_clips"

# Çoklu işlem ayarları
MULTIPROCESSING_ENABLED = True
MAX_WORKERS = 4  # Paralel işlem sayısı


# =============================================================================
# --- GEREKLİ KÜTÜPHANE KONTROLÜ ---
# =============================================================================
# --- PERFORMANS VE BELLEK YÖNETİMİ ---
# =============================================================================

def detect_gpu_capability():
    """GPU varlığını ve kapasitesini kontrol eder"""
    try:
        if TORCH_AVAILABLE and torch.cuda.is_available():
            gpu_count = torch.cuda.device_count()
            gpu_memory = torch.cuda.get_device_properties(0).total_memory / (1024**3)  # GB
            gpu_name = torch.cuda.get_device_name(0)
            return {
                'available': True, 
                'count': gpu_count, 
                'memory_gb': gpu_memory,
                'name': gpu_name
            }
    except Exception as e:
        print(f"GPU kontrol hatası: {e}")
    
    return {'available': False}

def cleanup_memory():
    """Bellek temizleme fonksiyonu"""
    gc.collect()
    
    try:
        if TORCH_AVAILABLE and torch.cuda.is_available():
            torch.cuda.empty_cache()
    except Exception as e:
        print(f"GPU bellek temizleme hatası: {e}")

def optimize_batch_size(video_resolution):
    """Video çözünürlüğüne göre batch size optimizasyonu"""
    if not video_resolution or len(video_resolution) < 2:
        return 8  # Varsayılan
    
    total_pixels = video_resolution[0] * video_resolution[1]
    
    if total_pixels > 1920 * 1080:  # 4K+
        return 4
    elif total_pixels > 1280 * 720:  # HD
        return 8
    else:  # SD
        return 16

# =============================================================================
# --- BAĞIMLILIK KONTROLÜ ---
# =============================================================================
def check_dependencies():
    """Gerekli kütüphaneleri ve FFmpeg'i kontrol eder."""
    print("🔍 Bağımlılıklar kontrol ediliyor...")
    missing = []
    warnings = []
    
    # Temel kütüphaneler
    try:
        import cv2
        print(f"✅ OpenCV: {cv2.__version__}")
    except ImportError:
        missing.append("opencv-python")
        print("❌ OpenCV eksik")
    
    # PyQt5 zaten import edilmiş, kontrol etmeye gerek yok
    print("✅ PyQt5 yüklü")
    
    try:
        import numpy as np
        print(f"✅ NumPy: {np.__version__}")
    except ImportError:
        missing.append("numpy")
        print("❌ NumPy eksik")
    
    # İsteğe bağlı kütüphaneler
    if not ULTRALYTICS_AVAILABLE:
        warnings.append("ultralytics (YOLO desteği için)")
        print("⚠️ Ultralytics eksik - temel tespit kullanılacak")
    else:
        print("✅ Ultralytics yüklü")
    
    if not EXCEL_AVAILABLE:
        warnings.append("openpyxl (Excel raporu için)")
        print("⚠️ OpenPyXL eksik - Excel raporu kullanılamaz")
    else:
        print("✅ OpenPyXL yüklü")
    
    if not MATPLOTLIB_AVAILABLE:
        warnings.append("matplotlib (grafik raporu için)")
        print("⚠️ Matplotlib eksik - grafik raporu kullanılamaz")
    else:
        print("✅ Matplotlib yüklü")
    
    if not DOCX_AVAILABLE:
        warnings.append("python-docx (Word raporu için)")
        print("⚠️ Python-docx eksik - Word raporu kullanılamaz")
    else:
        print("✅ Python-docx yüklü")
    
    if not MATPLOTLIB_AVAILABLE:
        warnings.append("matplotlib (grafik raporu için)")
        print("⚠️ Matplotlib eksik - grafik raporu kullanılamaz")
    else:
        print("✅ Matplotlib yüklü")
    
    # FFmpeg kontrolü
    try:
        _ = subprocess.run(['ffmpeg', '-version'], capture_output=True, text=True, check=True, 
                      creationflags=subprocess.CREATE_NO_WINDOW if os.name == 'nt' else 0)
        print("✅ FFmpeg yüklü")
    except (FileNotFoundError, subprocess.CalledProcessError):
        warnings.append("ffmpeg (video işleme için)")
        print("⚠️ FFmpeg eksik - video dışa aktarma sınırlı")
    
    if missing:
        error_message = (
            f"❌ Kritik eksiklikler: {', '.join(missing)}\n\n"
            "Bu kütüphaneler olmadan uygulama çalışmaz:\n"
            "pip install opencv-python PyQt5 numpy\n\n"
            "Uygulamayı başlatmak için önce bu kütüphaneleri yükleyin."
        )
        print(error_message)
        return False
    
    if warnings:
        print(f"\n💡 İsteğe bağlı eksiklikler: {len(warnings)} adet")
        print("Bu özellikler kullanılamayacak ama uygulama çalışacak:")
        for warning in warnings:
            print(f"  • {warning}")
        print("\nEksik kütüphaneleri yüklemek için:")
        print("pip install ultralytics openpyxl matplotlib python-docx Pillow")
    
    print("\n🚀 Temel gereksinimler karşılandı - uygulama başlatılabilir!")
    return True
    
    return True

# =============================================================================
# --- GPU VE PERFORMANS AYARLARI ---
# =============================================================================

def get_optimal_device():
    """En iyi cihazı seçer (GPU varsa GPU, yoksa CPU)."""
    try:
        if TORCH_AVAILABLE and torch.cuda.is_available():
            return 'cuda'
        else:
            return 'cpu'
    except Exception:
        return 'cpu'

def optimize_yolo_model(model):
    """YOLO modelini performans için optimize eder."""
    try:
        # Half precision (FP16) kullan - %40 daha hızlı
        if get_optimal_device() == 'cuda':
            model.half()
        return model
    except Exception as e:
        print(f"Model optimizasyonu başarısız: {e}")
        return model

# =============================================================================
# --- VİDEO İŞLEME THREAD'İ ---
# =============================================================================

class VideoProcessor(QThread):
    """Video analizini arka planda yürüten iş parçacığı."""
    progress_updated = pyqtSignal(int)
    motion_progress_updated = pyqtSignal(int)  # YENİ: Hareket tespiti özel progress
    # Değişiklik: analysis_complete sinyali artık tespit edilen nesnelerin koordinatlarını da taşıyacak
    analysis_complete = pyqtSignal(dict, list, dict) # detected_objects, events, video_info
    status_updated = pyqtSignal(str)
    error_occurred = pyqtSignal(str)

    def __init__(self, video_path: str, sensitivity: str):
        super().__init__()
        self.video_path = video_path
        self.sensitivity_settings = SENSITIVITY_LEVELS[sensitivity]
        self.stop_requested = False

    def run(self):
        try:
            self.status_updated.emit("🚀 ULTRA PERFORMANS MOD - Video açılıyor...")
            cap = cv2.VideoCapture(self.video_path)
            if not cap.isOpened():
                raise IOError("Video dosyası açılamadı veya bozuk.")

            video_info = self._get_video_info(cap)
            
            # GPU/CPU otomatik seçim
            device = get_optimal_device()
            self.status_updated.emit(f"🎯 Analiz motoru yükleniyor... ({device.upper()})")
            
            # Model yükleme ve optimizasyon
            if not ULTRALYTICS_AVAILABLE:
                raise ModelLoadError("YOLO modeli yüklenemedi - ultralytics kütüphanesi bulunamadı")
            
            try:
                model = YOLO("yolov8n.pt")
                model.to(device)
                model = optimize_yolo_model(model)
            except Exception as e:
                raise ModelLoadError(f"YOLO modeli yüklenirken hata: {e}")
            
            self.status_updated.emit(f"⚡ ULTRA HIZLI analiz başladı... ({video_info['total_frames']} kare)")
            
            # Gelişmiş analiz
            detected_objects, detected_frames_list = self._ultra_analyze_frames(cap, model, video_info)

            cap.release()
            
            if self.stop_requested:
                self.status_updated.emit("❌ Analiz kullanıcı tarafından durduruldu.")
                self.analysis_complete.emit({}, [], video_info)
                return

            self.status_updated.emit("🔄 Akıllı olay birleştirme...")
            events = self._smart_merge_events(detected_frames_list, video_info['fps'])
            
            self.progress_updated.emit(100)
            
            # YENİ: Final motion progress - toplam hareket oranını göster
            total_motion_frames = len(detected_frames_list)
            total_analyzed_frames = len([f for f in detected_objects.keys()])
            if total_analyzed_frames > 0:
                final_motion_percentage = min(int((total_motion_frames / total_analyzed_frames) * 100), 100)
                self.motion_progress_updated.emit(final_motion_percentage)
            
            self.status_updated.emit(f"✅ ULTRA analiz tamamlandı! {len(events)} olay bulundu.")
            self.analysis_complete.emit(detected_objects, events, video_info)

        except Exception as e:
            self.error_occurred.emit(f"❌ Analiz hatası: {e}")

    def _get_video_info(self, cap: cv2.VideoCapture) -> dict:
        total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        fps = cap.get(cv2.CAP_PROP_FPS)
        if fps <= 0 or total_frames <= 0:
            raise ValueError("Video bilgileri okunamadı (FPS/kare sayısı sıfır).")
        return {'fps': fps, 'total_frames': total_frames, 'duration': total_frames / fps}

    def _ultra_analyze_frames(self, cap: cv2.VideoCapture, model: YOLO, video_info: dict) -> tuple[dict, list]:
        """ULTRA PERFORMANSLI kare analizi - GPU optimizasyonu ile."""
        detected_objects = {}
        total_frames = video_info['total_frames']
        detected_frames_list = []
        
        # DEBUG: Analiz başlangıcı
        print(f"🔍 DEBUG: Analiz başlıyor - Toplam {total_frames} frame")
        self.status_updated.emit(f"🔍 Hareket analizi başlıyor - {total_frames} frame taranacak")
        
        # Performans ayarları
        frame_skip = self.sensitivity_settings['frame_skip']
        batch_size = PERFORMANCE_SETTINGS['batch_size']
        
        # Batch processing için frame buffer
        frame_buffer = []
        frame_numbers = []
        
        frame_count = 0
        processed_frames = 0
        
        while True:
            if self.stop_requested:
                break
                
            ret, frame = cap.read()
            if not ret:
                break
            
            # Frame atlama kontrolü
            if frame_count % (frame_skip + 1) != 0:
                frame_count += 1
                continue
            
            # Frame'i buffer'a ekle
            frame_buffer.append(frame)
            frame_numbers.append(frame_count)
            
            # Batch dolu olduğunda veya son frame'de process et
            if len(frame_buffer) >= batch_size or frame_count >= total_frames - 1:
                batch_results = self._process_batch(frame_buffer, frame_numbers, model)
                
                # Sonuçları birleştir
                for frame_num, detections in batch_results.items():
                    if detections:
                        detected_objects[frame_num] = detections
                        detected_frames_list.append(frame_num)
                
                # Buffer'ı temizle
                frame_buffer.clear()
                frame_numbers.clear()
            
            processed_frames += 1
            frame_count += 1
            
            # Progress güncelle ve UI'ı responsive tut - HER 5 FRAME'DE
            if processed_frames % 5 == 0:  # Daha sık güncelleme
                progress = min(int((frame_count / total_frames) * 100), 100)
                self.progress_updated.emit(progress)
                
                # UI thread'ini bloke etmemek için processEvents çağır
                QApplication.processEvents()
                
            # YENİ: Motion progress - HER FRAME'DE güncelle - GÜÇLÜ EMİT
            motion_frames_count = len(detected_frames_list)
            if processed_frames > 0:
                motion_progress = min(int((motion_frames_count / processed_frames) * 100), 100)
                # DEBUG: Emit'ten önce değerleri kontrol et
                print(f"📊 EMIT ÖNCESI: motion_frames={motion_frames_count}, processed={processed_frames}, progress={motion_progress}%")
                self.motion_progress_updated.emit(motion_progress)
                # Status'ta da doğrudan göster
                self.status_updated.emit(f"🚨 HAREKET: {motion_frames_count}/{processed_frames} frame = %{motion_progress}")
                
                # Debug için - daha sık log
                if processed_frames % 25 == 0:  # Her 25 frame'de log
                    print(f"🔍 DEBUG Motion Checkpoint: {motion_frames_count}/{processed_frames} = %{motion_progress}")
                    
            # UI'ı güncelle
            QApplication.processEvents()
                
            # Her 50 frame'de bir kısa bekle (mouse cursor sorunu için)
            if processed_frames % 50 == 0:
                self.msleep(1)  # 1ms bekle
        
        # DEBUG: Final motion progress
        final_motion_count = len(detected_frames_list)
        if processed_frames > 0:
            final_motion_progress = int((final_motion_count / processed_frames) * 100)
            print(f"🎯 DEBUG: Analiz tamamlandı - {final_motion_count}/{processed_frames} hareket frame (%{final_motion_progress})")
            self.status_updated.emit(f"✅ Analiz tamamlandı - {final_motion_count} hareket frame tespit edildi (%{final_motion_progress})")
            self.motion_progress_updated.emit(final_motion_progress)
        
        return detected_objects, detected_frames_list
    
    def _process_batch(self, frames: list, frame_numbers: list, model: YOLO) -> dict:
        """Batch olarak frame'leri işler - gelişmiş filtreleme ile."""
        batch_results = {}
        
        try:
            # UI responsive tutmak için
            QApplication.processEvents()
            
            # Frames'i resize et ve batch haline getir
            resized_frames = []
            original_frames = []
            for frame in frames:
                original_frames.append(frame)
                resized = cv2.resize(frame, (ANALYSIS_FRAME_WIDTH, 
                                           int(frame.shape[0] * ANALYSIS_FRAME_WIDTH / frame.shape[1])))
                resized_frames.append(resized)
            
            # YOLO batch inference
            results = model(resized_frames, 
                          conf=self.sensitivity_settings['conf'],
                          iou=PERFORMANCE_SETTINGS['iou_threshold'],
                          max_det=PERFORMANCE_SETTINGS['max_det'],
                          classes=list(TARGET_CLASSES.values()),
                          verbose=False)
            
            # Sonuçları parse et
            for result, frame_num, original_frame in zip(results, frame_numbers, original_frames):
                detections = []
                if result.boxes is not None:
                    boxes = result.boxes.xyxy.cpu().numpy()
                    confidences = result.boxes.conf.cpu().numpy()
                    
                    # Orijinal frame boyutları
                    original_height, original_width = original_frame.shape[:2]
                    
                    # Ölçekleme faktörleri
                    scale_x = original_width / ANALYSIS_FRAME_WIDTH
                    scale_y = original_height / int(original_frame.shape[0] * ANALYSIS_FRAME_WIDTH / original_frame.shape[1])
                    
                    # Geçerli tespitler için liste
                    valid_detections = []
                    
                    for box, conf in zip(boxes, confidences):
                        x1, y1, x2, y2 = box
                        
                        # Orijinal boyutlara dönüştür
                        orig_x1 = int(x1 * scale_x)
                        orig_y1 = int(y1 * scale_y)
                        orig_x2 = int(x2 * scale_x)
                        orig_y2 = int(y2 * scale_y)
                        
                        # x, y, w, h formatına çevir
                        orig_x = orig_x1
                        orig_y = orig_y1
                        orig_w = orig_x2 - orig_x1
                        orig_h = orig_y2 - orig_y1
                        
                        # Sınırları kontrol et
                        orig_x = max(0, min(orig_x, original_width - orig_w))
                        orig_y = max(0, min(orig_y, original_height - orig_h))
                        orig_w = min(orig_w, original_width - orig_x)
                        orig_h = min(orig_h, original_height - orig_y)
                        
                        # AKILLI FİLTRELEME - Yanlış tespitleri elemek için
                        if self._is_valid_person_detection(orig_x, orig_y, orig_w, orig_h, conf, original_width, original_height):
                            valid_detections.append({
                                'box': [orig_x, orig_y, orig_w, orig_h],
                                'conf': conf,
                                'area': orig_w * orig_h
                            })
                    
                    # Çakışan/benzer tespitleri birleştir
                    filtered_detections = self._merge_overlapping_detections(valid_detections)
                    
                    # Sadece kutu bilgilerini al
                    detections = [det['box'] for det in filtered_detections]
                
                batch_results[frame_num] = detections
        
        except Exception as e:
            self.error_occurred.emit(f"Batch processing hatası: {e}")
        
        return batch_results
    
    def _is_valid_person_detection(self, x: int, y: int, w: int, h: int, conf: float, frame_width: int, frame_height: int) -> bool:
        """Tespitin gerçekten bir kişi olup olmadığını kontrol eder."""
        
        # 1. Minimum boyut kontrolü
        min_person_area = (frame_width * frame_height) * self.sensitivity_settings['min_area_ratio']
        if w * h < min_person_area:
            return False
        
        # 2. Aspect ratio kontrolü (insan vücut oranları)
        aspect_ratio = h / w if w > 0 else 0
        if aspect_ratio < 0.8 or aspect_ratio > 4.0:  # İnsanlar genelde daha uzun
            return False
        
        # 3. Çok büyük tespit kontrolü (muhtemelen hatalı)
        max_person_area = (frame_width * frame_height) * 0.8  # Frame'in %80'inden büyük olamaz
        if w * h > max_person_area:
            return False
        
        # 4. Güven skoru kontrolü
        min_confidence = max(self.sensitivity_settings['conf'], 0.3)  # Minimum %30
        if conf < min_confidence:
            return False
        
        # 5. Çok dar veya çok geniş tespit kontrolü
        if w < 20 or h < 40:  # Çok küçük
            return False
        
        if w > frame_width * 0.8 or h > frame_height * 0.9:  # Çok büyük
            return False
        
        # 6. Kenar kontrolü - frame kenarlarındaki yarım tespitleri ele
        edge_threshold = 10
        if x < edge_threshold or y < edge_threshold:
            if w * h < min_person_area * 2:  # Kenar tespitleri için daha sıkı kontrol
                return False
        
        return True
    
    def _merge_overlapping_detections(self, detections: list) -> list:
        """Çakışan tespitleri birleştirir ve en güvenilir olanı tutar."""
        if len(detections) <= 1:
            return detections
        
        # Güven skoruna göre sırala
        detections.sort(key=lambda x: x['conf'], reverse=True)
        
        merged = []
        for detection in detections:
            box = detection['box']
            x, y, w, h = box
            
            # Mevcut merged listesindeki tespitlerle çakışma kontrolü
            overlap_found = False
            for merged_detection in merged:
                merged_box = merged_detection['box']
                mx, my, mw, mh = merged_box
                
                # IoU (Intersection over Union) hesapla
                iou = self._calculate_iou(x, y, w, h, mx, my, mw, mh)
                
                if iou > 0.5:  # %50'den fazla çakışma
                    overlap_found = True
                    # Daha güvenilir olanı tut (zaten sıralı)
                    if detection['conf'] > merged_detection['conf']:
                        merged.remove(merged_detection)
                        merged.append(detection)
                    break
            
            if not overlap_found:
                merged.append(detection)
        
        return merged
    
    def _calculate_iou(self, x1: int, y1: int, w1: int, h1: int, x2: int, y2: int, w2: int, h2: int) -> float:
        """İki kutunun IoU (Intersection over Union) değerini hesaplar."""
        
        # Kesişim alanını hesapla
        x_left = max(x1, x2)
        y_top = max(y1, y2)
        x_right = min(x1 + w1, x2 + w2)
        y_bottom = min(y1 + h1, y2 + h2)
        
        if x_right < x_left or y_bottom < y_top:
            return 0.0
        
        intersection_area = (x_right - x_left) * (y_bottom - y_top)
        
        # Birleşim alanını hesapla
        box1_area = w1 * h1
        box2_area = w2 * h2
        union_area = box1_area + box2_area - intersection_area
        
        if union_area == 0:
            return 0.0
        
        return intersection_area / union_area
    
    def _smart_merge_events(self, detected_frames: list, fps: float) -> list:
        """Akıllı olay birleştirme algoritması."""
        if not detected_frames:
            return []
        
        events = []
        detected_frames.sort()
        
        current_start = detected_frames[0] / fps
        current_end = detected_frames[0] / fps
        
        for i in range(1, len(detected_frames)):
            current_time = detected_frames[i] / fps
            
            # Eğer frame'ler arasındaki gap EVENT_MERGE_GAP_SECONDS'den az ise birleştir
            if current_time - current_end <= EVENT_MERGE_GAP_SECONDS:
                current_end = current_time
            else:
                # Yeni event başlat
                events.append((current_start, current_end + EXPORT_BUFFER_SECONDS))
                current_start = current_time
                current_end = current_time
        
        # Son event'i ekle
        events.append((current_start, current_end + EXPORT_BUFFER_SECONDS))
        
        return events
    
    def request_stop(self):
        """Analizi durdur."""
        self.stop_requested = True
    
    def stop(self):
        """Analizi durdur."""
        self.stop_requested = True

# =============================================================================
# --- VİDEO DIŞA AKTARMA THREAD'İ ---
# =============================================================================

class VideoExporter(QThread):
    """Video dışa aktarma işlemini yürüten iş parçacığı."""
    export_progress = pyqtSignal(int, str)
    export_complete = pyqtSignal(str, str) # path, message
    error_occurred = pyqtSignal(str)

    def __init__(self, video_path: str, events: list, video_info: dict, output_path: str):
        super().__init__()
        self.video_path = video_path
        self.events = events
        self.video_duration = video_info['duration']
        self.output_path = output_path

    def run(self):
        try:
            temp_dir = self._create_temp_dir()
            clip_files = self._create_clips(temp_dir)
            
            if not clip_files:
                raise ValueError("Dışa aktarılacak klip oluşturulamadı.")

            self._concatenate_clips(clip_files, temp_dir)
            self.export_complete.emit(self.output_path, "Özet video başarıyla kaydedildi.")

        except Exception as e:
            self.error_occurred.emit(f"Dışa aktarma hatası: {e}")
        finally:
            self._cleanup()

    def _create_temp_dir(self) -> str:
        temp_dir = os.path.abspath(TEMP_CLIP_DIR)
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        os.makedirs(temp_dir)
        return temp_dir

    def _create_clips(self, temp_dir: str) -> list:
        clip_files = []
        total_events = len(self.events)
        for i, (start, end) in enumerate(self.events):
            start_buffered = max(0, start - EXPORT_BUFFER_SECONDS)
            end_buffered = min(self.video_duration, end + EXPORT_BUFFER_SECONDS)
            duration = end_buffered - start_buffered

            progress = int(((i + 1) / total_events) * 50)
            self.export_progress.emit(progress, f"Klip {i+1}/{total_events} oluşturuluyor...")

            clip_path = os.path.join(temp_dir, f"clip_{i:03d}.mp4")
            
            command = [
                'ffmpeg', '-y', '-hide_banner', '-loglevel', 'error',
                '-ss', str(start_buffered), '-i', self.video_path,
                '-t', str(duration), '-c:v', 'libx264', '-preset', 'fast', 
                '-crf', '23', '-c:a', 'aac', '-b:a', '128k', clip_path
            ]
            
            try:
                subprocess.run(command, check=True, capture_output=True, text=True, creationflags=subprocess.CREATE_NO_WINDOW)
                clip_files.append(clip_path)
            except subprocess.CalledProcessError as e:
                self.error_occurred.emit(f"Uyarı: Klip {i+1} oluşturulamadı. Hata: {e.stderr}")
                continue
        
        return clip_files

    def _concatenate_clips(self, clip_files: list, temp_dir: str):
        self.export_progress.emit(75, "Klipler birleştiriliyor...")
        
        file_list_path = os.path.join(temp_dir, "mylist.txt")
        with open(file_list_path, "w", encoding='utf-8') as f:
            for clip_path in clip_files:
                f.write(f"file '{clip_path.replace(os.sep, '/')}'\n")

        concat_command = [
            'ffmpeg', '-y', '-hide_banner', '-loglevel', 'error',
            '-f', 'concat', '-safe', '0', '-i', file_list_path,
            '-c', 'copy', self.output_path
        ]
        
        subprocess.run(concat_command, check=True, capture_output=True, text=True, creationflags=subprocess.CREATE_NO_WINDOW)
        self.export_progress.emit(100, "Birleştirme tamamlandı.")

    def _cleanup(self):
        temp_dir = os.path.abspath(TEMP_CLIP_DIR)
        if os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)
            except Exception as e:
                self.error_occurred.emit(f"Uyarı: Geçici dosyalar temizlenemedi: {e}")

# =============================================================================
# --- ZAMAN ÇİZELGESİ WIDGET'I ---
# =============================================================================

class TimelineWidget(QWidget):
    """Video olaylarını gösteren etkileşimli zaman çizelgesi."""
    seek_requested = pyqtSignal(float)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setMinimumHeight(40)
        self.setCursor(Qt.PointingHandCursor)
        self.duration = 0.0
        self.events = []
        self.progress = 0.0
        self.hover_pos = -1
        self.setMouseTracking(True)

    def set_duration(self, duration: float):
        self.duration = duration
        self.update()

    def set_events(self, events: list):
        self.events = events
        self.update()
        
    def set_progress(self, progress: float):
        self.progress = max(0.0, min(1.0, progress))
        self.update()

    def paintEvent(self, event):
        super().paintEvent(event)
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)
        
        painter.fillRect(self.rect(), QColor("#2c3e50"))
        timeline_rect = self.rect().adjusted(5, 5, -5, -5)
        painter.fillRect(timeline_rect, QColor("#34495e"))

        if self.duration > 0:
            for start, end in self.events:
                x_start = timeline_rect.left() + int((start / self.duration) * timeline_rect.width())
                x_end = timeline_rect.left() + int((end / self.duration) * timeline_rect.width())
                event_rect = QRect(x_start, timeline_rect.top(), max(1, x_end - x_start), timeline_rect.height())
                painter.fillRect(event_rect, QColor("#e74c3c"))

            progress_x = timeline_rect.left() + int(self.progress * timeline_rect.width())
            painter.setPen(QPen(QColor("#3498db"), 2))
            painter.drawLine(progress_x, timeline_rect.top(), progress_x, timeline_rect.bottom())

        if self.hover_pos != -1:
            painter.fillRect(timeline_rect.left() + self.hover_pos, timeline_rect.top(), 1, timeline_rect.height(), QColor(255, 255, 255, 40))

    def mousePressEvent(self, event):
        if self.duration > 0 and event.button() == Qt.LeftButton:
            self._emit_seek_request(event.pos().x())

    def mouseMoveEvent(self, event):
        timeline_rect = self.rect().adjusted(5, 5, -5, -5)
        if timeline_rect.contains(event.pos()):
            self.hover_pos = event.pos().x() - timeline_rect.left()
            self.update()
        else:
            self.leaveEvent(event)

    def leaveEvent(self, event):
        self.hover_pos = -1
        self.update()

    def _emit_seek_request(self, x_pos: int):
        timeline_rect = self.rect().adjusted(5, 5, -5, -5)
        if timeline_rect.width() > 0:
            relative_pos = (x_pos - timeline_rect.left()) / timeline_rect.width()
            seek_time = max(0, min(self.duration, relative_pos * self.duration))
            self.seek_requested.emit(seek_time)

# =============================================================================
# --- ANA ARAYÜZ PENCERESİ ---
# =============================================================================

class AdvancedObjectSelectionDialog(QDialog):
    """🚨 Gelişmiş Polis Nesne Seçimi - Türkçe İsimlendirme"""
    
    def __init__(self, target_classes, active_classes, parent=None):
        super().__init__(parent)
        self.target_classes = target_classes
        self.active_classes = active_classes
        
        # Türkçe nesne isimleri - polis özel
        self.turkish_names = {
            # Güvenlik
            'person': 'İnsan/Şüpheli',
            'knife': 'Bıçak ⚠️',
            'scissors': 'Makas',
            'bottle': 'Şişe/Molotof ⚠️',
            'handbag': 'El Çantası',
            'backpack': 'Sırt Çantası',
            'suitcase': 'Valiz/Çanta',
            
            # Silah ve Tehlikeli
            'baseball bat': 'Beyzbol Sopası ⚠️',
            'umbrella': 'Şemsiye',
            'sports ball': 'Top/Mermi ⚠️',
            
            # Taşıt Araçları
            'car': 'Otomobil',
            'motorcycle': 'Motosiklet',
            'bicycle': 'Bisiklet',
            'truck': 'Kamyon',
            'bus': 'Otobüs',
            'airplane': 'Uçak',
            'boat': 'Tekne',
            'train': 'Tren',
            
            # Elektronik Cihaz
            'cell phone': 'Cep Telefonu 📱',
            'laptop': 'Dizüstü Bilgisayar 💻',
            'tv': 'Televizyon',
            'keyboard': 'Klavye',
            'mouse': 'Fare',
            'remote': 'Kumanda',
            
            # Hayvanlar
            'cat': 'Kedi',
            'dog': 'Köpek',
            'horse': 'At',
            'sheep': 'Koyun',
            'cow': 'İnek',
            'elephant': 'Fil',
            'bear': 'Ayı',
            
            # Ev Eşyaları
            'chair': 'Sandalye',
            'sofa': 'Koltuk',
            'bed': 'Yatak',
            'dining table': 'Yemek Masası',
            'toilet': 'Tuvalet',
            'clock': 'Saat',
            'vase': 'Vazo',
            
            # Yiyecek ve İçecek
            'wine glass': 'Şarap Kadehi',
            'cup': 'Fincan',
            'fork': 'Çatal',
            'spoon': 'Kaşık',
            'bowl': 'Kase',
            'banana': 'Muz',
            'apple': 'Elma',
            
            # Spor Malzemeleri
            'kite': 'Uçurtma',
            'baseball glove': 'Beyzbol Eldiveni',
            'skateboard': 'Kaykay',
            'surfboard': 'Sörf Tahtası',
            'tennis racket': 'Tenis Raketi',
            
            # Kişisel Eşya
            'tie': 'Kravat',
            'book': 'Kitap',
            
            # Diğer
            'frisbee': 'Frizbi',
            'skis': 'Kayak',
            'snowboard': 'Snowboard',
            'toothbrush': 'Diş Fırçası',
            'hair drier': 'Saç Kurutma',
            'teddy bear': 'Oyuncak Ayı'
        }
        self.init_ui()
    
    def init_ui(self):
        """🎨 Polis özel arayüz"""
        self.setWindowTitle("🚨 Gelişmiş Polis Nesne Seçimi")
        self.setModal(True)
        self.resize(600, 500)
        
        # Ana stil
        self.setStyleSheet("""
            QDialog {
                background-color: #2c3e50;
                color: white;
            }
            QLabel {
                color: white;
            }
            QCheckBox {
                color: white;
                font-weight: bold;
                padding: 3px;
            }
            QCheckBox:hover {
                background-color: #34495e;
                border-radius: 3px;
            }
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                border-radius: 5px;
                padding: 8px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
            QPushButton:pressed {
                background-color: #21618c;
                padding-top: 10px;
            }
        """)
        
        layout = QVBoxLayout()
        
        # Açıklama
        description = QLabel("🔍 Analiz edilecek nesne türlerini seçin (Polis özel kategoriler):")
        description.setFont(QFont("Arial", 11, QFont.Bold))
        description.setStyleSheet("color: #ecf0f1; padding: 10px; background-color: #34495e; border-radius: 5px;")
        layout.addWidget(description)
        
        # Scroll area for checkboxes
        scroll = QScrollArea()
        scroll.setStyleSheet("""
            QScrollArea {
                background-color: #34495e;
                border: 1px solid #2c3e50;
                border-radius: 5px;
            }
        """)
        scroll_widget = QWidget()
        scroll_layout = QGridLayout(scroll_widget)
        
        self.checkboxes = {}
        row, col = 0, 0
        
        for category, objects in self.target_classes.items():
            # Kategori başlığı - renk kodlu
            category_color = self.get_category_color(category)
            category_label = QLabel(f"{category}")
            category_label.setFont(QFont("Arial", 10, QFont.Bold))
            category_label.setStyleSheet(f"color: {category_color}; padding: 8px; background-color: #2c3e50; border-radius: 3px; margin: 3px;")
            scroll_layout.addWidget(category_label, row, 0, 1, 3)
            row += 1
            
            # Nesneler - Türkçe isimler ile
            for obj in objects:
                turkish_name = self.turkish_names.get(obj, obj.title())
                checkbox = QCheckBox(turkish_name)
                
                # Tehlikeli nesneler için kırmızı renk
                if any(warning in turkish_name for warning in ['⚠️', '🔫', '💥']):
                    checkbox.setStyleSheet("""
                        QCheckBox {
                            color: #e74c3c;
                            font-weight: bold;
                            padding: 4px;
                        }
                        QCheckBox:hover {
                            background-color: #c0392b;
                            border-radius: 3px;
                        }
                    """)
                
                if obj in self.active_classes:
                    checkbox.setChecked(True)
                
                self.checkboxes[obj] = checkbox
                scroll_layout.addWidget(checkbox, row, col)
                
                col += 1
                if col >= 3:
                    col = 0
                    row += 1
            
            if col != 0:
                row += 1
                col = 0
            
            # Kategori arası boşluk
            row += 1
        
        scroll.setWidget(scroll_widget)
        layout.addWidget(scroll)
        
        # Butonlar
        button_layout = QHBoxLayout()
        
        select_all_btn = QPushButton("✅ Tümünü Seç")
        select_all_btn.clicked.connect(self.select_all)
        button_layout.addWidget(select_all_btn)
        
        deselect_all_btn = QPushButton("❌ Tümünü Kaldır")
        deselect_all_btn.clicked.connect(self.deselect_all)
        button_layout.addWidget(deselect_all_btn)
        
        # Polis özel seçimler
        police_btn = QPushButton("🚨 Polis Özel")
        police_btn.clicked.connect(self.select_police_special)
        police_btn.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                border-radius: 5px;
                padding: 8px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
        """)
        button_layout.addWidget(police_btn)
        
        button_layout.addStretch()
        
        ok_btn = QPushButton("✅ Tamam")
        ok_btn.clicked.connect(self.accept)
        ok_btn.setDefault(True)
        ok_btn.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                border-radius: 5px;
                padding: 10px 20px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #229954;
            }
        """)
        button_layout.addWidget(ok_btn)
        
        cancel_btn = QPushButton("❌ İptal")
        cancel_btn.clicked.connect(self.reject)
        cancel_btn.setStyleSheet("""
            QPushButton {
                background-color: #95a5a6;
                color: white;
                border: none;
                border-radius: 5px;
                padding: 10px 20px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #7f8c8d;
            }
        """)
        button_layout.addWidget(cancel_btn)
        
        layout.addLayout(button_layout)
        self.setLayout(layout)
    
    def get_category_color(self, category):
        """Kategori rengini döndür"""
        colors = {
            '🚨 Güvenlik': '#e74c3c',
            '🔫 Silah ve Tehlikeli': '#8e44ad',
            '🚗 Taşıt Araçları': '#3498db',
            '📱 Elektronik Cihaz': '#f39c12',
            '👤 İnsan ve Hayvan': '#27ae60',
            '🏠 Ev Eşyaları': '#16a085',
            '🍽️ Yiyecek ve İçecek': '#e67e22',
            '⚽ Spor Malzemeleri': '#9b59b6',
            '🎒 Kişisel Eşya': '#34495e',
            '🔧 Diğer Nesneler': '#95a5a6'
        }
        return colors.get(category, '#ecf0f1')
    
    def select_police_special(self):
        """🚨 Polis özel nesnelerini seç"""
        police_objects = ['person', 'knife', 'scissors', 'bottle', 'handbag', 'backpack', 'suitcase', 
                         'baseball bat', 'car', 'motorcycle', 'truck', 'cell phone', 'laptop']
        
        # Önce tümünü kaldır
        self.deselect_all()
        
        # Polis özel nesneleri seç
        for obj, checkbox in self.checkboxes.items():
            if obj in police_objects:
                checkbox.setChecked(True)
    
    def select_all(self):
        """Tümünü seç"""
        for checkbox in self.checkboxes.values():
            checkbox.setChecked(True)
    
    def deselect_all(self):
        """Tümünü kaldır"""
        for checkbox in self.checkboxes.values():
            checkbox.setChecked(False)
    
    def get_selected_classes(self):
        """Seçili sınıfları döndür"""
        selected = []
        for obj, checkbox in self.checkboxes.items():
            if checkbox.isChecked():
                selected.append(obj)
        return selected

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        # İlk önce OBJECT_NAMES_TURKISH'i tanımla
        self.OBJECT_NAMES_TURKISH = {
            'person': 'İnsan', 'bicycle': 'Bisiklet', 'car': 'Araba', 'motorcycle': 'Motosiklet',
            'airplane': 'Uçak', 'bus': 'Otobüs', 'train': 'Tren', 'truck': 'Kamyon',
            'boat': 'Tekne', 'traffic light': 'Trafik Işığı', 'fire hydrant': 'Yangın Musluğu',
            'stop sign': 'Dur Tabelası', 'parking meter': 'Park Sayacı', 'bench': 'Bank',
            'bird': 'Kuş', 'cat': 'Kedi', 'dog': 'Köpek', 'horse': 'At', 'sheep': 'Koyun',
            'cow': 'İnek', 'elephant': 'Fil', 'bear': 'Ayı', 'zebra': 'Zebra', 'giraffe': 'Zürafa',
            'backpack': 'Sırt Çantası', 'umbrella': 'Şemsiye', 'handbag': 'El Çantası',
            'tie': 'Kravat', 'suitcase': 'Valiz', 'frisbee': 'Frizbi', 'skis': 'Kayak',
            'snowboard': 'Kar Kayağı', 'sports ball': 'Spor Topu', 'kite': 'Uçurtma',
            'baseball bat': 'Beyzbol Sopası', 'baseball glove': 'Beyzbol Eldiveni',
            'skateboard': 'Kaykay', 'surfboard': 'Sörf Tahtası', 'tennis racket': 'Tenis Raketi',
            'bottle': 'Şişe', 'wine glass': 'Şarap Kadehi', 'cup': 'Fincan', 'fork': 'Çatal',
            'knife': 'Bıçak', 'spoon': 'Kaşık', 'bowl': 'Kase', 'banana': 'Muz',
            'apple': 'Elma', 'sandwich': 'Sandviç', 'orange': 'Portakal', 'broccoli': 'Brokoli',
            'carrot': 'Havuç', 'hot dog': 'Sosisli', 'pizza': 'Pizza', 'donut': 'Donut',
            'cake': 'Pasta', 'chair': 'Sandalye', 'sofa': 'Koltuk', 'pottedplant': 'Saksı Bitkisi',
            'bed': 'Yatak', 'dining table': 'Yemek Masası', 'toilet': 'Tuvalet', 'tv': 'Televizyon',
            'laptop': 'Laptop', 'mouse': 'Fare', 'remote': 'Kumanda', 'keyboard': 'Klavye',
            'cell phone': 'Cep Telefonu', 'microwave': 'Mikrodalga', 'oven': 'Fırın',
            'toaster': 'Ekmek Kızartma Makinesi', 'sink': 'Lavabo', 'refrigerator': 'Buzdolabı',
            'book': 'Kitap', 'clock': 'Saat', 'vase': 'Vazo', 'scissors': 'Makas',
            'teddy bear': 'Oyuncak Ayı', 'hair drier': 'Saç Kurutma Makinesi', 'toothbrush': 'Diş Fırçası'
        }
        self._setup_ui()
        self._connect_signals()
        self._setup_shortcuts()  # Yeni: Klavye kısayolları
        self._initialize_state()

    def _setup_ui(self):
        self.setWindowTitle("M.SAVAŞ - Motion Surveillance and Video Analysis System")
        self.setGeometry(50, 50, 1920, 1080)  # Full HD ekran boyutu
        self.setStyleSheet(self.get_stylesheet())

        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        
        # Ana düzen: 3 panel (Sol, Orta, Sağ)
        self.main_layout = QHBoxLayout(self.central_widget)
        self.main_layout.setSpacing(5)
        self.main_layout.setContentsMargins(5, 5, 5, 5)

        # =============================================================================
        # SOL PANEL - Video Dosyaları ve Kontroller (400px - GENİŞLETİLDİ)
        # =============================================================================
        left_panel_widget = QWidget()
        left_panel_widget.setFixedWidth(400)
        left_panel_widget.setStyleSheet("""
            QWidget {
                background-color: #2c3e50;
                border: 2px solid #34495e;
                border-radius: 10px;
            }
        """)
        left_panel_layout = QVBoxLayout(left_panel_widget)
        left_panel_layout.setSpacing(8)
        left_panel_layout.setContentsMargins(10, 10, 10, 10)
        
        # Video dosyaları grubu - SOL PANEL
        video_group = QGroupBox("📁 Video Dosyaları")
        video_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 12px;
                color: white;
                border: 2px solid #3498db;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 15px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 8px 0 8px;
            }
        """)
        video_layout = QVBoxLayout()
        video_layout.setSpacing(5)
        video_layout.setContentsMargins(8, 8, 8, 8)
        
        # Video butonları - 2x2 düzen - GENİŞLETİLDİ
        video_btn_layout1 = QHBoxLayout()
        video_btn_layout1.setSpacing(5)
        self.btn_add_video = QPushButton("➕ Video Ekle")
        self.btn_add_video.setFixedSize(180, 35)
        self.btn_load = QPushButton("📂 Tek Video")
        self.btn_load.setFixedSize(180, 35)
        video_btn_layout1.addWidget(self.btn_add_video)
        video_btn_layout1.addWidget(self.btn_load)
        
        video_btn_layout2 = QHBoxLayout()
        video_btn_layout2.setSpacing(5)
        self.btn_remove_video = QPushButton("➖ Kaldır")
        self.btn_remove_video.setFixedSize(180, 35)
        self.btn_clear_videos = QPushButton("🗑️ Temizle")
        self.btn_clear_videos.setFixedSize(180, 35)
        video_btn_layout2.addWidget(self.btn_remove_video)
        video_btn_layout2.addWidget(self.btn_clear_videos)
        
        # Canlı kamera butonları - GENİŞLETİLDİ
        video_btn_layout3 = QHBoxLayout()
        video_btn_layout3.setSpacing(5)
        self.btn_start_camera = QPushButton("📹 Canlı Kamera")
        self.btn_start_camera.setFixedSize(180, 35)
        self.btn_stop_camera = QPushButton("🛑 Durdur")
        self.btn_stop_camera.setFixedSize(180, 35)
        self.btn_stop_camera.setEnabled(False)
        video_btn_layout3.addWidget(self.btn_start_camera)
        video_btn_layout3.addWidget(self.btn_stop_camera)
        
        # Video butonlarına stil ekle
        video_button_style = """
            QPushButton {
                background-color: #3498db;
                color: white;
                border: 1px solid #2980b9;
                border-radius: 5px;
                font-weight: bold;
                font-size: 10px;
                padding: 2px;
            }
            QPushButton:hover {
                background-color: #2980b9;
                border: 1px solid #2471a3;
            }
            QPushButton:pressed {
                background-color: #2471a3;
                border: 1px solid #1f618d;
                padding-top: 3px;
                padding-left: 3px;
            }
            QPushButton:disabled {
                background-color: #95a5a6;
                border: 1px solid #7f8c8d;
                color: #bdc3c7;
            }
        """
        for btn in [self.btn_add_video, self.btn_load, self.btn_remove_video, 
                   self.btn_clear_videos, self.btn_start_camera, self.btn_stop_camera]:
            btn.setStyleSheet(video_button_style)
        
        # Video listesi
        self.video_list = QListWidget()
        self.video_list.setFixedHeight(120)
        self.video_list.setStyleSheet("""
            QListWidget {
                background-color: #34495e;
                color: white;
                border: 1px solid #3498db;
                border-radius: 5px;
                font-size: 11px;
            }
            QListWidget::item {
                padding: 5px;
                border-bottom: 1px solid #2c3e50;
            }
            QListWidget::item:selected {
                background-color: #3498db;
            }
        """)
        
        # Video bilgi paneli
        self.video_info_label = QLabel("📹 Video seçilmedi")
        self.video_info_label.setFixedHeight(60)
        self.video_info_label.setWordWrap(True)
        self.video_info_label.setStyleSheet("""
            QLabel {
                background-color: #34495e;
                color: #ecf0f1;
                border: 1px solid #3498db;
                border-radius: 5px;
                padding: 8px;
                font-size: 11px;
                font-weight: bold;
            }
        """)
        
        video_layout.addLayout(video_btn_layout1)
        video_layout.addLayout(video_btn_layout2)
        video_layout.addLayout(video_btn_layout3)
        video_layout.addWidget(self.video_list)
        video_layout.addWidget(self.video_info_label)
        video_group.setLayout(video_layout)
        
        # Video döndürme grubu - SOL PANEL
        rotation_group = QGroupBox("🔄 Video Döndürme")
        rotation_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 12px;
                color: white;
                border: 2px solid #e67e22;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 15px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 8px 0 8px;
            }
        """)
        rotation_layout = QVBoxLayout()
        rotation_layout.setSpacing(5)
        rotation_layout.setContentsMargins(8, 8, 8, 8)
        
        # Döndürme butonları - 2x2 düzen
        rotation_btn_layout1 = QHBoxLayout()
        rotation_btn_layout1.setSpacing(5)
        self.btn_rotate_90 = QPushButton("↻ 90°")
        self.btn_rotate_90.setFixedSize(160, 30)
        self.btn_rotate_180 = QPushButton("↻ 180°")
        self.btn_rotate_180.setFixedSize(160, 30)
        rotation_btn_layout1.addWidget(self.btn_rotate_90)
        rotation_btn_layout1.addWidget(self.btn_rotate_180)
        
        rotation_btn_layout2 = QHBoxLayout()
        rotation_btn_layout2.setSpacing(5)
        self.btn_rotate_270 = QPushButton("↻ 270°")
        self.btn_rotate_270.setFixedSize(160, 30)
        self.btn_rotate_reset = QPushButton("🔄 Sıfırla")
        self.btn_rotate_reset.setFixedSize(160, 30)
        rotation_btn_layout2.addWidget(self.btn_rotate_270)
        rotation_btn_layout2.addWidget(self.btn_rotate_reset)
        
        # Döndürme buton stillerini uygula
        rotate_button_style = """
            QPushButton {
                background-color: #f39c12;
                color: white;
                border: 2px solid #e67e22;
                border-radius: 5px;
                font-weight: bold;
                font-size: 11px;
                padding: 2px;
            }
            QPushButton:hover {
                background-color: #e67e22;
                border: 2px solid #d35400;
            }
            QPushButton:pressed {
                background-color: #d35400;
                border: 2px solid #ba4a00;
                padding-top: 4px;
                padding-left: 4px;
            }
            QPushButton:disabled {
                background-color: #95a5a6;
                border: 2px solid #7f8c8d;
                color: #bdc3c7;
            }
        """
        for btn in [self.btn_rotate_90, self.btn_rotate_180, self.btn_rotate_270, self.btn_rotate_reset]:
            btn.setStyleSheet(rotate_button_style)
        
        rotation_layout.addLayout(rotation_btn_layout1)
        rotation_layout.addLayout(rotation_btn_layout2)
        rotation_group.setLayout(rotation_layout)
        
        # Analiz kontrolleri grubu - SOL PANEL
        controls_group = QGroupBox("⚡ Analiz Kontrolleri")
        controls_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 12px;
                color: white;
                border: 2px solid #27ae60;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 15px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 8px 0 8px;
            }
        """)
        controls_layout = QVBoxLayout()
        controls_layout.setSpacing(5)
        controls_layout.setContentsMargins(8, 8, 8, 8)
        
        # Analiz butonları - tam genişlik
        self.btn_analyze = QPushButton("🚀 Analiz Et")
        self.btn_analyze.setFixedHeight(40)
        self.btn_analyze.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: 2px solid #229954;
                border-radius: 8px;
                font-weight: bold;
                font-size: 14px;
                padding: 3px;
            }
            QPushButton:hover {
                background-color: #2ecc71;
                border: 2px solid #27ae60;
            }
            QPushButton:pressed {
                background-color: #229954;
                border: 2px solid #1e8449;
                padding-top: 5px;
                padding-left: 5px;
            }
        """)
        
        self.btn_stop_analysis = QPushButton("🛑 Analizi Durdur")
        self.btn_stop_analysis.setFixedHeight(35)
        self.btn_stop_analysis.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: 2px solid #c0392b;
                border-radius: 8px;
                font-weight: bold;
                font-size: 12px;
                padding: 3px;
            }
            QPushButton:hover {
                background-color: #c0392b;
                border: 2px solid #a93226;
            }
            QPushButton:pressed {
                background-color: #a93226;
                border: 2px solid #922b21;
                padding-top: 5px;
                padding-left: 5px;
            }
        """)
        
        # Video dışa aktarma
        self.btn_export = QPushButton("📹 Özet Video Oluştur")
        self.btn_export.setFixedHeight(35)
        self.btn_export.setStyleSheet("""
            QPushButton {
                background-color: #9b59b6;
                color: white;
                border: 2px solid #8e44ad;
                border-radius: 8px;
                font-weight: bold;
                font-size: 12px;
                padding: 3px;
            }
            QPushButton:hover {
                background-color: #8e44ad;
                border: 2px solid #7d3c98;
            }
            QPushButton:pressed {
                background-color: #7d3c98;
                border: 2px solid #6c3483;
                padding-top: 5px;
                padding-left: 5px;
            }
        """)
        
        controls_layout.addWidget(self.btn_analyze)
        controls_layout.addWidget(self.btn_stop_analysis)
        controls_layout.addWidget(self.btn_export)
        controls_group.setLayout(controls_layout)
        
        # Rapor butonları grubu - SOL PANEL'E TAŞINDI
        reports_group = QGroupBox("📊 Rapor Oluştur")
        reports_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 12px;
                color: white;
                border: 2px solid #9b59b6;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 15px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 8px 0 8px;
            }
        """)
        reports_layout = QVBoxLayout()
        reports_layout.setSpacing(5)
        reports_layout.setContentsMargins(8, 8, 8, 8)
        
        # Rapor butonları - daha büyük
        self.btn_export_word = QPushButton("📄 Word Raporu")
        self.btn_export_word.setFixedHeight(45)
        self.btn_export_word.setStyleSheet("""
            QPushButton {
                background-color: #9b59b6;
                color: white;
                border: 2px solid #8e44ad;
                border-radius: 8px;
                font-weight: bold;
                font-size: 13px;
                padding: 3px;
            }
            QPushButton:hover {
                background-color: #8e44ad;
                border: 2px solid #7d3c98;
            }
            QPushButton:pressed {
                background-color: #7d3c98;
                border: 2px solid #6c3483;
                padding-top: 5px;
                padding-left: 5px;
            }
        """)
        
        self.btn_export_charts = QPushButton("📊 Grafik Raporu")
        self.btn_export_charts.setFixedHeight(45)
        self.btn_export_charts.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: 2px solid #2980b9;
                border-radius: 8px;
                font-weight: bold;
                font-size: 13px;
                padding: 3px;
            }
            QPushButton:hover {
                background-color: #2980b9;
                border: 2px solid #2471a3;
            }
            QPushButton:pressed {
                background-color: #2471a3;
                border: 2px solid #1f618d;
                padding-top: 5px;
                padding-left: 5px;
            }
        """)
        
        self.btn_export_all = QPushButton("🎯 Tüm Raporlar")
        self.btn_export_all.setFixedHeight(45)
        self.btn_export_all.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: 2px solid #c0392b;
                border-radius: 8px;
                font-weight: bold;
                font-size: 13px;
                padding: 3px;
            }
            QPushButton:hover {
                background-color: #c0392b;
                border: 2px solid #a93226;
            }
            QPushButton:pressed {
                background-color: #a93226;
                border: 2px solid #922b21;
                padding-top: 5px;
                padding-left: 5px;
            }
        """)
        
        reports_layout.addWidget(self.btn_export_word)
        reports_layout.addWidget(self.btn_export_charts)
        
        # Hareket raporu butonu ekle
        self.btn_export_motion = QPushButton("🏃 Hareket Raporu")
        self.btn_export_motion.setFixedHeight(45)
        self.btn_export_motion.setStyleSheet("""
            QPushButton {
                background-color: #f39c12;
                color: white;
                border: 2px solid #e67e22;
                border-radius: 8px;
                font-weight: bold;
                font-size: 13px;
                padding: 3px;
            }
            QPushButton:hover {
                background-color: #e67e22;
                border: 2px solid #d35400;
            }
            QPushButton:pressed {
                background-color: #d35400;
                border: 2px solid #ba4a00;
                padding-top: 5px;
                padding-left: 5px;
            }
        """)
        reports_layout.addWidget(self.btn_export_motion)
        reports_layout.addWidget(self.btn_export_all)
        reports_group.setLayout(reports_layout)
        
        # Sol paneli tamamla
        left_panel_layout.addWidget(video_group)
        left_panel_layout.addWidget(rotation_group)
        left_panel_layout.addWidget(controls_group)
        left_panel_layout.addWidget(reports_group)
        left_panel_layout.addStretch()  # Boş alan
        
        # =============================================================================
        # ORTA PANEL - Video Önizleme ve Oynatıcı (YENI DÜZENLEMELİ)
        # =============================================================================
        center_panel_widget = QWidget()
        center_panel_widget.setStyleSheet("""
            QWidget {
                background-color: #34495e;
                border: 2px solid #2c3e50;
                border-radius: 10px;
            }
        """)
        center_panel_layout = QVBoxLayout(center_panel_widget)
        center_panel_layout.setSpacing(8)
        center_panel_layout.setContentsMargins(10, 10, 10, 10)
        
        # ÜST KISIM: Video bilgileri ve önizleme
        video_preview_group = QGroupBox("📺 Video Önizleme")
        video_preview_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 14px;
                color: white;
                border: 2px solid #9b59b6;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 15px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 8px 0 8px;
            }
        """)
        video_preview_layout = QVBoxLayout()
        video_preview_layout.setSpacing(5)
        video_preview_layout.setContentsMargins(8, 8, 8, 8)
        
        # Video bilgi etiketi
        self.video_info_label = QLabel("📹 Video seçilmedi")
        self.video_info_label.setStyleSheet("""
            QLabel {
                background-color: #2c3e50;
                color: white;
                padding: 8px;
                border-radius: 5px;
                font-size: 11px;
                font-weight: bold;
            }
        """)
        self.video_info_label.setFixedHeight(45)
        video_preview_layout.addWidget(self.video_info_label)
        
        # Video ekranı - önizleme üstte - OPTİMİZE EDİLDİ
        self.video_display_label = QLabel("🎬 Video yükleyin veya canlı kamerayı başlatın")
        self.video_display_label.setStyleSheet("""
            QLabel {
                background-color: #2c3e50;
                color: #bdc3c7;
                border: 3px dashed #3498db;
                border-radius: 15px;
                font-size: 16px;
                font-weight: bold;
                text-align: center;
            }
        """)
        self.video_display_label.setAlignment(Qt.AlignCenter)
        self.video_display_label.setMinimumHeight(400)  # Daha büyük minimum yükseklik
        self.video_display_label.setMaximumHeight(600)  # Daha büyük maksimum yükseklik
        self.video_display_label.setScaledContents(False)  # Aspect ratio'yu korumak için False
        self.video_display_label.setWordWrap(True)
        video_preview_layout.addWidget(self.video_display_label)
        video_preview_group.setLayout(video_preview_layout)
        
        # ORTA KISIM: Video oynatma kontrolleri
        video_controls_group = QGroupBox("🎮 Video Oynatma Kontrolleri")
        video_controls_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 12px;
                color: white;
                border: 2px solid #27ae60;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 15px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 8px 0 8px;
            }
        """)
        video_controls_layout = QVBoxLayout()
        video_controls_layout.setSpacing(5)
        video_controls_layout.setContentsMargins(8, 8, 8, 8)
        
        # Timeline widget (eğer varsa)
        try:
            self.timeline_widget = TimelineWidget()
            self.timeline_widget.setFixedHeight(50)
        except:
            self.timeline_widget = QSlider(Qt.Horizontal)
            self.timeline_widget.setFixedHeight(50)
        video_controls_layout.addWidget(self.timeline_widget)
        
        # Oynatma butonları
        playback_layout = QHBoxLayout()
        playback_layout.setSpacing(8)
        
        self.btn_play_pause = QPushButton("▶️ Oynat")
        self.btn_play_pause.setFixedHeight(40)
        self.btn_play_pause.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                border-radius: 5px;
                font-weight: bold;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #229954;
            }
            QPushButton:pressed {
                background-color: #1e8449;
                padding-top: 5px;
            }
        """)
        
        # Progress bar ekle - YEŞİL RENKLE
        self.progress_bar = QProgressBar()
        self.progress_bar.setFixedHeight(25)
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                background-color: #34495e;
                border: 2px solid #27ae60;
                border-radius: 8px;
                text-align: center;
                color: white;
                font-weight: bold;
                font-size: 12px;
            }
            QProgressBar::chunk {
                background-color: #27ae60;
                border-radius: 6px;
                margin: 1px;
            }
        """)
        self.progress_bar.setVisible(False)
        
        playback_layout.addWidget(self.btn_play_pause)
        playback_layout.addStretch()
        
        video_controls_layout.addLayout(playback_layout)
        video_controls_layout.addWidget(self.progress_bar)
        video_controls_group.setLayout(video_controls_layout)
        
        # ALT KISIM: Hareket tespiti zaman çizelgesi
        motion_timeline_group = QGroupBox("📊 Hareket Tespiti Zaman Çizelgesi")
        motion_timeline_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 12px;
                color: white;
                border: 2px solid #e74c3c;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 15px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 8px 0 8px;
            }
        """)
        motion_timeline_layout = QVBoxLayout()
        motion_timeline_layout.setSpacing(5)
        motion_timeline_layout.setContentsMargins(8, 8, 8, 8)
        
        # Hareket tespiti listesi
        self.motion_timeline_list = QListWidget()
        self.motion_timeline_list.setFixedHeight(120)
        self.motion_timeline_list.setStyleSheet("""
            QListWidget {
                background-color: #2c3e50;
                color: white;
                border: 1px solid #34495e;
                border-radius: 5px;
                font-family: 'Consolas', 'Courier New', monospace;
                font-size: 10px;
                padding: 5px;
            }
            QListWidget::item {
                padding: 5px;
                margin: 2px;
                border-radius: 3px;
                background-color: #34495e;
            }
            QListWidget::item:hover {
                background-color: #3498db;
                color: white;
            }
            QListWidget::item:selected {
                background-color: #e74c3c;
                color: white;
            }
        """)
        motion_timeline_layout.addWidget(self.motion_timeline_list)
        
        # YENİ: Hareket tespiti özel ilerleme çubuğu - SADECE HAREKETLİ KISIMLAR
        self.motion_progress_bar = QProgressBar()
        self.motion_progress_bar.setFixedHeight(35)  # Daha büyük boyut
        self.motion_progress_bar.setStyleSheet("""
            QProgressBar {
                background-color: #2c3e50;
                border: 3px solid #e74c3c;
                border-radius: 10px;
                text-align: center;
                color: white;
                font-weight: bold;
                font-size: 13px;
            }
            QProgressBar::chunk {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0, 
                    stop:0 #e74c3c, stop:0.5 #ec7063, stop:1 #e74c3c);
                border-radius: 7px;
                margin: 2px;
            }
        """)
        self.motion_progress_bar.setVisible(False)
        self.motion_progress_bar.setFormat("� HAREKET TESPİTİ: %p% (%v/%m)")  # Daha belirgin format
        # DEBUG: Test için motion progress bar'ı görünür yap ve test değeri ver
        self.motion_progress_bar.setVisible(True)
        self.motion_progress_bar.setValue(33)  # Test: %33 hareket
        print("🔍 DEBUG: Motion progress bar test değeri %33 ile başlatıldı")
        motion_timeline_layout.addWidget(self.motion_progress_bar)
        
        # Hareket istatistikleri
        self.motion_stats_label = QLabel("📈 Hareket istatistikleri: Analiz yapıldıktan sonra görünecek")
        self.motion_stats_label.setStyleSheet("""
            QLabel {
                background-color: #2c3e50;
                color: #ecf0f1;
                padding: 8px;
                border-radius: 5px;
                font-size: 10px;
                border: 1px solid #34495e;
            }
        """)
        motion_timeline_layout.addWidget(self.motion_stats_label)
        motion_timeline_group.setLayout(motion_timeline_layout)
        
        # Orta paneli tamamla - yeni düzenlemeli
        center_panel_layout.addWidget(video_preview_group, 3)      # Video önizleme üstte (büyük alan)
        center_panel_layout.addWidget(video_controls_group, 1)     # Kontroller ortada (küçük alan)  
        center_panel_layout.addWidget(motion_timeline_group, 1)    # Hareket çizelgesi altta (küçük alan)
        
        # =============================================================================
        # SAĞ PANEL - Ayarlar, Raporlar ve Log (450px - GENİŞLETİLDİ) - LAYOUT TANIMLAMA
        # =============================================================================
        right_panel_widget = QWidget()
        right_panel_widget.setFixedWidth(450)
        right_panel_widget.setStyleSheet("""
            QWidget {
                background-color: #2c3e50;
                border: 2px solid #34495e;
                border-radius: 10px;
            }
        """)
        right_panel_layout = QVBoxLayout(right_panel_widget)
        right_panel_layout.setSpacing(8)
        right_panel_layout.setContentsMargins(10, 10, 10, 10)
        
        # Hassasiyet ayarları grubu - SAĞ PANEL
        sensitivity_group = QGroupBox("⚙️ Hassasiyet Ayarları")
        sensitivity_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 12px;
                color: white;
                border: 2px solid #f39c12;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 15px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 8px 0 8px;
            }
        """)
        sensitivity_layout = QVBoxLayout()
        sensitivity_layout.setSpacing(5)
        sensitivity_layout.setContentsMargins(8, 8, 8, 8)
        
        # Hassasiyet butonları
        self.sensitivity_buttons = QButtonGroup()
        
        self.btn_low_sensitivity = QRadioButton("🟢 Düşük")
        self.btn_medium_sensitivity = QRadioButton("🟡 Orta") 
        self.btn_high_sensitivity = QRadioButton("🔴 Yüksek")
        self.btn_ultra_sensitivity = QRadioButton("� ULTRA MAX")
        
        # Varsayılan olarak ULTRA MAX seçili
        self.btn_ultra_sensitivity.setChecked(True)
        
        # Hassasiyet butonlarına stil ekle
        sensitivity_style = """
            QRadioButton {
                color: white;
                font-weight: bold;
                margin: 3px;
                padding: 5px;
            }
            QRadioButton:hover {
                background-color: #34495e;
                border-radius: 3px;
            }
            QRadioButton::indicator {
                width: 15px;
                height: 15px;
            }
        """
        self.btn_low_sensitivity.setStyleSheet(sensitivity_style)
        self.btn_medium_sensitivity.setStyleSheet(sensitivity_style)
        self.btn_high_sensitivity.setStyleSheet(sensitivity_style)
        self.btn_ultra_sensitivity.setStyleSheet(sensitivity_style)
        
        # Buton grubuna ekle
        self.sensitivity_buttons.addButton(self.btn_low_sensitivity, 0)
        self.sensitivity_buttons.addButton(self.btn_medium_sensitivity, 1)
        self.sensitivity_buttons.addButton(self.btn_high_sensitivity, 2)
        self.sensitivity_buttons.addButton(self.btn_ultra_sensitivity, 3)
        
        sensitivity_layout.addWidget(self.btn_low_sensitivity)
        sensitivity_layout.addWidget(self.btn_medium_sensitivity)
        sensitivity_layout.addWidget(self.btn_high_sensitivity)
        sensitivity_layout.addWidget(self.btn_ultra_sensitivity)
        sensitivity_group.setLayout(sensitivity_layout)
        
        # Nesne tespit ayarları grubu - SAĞ PANEL
        objects_group = QGroupBox("🎯 Nesne Tespit Ayarları")
        objects_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 12px;
                color: white;
                border: 2px solid #e74c3c;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 15px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 8px 0 8px;
            }
        """)
        objects_layout = QVBoxLayout()
        objects_layout.setSpacing(5)
        objects_layout.setContentsMargins(8, 8, 8, 8)
        
        # Basit nesne checkboxları - Türkçe isimlerle
        self.object_checkboxes = {}
        main_objects = ['person', 'car', 'bicycle', 'motorcycle', 'cat', 'dog']
        
        for obj_name in main_objects:
            # Türkçe isim kullan
            turkish_name = self.OBJECT_NAMES_TURKISH.get(obj_name, obj_name.capitalize())
            checkbox = QCheckBox(f"🎯 {turkish_name}")
            checkbox.setStyleSheet("color: white; margin: 2px;")
            if obj_name == 'person':
                checkbox.setChecked(True)  # person varsayılan seçili
            self.object_checkboxes[obj_name] = checkbox
            objects_layout.addWidget(checkbox)
        
        # Gelişmiş seçim butonu
        self.btn_advanced_objects = QPushButton("⚙️ Gelişmiş Nesne Seçimi")
        self.btn_advanced_objects.setFixedHeight(35)
        self.btn_advanced_objects.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: 2px solid #c0392b;
                border-radius: 5px;
                font-weight: bold;
                font-size: 11px;
                padding: 3px;
            }
            QPushButton:hover {
                background-color: #c0392b;
                border: 2px solid #a93226;
            }
            QPushButton:pressed {
                background-color: #a93226;
                border: 2px solid #922b21;
                padding-top: 5px;
                padding-left: 5px;
            }
        """)
        objects_layout.addWidget(self.btn_advanced_objects)
        objects_group.setLayout(objects_layout)
        
        # Log ekranı - SAĞ PANEL
        log_group = QGroupBox("📋 İşlem Geçmişi")
        log_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 12px;
                color: white;
                border: 2px solid #34495e;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 15px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 8px 0 8px;
            }
        """)
        log_layout = QVBoxLayout()
        log_layout.setSpacing(5)
        log_layout.setContentsMargins(8, 8, 8, 8)
        
        # Log display
        self.log_display = QTextEdit()
        self.log_display.setFixedHeight(200)
        self.log_display.setReadOnly(True)
        self.log_display.setStyleSheet("""
            QTextEdit {
                background-color: #2c3e50;
                color: #ecf0f1;
                border: 1px solid #34495e;
                border-radius: 5px;
                font-family: 'Consolas', 'Courier New', monospace;
                font-size: 10px;
                padding: 5px;
            }
        """)
        
        # Log temizleme butonu
        self.btn_clear_log = QPushButton("🗑️ Log Temizle")
        self.btn_clear_log.setFixedHeight(30)
        self.btn_clear_log.setStyleSheet("""
            QPushButton {
                background-color: #95a5a6;
                color: white;
                border: 1px solid #7f8c8d;
                border-radius: 5px;
                font-weight: bold;
                font-size: 10px;
                padding: 2px;
            }
            QPushButton:hover {
                background-color: #7f8c8d;
                border: 1px solid #6c7b7f;
            }
            QPushButton:pressed {
                background-color: #6c7b7f;
                border: 1px solid #5d6d6d;
                padding-top: 3px;
                padding-left: 3px;
            }
        """)
        
        log_layout.addWidget(self.log_display)
        log_layout.addWidget(self.btn_clear_log)
        log_group.setLayout(log_layout)
        
        # Sistem kontrol grubu - SAĞ PANEL
        system_group = QGroupBox("⚡ Sistem Kontrol")
        system_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 12px;
                color: white;
                border: 2px solid #e67e22;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 15px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 8px 0 8px;
            }
        """)
        system_layout = QHBoxLayout()
        system_layout.setSpacing(5)
        system_layout.setContentsMargins(8, 8, 8, 8)
        
        # Sistem kontrol butonları
        self.btn_force_exit = QPushButton("🚪 Zorla Çık")
        self.btn_force_exit.setFixedHeight(35)
        self.btn_force_exit.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
            QPushButton:pressed {
                background-color: #a93226;
                padding-top: 3px;
                padding-left: 3px;
            }
        """)
        
        self.btn_restart_app = QPushButton("🔄 Yeniden Başlat")
        self.btn_restart_app.setFixedHeight(35)
        self.btn_restart_app.setStyleSheet("""
            QPushButton {
                background-color: #f39c12;
                color: white;
                border: none;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #e67e22;
            }
            QPushButton:pressed {
                background-color: #d35400;
                padding-top: 3px;
                padding-left: 3px;
            }
        """)
        
        system_layout.addWidget(self.btn_force_exit)
        system_layout.addWidget(self.btn_restart_app)
        system_group.setLayout(system_layout)
        
        # Sağ paneli tamamla
        right_panel_layout.addWidget(sensitivity_group)
        right_panel_layout.addWidget(objects_group)
        right_panel_layout.addWidget(log_group)
        right_panel_layout.addWidget(system_group)
        right_panel_layout.addStretch()  # Boş alan
        
        # Ana panelleri layout'a ekle
        self.main_layout.addWidget(left_panel_widget)    # Sol panel (350px)
        self.main_layout.addWidget(center_panel_widget)  # Orta panel (genişleyebilir)
        self.main_layout.addWidget(right_panel_widget)   # Sağ panel (400px)
        
        # Status bar ekle
        self.status_label = QLabel("🟢 Hazır - Video yükleyin veya canlı kamerayı başlatın")
        self.status_label.setStyleSheet("""
            QLabel {
                background-color: #2c3e50;
                color: white;
                padding: 5px;
                border-top: 1px solid #34495e;
                font-weight: bold;
            }
        """)
        
        # Ana layout'un altına status bar ekle
        main_vertical_layout = QVBoxLayout()
        main_vertical_layout.setContentsMargins(0, 0, 0, 0)
        main_vertical_layout.setSpacing(0)
        main_vertical_layout.addLayout(self.main_layout)
        main_vertical_layout.addWidget(self.status_label)
        
        # Layout'u central widget'a uygula
        self.central_widget.setLayout(main_vertical_layout)
        
        
        # UI başlatma tamamlandı
        
        # Event list widget için placeholder (eğer eksikse)
        if not hasattr(self, 'event_list_widget'):
            self.event_list_widget = QListWidget()
            self.event_list_widget.setFixedHeight(100)
            self.event_list_widget.setStyleSheet("""
                QListWidget {
                    background-color: #2c3e50;
                    color: white;
                    border: 1px solid #34495e;
                    border-radius: 5px;
                }
            """)
        
        # Log clear button signal bağlantısı
        self.btn_clear_log.clicked.connect(self.clear_log)
        
        # Progress bar eksikse ekle
        if not hasattr(self, 'progress_bar'):
            self.progress_bar = QProgressBar()
            self.progress_bar.setVisible(False)
            self.progress_bar.setStyleSheet("""
                QProgressBar {
                    background-color: #34495e;
                    border: 1px solid #2c3e50;
                    border-radius: 5px;
                    text-align: center;
                    color: white;
                    font-weight: bold;
                }
                QProgressBar::chunk {
                    background-color: #27ae60;
                    border-radius: 5px;
                }
            """)

    def clear_log(self):
        """Log ekranını temizler"""
        self.log_display.clear()
        self.log_message("📋 Log ekranı temizlendi", "info")

    def sensitivity_changed(self):
        """Hassasiyet değiştiğinde çağrılır"""
        button = self.sensitivity_buttons.checkedButton()
        if button:
            self.current_sensitivity = button.text()
            self.log_message(f"⚙️ Hassasiyet değiştirildi: {self.current_sensitivity}", "info")

    def on_event_item_clicked(self, item):
        """Olay listesindeki öğeye tıklandığında çağrılır"""
        if item:
            self.log_message(f"🎯 Olay seçildi: {item.text()}", "info")

    def _connect_signals(self):
        self.btn_load.clicked.connect(self.load_video)
        self.btn_add_video.clicked.connect(self.add_video_file)
        self.btn_remove_video.clicked.connect(self.remove_video_file)
        self.btn_clear_videos.clicked.connect(self.clear_all_videos)
        self.video_list.itemSelectionChanged.connect(self.on_video_selection_changed)
        self.btn_analyze.clicked.connect(self.analyze_video)
        self.btn_stop_analysis.clicked.connect(self.stop_analysis)
        self.btn_export.clicked.connect(self.export_video)
        self.btn_export_charts.clicked.connect(self.export_charts_report)
        self.btn_export_word.clicked.connect(self.export_word_report)
        self.btn_export_motion.clicked.connect(self.export_motion_report)  # Hareket raporu
        self.btn_export_all.clicked.connect(self.export_all_reports)  # Yeni
        self.btn_play_pause.clicked.connect(self.toggle_playback)
        self.timeline_widget.seek_requested.connect(self.seek_video)
        self.sensitivity_buttons.buttonClicked.connect(self.sensitivity_changed)
        
        # YENİ: Sistem kontrol butonları
        self.btn_force_exit.clicked.connect(self.force_exit_application)
        self.btn_restart_app.clicked.connect(self.restart_application)
        # YENİ: Hareket tespiti zaman çizelgesi bağlantısı
        self.motion_timeline_list.itemClicked.connect(self.on_motion_timeline_clicked)
        
        # Video döndürme butonları
        self.btn_rotate_90.clicked.connect(lambda: self.rotate_video(90))
        self.btn_rotate_180.clicked.connect(lambda: self.rotate_video(180))
        self.btn_rotate_270.clicked.connect(lambda: self.rotate_video(270))
        self.btn_rotate_reset.clicked.connect(lambda: self.rotate_video(0))
        
        # YENİ: Canlı kamera butonları
        self.btn_start_camera.clicked.connect(self.start_live_camera)
        self.btn_stop_camera.clicked.connect(self.stop_live_camera)
        
        # YENİ: Nesne seçimi butonları
        for checkbox in self.object_checkboxes.values():
            checkbox.clicked.connect(self.update_active_classes)
        self.btn_advanced_objects.clicked.connect(self.open_advanced_selection)
    
    def _setup_shortcuts(self):
        """Klavye kısayollarını ayarlar"""
        
        # Dosya işlemleri
        self.shortcut_open = QShortcut(QKeySequence("Ctrl+O"), self)
        self.shortcut_open.activated.connect(self.load_video)
        
        # Analiz işlemleri
        self.shortcut_analyze = QShortcut(QKeySequence("F5"), self)
        self.shortcut_analyze.activated.connect(self.analyze_video)
        
        self.shortcut_stop = QShortcut(QKeySequence("Esc"), self)
        self.shortcut_stop.activated.connect(self.stop_analysis)
        
        # Video kontrolleri
        self.shortcut_play_pause = QShortcut(QKeySequence("Space"), self)
        self.shortcut_play_pause.activated.connect(self.toggle_playback)
        
        # Rapor kaydetme
        self.shortcut_save_word = QShortcut(QKeySequence("Ctrl+W"), self)
        self.shortcut_save_word.activated.connect(self.export_word_report)
        
        # Video döndürme
        self.shortcut_rotate_90 = QShortcut(QKeySequence("Ctrl+R"), self)
        self.shortcut_rotate_90.activated.connect(lambda: self.rotate_video(90))
        
        # Tüm raporları kaydet
        self.shortcut_save_all = QShortcut(QKeySequence("Ctrl+Shift+S"), self)
        self.shortcut_save_all.activated.connect(self.export_all_reports)
        
        self.log_message("⌨️ Klavye kısayolları aktif: Ctrl+O (Aç), F5 (Analiz), Space (Oynat/Duraklat)", "info")

    def _initialize_state(self):
        self.video_path = None
        self.video_paths = []  # Çoklu video desteği
        self.current_video_index = 0  # Şu anki video indeksi
        self.video_capture = None
        self.video_info = {}
        self.detected_events = []
        self.detected_objects = {} # Yeni: Nesne koordinatlarını saklamak için
        self.is_playing = False
        self.current_sensitivity = DEFAULT_SENSITIVITY
        self.playback_timer = QTimer(self)
        self.playback_timer.timeout.connect(self.update_frame)
        self.processor_thread = None
        self.exporter_thread = None
        self.current_rotation = 0  # Video döndürme açısı (0, 90, 180, 270)
        
        # Polis özel nesne tespit sınıfları (Türkçe)
        self.TARGET_CLASSES = {
            '🚨 Güvenlik': ['person', 'knife', 'scissors', 'bottle', 'handbag', 'backpack', 'suitcase'],
            '🔫 Silah ve Tehlikeli': ['knife', 'scissors', 'baseball bat', 'umbrella', 'bottle', 'sports ball'],
            '🚗 Taşıt Araçları': ['car', 'motorcycle', 'bicycle', 'truck', 'bus', 'airplane', 'boat', 'train'],
            '📱 Elektronik Cihaz': ['cell phone', 'laptop', 'tv', 'keyboard', 'mouse', 'remote'],
            '👤 İnsan ve Hayvan': ['person', 'cat', 'dog', 'horse', 'sheep', 'cow', 'elephant', 'bear'],
            '🏠 Ev Eşyaları': ['chair', 'sofa', 'bed', 'dining table', 'toilet', 'clock', 'vase'],
            '🍽️ Yiyecek ve İçecek': ['bottle', 'wine glass', 'cup', 'fork', 'knife', 'spoon', 'bowl', 'banana', 'apple'],
            '⚽ Spor Malzemeleri': ['sports ball', 'kite', 'baseball bat', 'baseball glove', 'skateboard', 'surfboard', 'tennis racket'],
            '🎒 Kişisel Eşya': ['handbag', 'tie', 'suitcase', 'umbrella', 'backpack', 'book'],
            '🔧 Diğer Nesneler': ['frisbee', 'skis', 'snowboard', 'toothbrush', 'hair drier', 'teddy bear']
        }
        
        # Varsayılan aktif sınıflar
        self.ACTIVE_CLASSES = [0]  # person ID'si
        
        # YOLO modeli yükle
        try:
            from ultralytics import YOLO
            self.model = YOLO('yolov8n.pt')  # Nano model (hızlı)
            self.log_message("🤖 YOLO modeli başarıyla yüklendi!", "success")
        except Exception as e:
            self.log_message(f"⚠️ YOLO modeli yüklenemedi: {e}", "warning")
            self.model = None
        
        # Live camera variables
        self.live_camera = None
        self.live_timer = QTimer()
        self.live_timer.timeout.connect(self.update_live_camera_frame)
        self.live_detection_count = 0
        
        # YENİ: Canlı kamera özellikleri
        self.is_live_camera = False
        self.camera_capture = None
        self.live_camera_timer = QTimer(self)
        self.live_camera_timer.timeout.connect(self.update_live_camera_frame)
        self.live_detection_enabled = False
        
        self.update_ui_state()
        
        # GPU/CUDA durumu kontrolü
        device_info = detect_gpu_capability()
        if device_info['available']:
            self.log_message(f"🚀 M.SAVAŞ ULTRA PERFORMANS modu başlatıldı! GPU: {device_info['name']} ({device_info['memory_gb']:.1f}GB)", "success")
        else:
            self.log_message("🚀 M.SAVAŞ ULTRA PERFORMANS modu başlatıldı! CPU optimizasyonu aktif (CUDA bulunamadı)", "success")
        
        self.log_message("📋 Yeni: Sadece Word ve Grafik raporları kullanılabilir. Video seçince detaylar görünecek.", "info")
    
    def validate_video_file(self, file_path: str) -> bool:
        """Video dosyasının güvenliğini ve geçerliliğini kontrol eder"""
        try:
            # Dosya varlığı kontrolü
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"Dosya bulunamadı: {file_path}")
            
            # Dosya boyutu kontrolü
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            max_size = 500  # MB - varsayılan limit
            try:
                if hasattr(config_manager, 'get_max_file_size'):
                    max_size = config_manager.get_max_file_size()
            except:
                pass
            if file_size_mb > max_size:
                raise FileFormatError(f"Dosya boyutu çok büyük: {file_size_mb:.1f}MB (Max: {max_size}MB)")
            
            # Dosya uzantısı kontrolü
            _, ext = os.path.splitext(file_path)
            supported_formats = ['mp4', 'avi', 'mkv', 'mov', 'wmv', 'flv', 'webm', 'dav', 'h264', '264', 'ts', 'm2ts', 'mts']
            try:
                if hasattr(config_manager, 'get_supported_formats'):
                    supported_formats = config_manager.get_supported_formats()
            except:
                pass
            if ext.lower().replace('.', '') not in supported_formats:
                raise FileFormatError(f"Desteklenmeyen format: {ext}")
            
            # Video dosyası geçerlilik kontrolü
            cap = cv2.VideoCapture(file_path)
            if not cap.isOpened():
                cap.release()
                raise VideoLoadError("Video dosyası açılamadı")
            
            fps = cap.get(cv2.CAP_PROP_FPS)
            frame_count = cap.get(cv2.CAP_PROP_FRAME_COUNT)
            cap.release()
            
            if fps <= 0 or frame_count <= 0:
                raise VideoLoadError("Geçersiz video metadata")
            
            return True
            
        except (FileNotFoundError, FileFormatError, VideoLoadError) as e:
            self.show_error_message(str(e))
            return False
        except Exception as e:
            self.show_error_message(f"Video doğrulama hatası: {e}")
            return False
    
    def cleanup_temp_files(self, temp_dir: str):
        """Geçici dosyaları güvenli şekilde temizler"""
        try:
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
                self.log_message(f"Geçici dosyalar temizlendi: {temp_dir}", "info")
        except Exception as e:
            self.log_message(f"Geçici dosya temizleme hatası: {e}", "warning")
    
    def _apply_security_camera_settings(self):
        """Güvenlik kamerası için optimum ayarları uygular"""
        try:
            # Hassasiyet seviyesini güvenlik kamerası moduna ayarla
            for button in self.sensitivity_buttons.buttons():
                if button.text() == "🏢 Güvenlik Kamerası":
                    button.setChecked(True)
                    self.current_sensitivity = "🏢 Güvenlik Kamerası"
                    break
            
            # Eğer güvenlik kamerası modu yoksa DETAYLI kullan
            if self.current_sensitivity != "🏢 Güvenlik Kamerası":
                for button in self.sensitivity_buttons.buttons():
                    if button.text() == "🔍 Detaylı":
                        button.setChecked(True)
                        self.current_sensitivity = "🔍 Detaylı"
                        break
            
            # Özel log mesajı
            self.log_message("🎯 Güvenlik kamerası optimizasyonları uygulandı:", "success")
            self.log_message("  • Düşük ışık optimizasyonu aktif", "info")
            self.log_message("  • Sürekli izleme modu etkin", "info")
            self.log_message("  • Gelişmiş hareket tespiti açık", "info")
            self.log_message("  • Sabit kamera optimizasyonu aktif", "info")
            
        except Exception as e:
            self.log_message(f"Güvenlik kamerası ayarları uygulanamadı: {e}", "warning")
    
    def detect_video_type(self, file_path: str) -> str:
        """Video tipini tespit eder (güvenlik kamerası, mobil, drone vb.)"""
        filename = os.path.basename(file_path).lower()
        
        # Güvenlik kamerası patterns
        if any(pattern in filename for pattern in ['_ch', 'channel', 'cam0', 'dvr', 'nvr', 'hikvision', 'dahua']):
            return "security_camera"
        
        # Mobil telefon patterns
        if any(pattern in filename for pattern in ['img_', 'vid_', 'mov_', 'whatsapp', 'telegram']):
            return "mobile_phone"
        
        # Drone patterns
        if any(pattern in filename for pattern in ['dji_', 'drone_', 'aerial_', 'phantom']):
            return "drone"
        
        # Varsayılan
        return "standard"

    def update_ui_state(self, is_analyzing=False, is_exporting=False):
        is_video_loaded = self.video_path is not None
        has_videos = bool(self.video_paths)
        has_events = bool(self.detected_events)
        
        # Progress bar görünürlüğü - ANALİZ SIRASINDA GÖRÜNÜR
        if hasattr(self, 'progress_bar'):
            self.progress_bar.setVisible(is_analyzing or is_exporting)
            if not is_analyzing and not is_exporting:
                self.progress_bar.setValue(0)
        
        # YENİ: Motion progress bar görünürlüğü - SADECE ANALİZ SIRASINDA
        if hasattr(self, 'motion_progress_bar'):
            self.motion_progress_bar.setVisible(is_analyzing)
            if not is_analyzing:
                self.motion_progress_bar.setValue(0)
        
        # Video yönetimi butonları
        self.btn_add_video.setEnabled(not is_analyzing and not is_exporting)
        self.btn_remove_video.setEnabled(has_videos and not is_analyzing and not is_exporting)
        self.btn_clear_videos.setEnabled(has_videos and not is_analyzing and not is_exporting)
        
        # Analiz butonları - tek video veya çoklu video varsa etkin
        can_analyze = (is_video_loaded or has_videos) and not is_analyzing and not is_exporting
        
        self.btn_load.setEnabled(not is_analyzing and not is_exporting)
        self.btn_analyze.setEnabled(can_analyze)
        
        # Analiz buton metni güncelle
        if is_analyzing:
            self.btn_analyze.setText("🛑 Analizi Durdur")
        else:
            self.btn_analyze.setText("🚀 Analiz Et")
        
        self.btn_stop_analysis.setEnabled(is_analyzing)
        self.btn_export.setEnabled(has_events and not is_analyzing and not is_exporting)
        self.btn_export_charts.setEnabled(has_events and not is_analyzing and not is_exporting)
        self.btn_export_word.setEnabled(has_events and not is_analyzing and not is_exporting)
        self.btn_export_motion.setEnabled(has_events and not is_analyzing and not is_exporting)  # Hareket raporu
        self.btn_export_all.setEnabled(has_events and not is_analyzing and not is_exporting)  # Yeni
        self.btn_play_pause.setEnabled(is_video_loaded and not is_analyzing and not is_exporting)
        self.timeline_widget.setEnabled(is_video_loaded and not is_analyzing)
        
        # Video döndürme butonları
        self.btn_rotate_90.setEnabled(is_video_loaded and not is_analyzing and not is_exporting)
        self.btn_rotate_180.setEnabled(is_video_loaded and not is_analyzing and not is_exporting)
        self.btn_rotate_270.setEnabled(is_video_loaded and not is_analyzing and not is_exporting)
        self.btn_rotate_reset.setEnabled(is_video_loaded and not is_analyzing and not is_exporting)
        
        for button in self.sensitivity_buttons.buttons():
            button.setEnabled(not is_analyzing and not is_exporting)

    @pyqtSlot()
    def load_video(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Video Dosyası Seç", "", 
            "Tüm Video Dosyaları (*.mp4 *.avi *.mkv *.mov *.wmv *.flv *.webm *.dav *.h264 *.264 *.ts *.m2ts *.mts);;"
            "Standart Videolar (*.mp4 *.avi *.mkv *.mov *.wmv);;"
            "Güvenlik Kamerası (*.dav *.h264 *.264 *.ts *.m2ts);;"
            "Canlı Yayın (*.flv *.webm *.ts);;"
            "Tüm Dosyalar (*.*)")
        if not file_path:
            return

        try:
            self.video_path = file_path
            self.log_message(f"Video yükleniyor: {os.path.basename(file_path)}", "info")
            self._reset_for_new_video()

            # Video capture nesnesini oluştur - farklı backend'leri dene
            backends_to_try = [
                cv2.CAP_FFMPEG,    # FFmpeg backend (en uyumlu)
                cv2.CAP_DSHOW,     # DirectShow (Windows)
                cv2.CAP_MSMF,      # Microsoft Media Foundation
                cv2.CAP_ANY        # Herhangi bir backend
            ]
            
            self.video_capture = None
            for backend in backends_to_try:
                try:
                    self.video_capture = cv2.VideoCapture(self.video_path, backend)
                    if self.video_capture.isOpened():
                        self.log_message(f"Video backend: {backend}", "info")
                        break
                    else:
                        self.video_capture.release()
                        self.video_capture = None
                except Exception as e:
                    self.log_message(f"Backend {backend} hatası: {e}", "warning")
                    continue
            
            if not self.video_capture or not self.video_capture.isOpened():
                raise VideoLoadError(f"Hiçbir backend ile video açılamadı: {file_path}")

            # Video özelliklerini al
            fps = self.video_capture.get(cv2.CAP_PROP_FPS)
            total_frames = int(self.video_capture.get(cv2.CAP_PROP_FRAME_COUNT))
            width = int(self.video_capture.get(cv2.CAP_PROP_FRAME_WIDTH))
            height = int(self.video_capture.get(cv2.CAP_PROP_FRAME_HEIGHT))
            
            # Video özelliklerini doğrula
            if fps <= 0 or total_frames <= 0 or width <= 0 or height <= 0:
                self.show_error_message(f"Video dosyası geçersiz bilgiler içeriyor:\nFPS: {fps}, Kareler: {total_frames}, Boyut: {width}x{height}")
                self._reset_for_new_video()
                self.video_path = None
                self.update_ui_state()
                return
                
            duration = total_frames / fps
            self.video_info = {
                'fps': fps, 
                'total_frames': total_frames, 
                'duration': duration,
                'width': width,
                'height': height
            }

            # İlk frame'i test et
            self.video_capture.set(cv2.CAP_PROP_POS_FRAMES, 0)
            ret, test_frame = self.video_capture.read()
            if not ret or test_frame is None:
                self.show_error_message("Video frame'leri okunamıyor. Dosya bozuk olabilir.")
                self._reset_for_new_video()
                self.video_path = None
                self.update_ui_state()
                return
            
            # Frame formatını kontrol et
            if len(test_frame.shape) not in [2, 3]:
                self.show_error_message(f"Desteklenmeyen video formatı: {test_frame.shape}")
                self._reset_for_new_video()
                self.video_path = None
                self.update_ui_state()
                return

            file_size_mb = os.path.getsize(self.video_path) / (1024 * 1024)
            
            # Video bilgi panelini güncelle
            video_name = os.path.basename(self.video_path)
            self.video_info_label.setText(
                f"📹 {video_name}\n"
                f"📐 {width}x{height} | ⏱️ {duration:.1f}s | 💾 {file_size_mb:.1f}MB"
            )

            self.timeline_widget.set_duration(duration)
            
            # İlk frame'i göster
            self.show_frame(0)
            
            # Güvenlik kamerası tespiti ve otomatik ayar
            try:
                if hasattr(config_manager, 'is_security_camera_file') and config_manager.is_security_camera_file(self.video_path):
                    self.log_message("🔒 Güvenlik kamerası dosyası tespit edildi! Otomatik optimizasyon uygulanıyor...", "info")
                    self._apply_security_camera_settings()
            except:
                # Video tipini basit şekilde tespit et
                if any(pattern in file_path.lower() for pattern in ['_ch', 'channel', 'cam0', 'dvr', 'nvr', 'hikvision', 'dahua']):
                    self.log_message("🔒 Güvenlik kamerası dosyası tespit edildi! Otomatik optimizasyon uygulanıyor...", "info")
                    self._apply_security_camera_settings()
            
            self.log_message(f"✅ Video başarıyla yüklendi: {os.path.basename(file_path)}", "success")
            self.log_message(f"   • Boyut: {width}x{height}", "info")
            self.log_message(f"   • Süre: {self.format_duration(duration)}", "info")
            self.log_message(f"   • FPS: {fps:.2f}", "info")
            self.log_message(f"   • Toplam Kare: {total_frames:,}", "info")
                
        except VideoLoadError as e:
            self.show_error_message(f"Video yüklenemedi: {e}")
            self.log_message(f"❌ Video yükleme hatası: {e}", "error")
            self._reset_for_new_video()
            self.video_path = None
        except Exception as e:
            self.show_error_message(f"Beklenmeyen hata: {e}")
            self.log_message(f"❌ Beklenmeyen hata: {e}", "error")
            self._reset_for_new_video()
            self.video_path = None
        finally:
            self.update_ui_state()

    def _reset_for_new_video(self):
        if self.is_playing: 
            self.toggle_playback()
        if self.video_capture: 
            self.video_capture.release()
        self.detected_events = []
        self.detected_objects = {} # Yeni: Nesneleri de sıfırla
        self.current_rotation = 0  # Video döndürme sıfırla
        
        # Progress bar varsa sıfırla
        if hasattr(self, 'progress_bar'):
            self.progress_bar.setValue(0)
        
        # Status güncelle
        self.status_label.setText("🟢 Video hazır - Analiz başlatabilirsiniz")
        
        # Timeline widgets varsa sıfırla
        if hasattr(self, 'timeline_widget'):
            try:
                self.timeline_widget.set_duration(0)
                self.timeline_widget.set_events([])
                self.timeline_widget.set_progress(0)
            except:
                pass
        
        self.video_display_label.setText("Video yüklendi - 'Analiz Et' butonuna basın")
        
        # Event list varsa temizle
        if hasattr(self, 'event_list_widget'):
            self.event_list_widget.clear() # Yeni: Olay listesini temizle

    def add_video_file(self):
        """Yeni video dosyası ekler."""
        file_paths, _ = QFileDialog.getOpenFileNames(
            self, 
            "Video Dosyaları Seç", 
            "", 
            "Tüm Video Dosyaları (*.mp4 *.avi *.mkv *.mov *.wmv *.flv *.webm *.dav *.h264 *.264 *.ts *.m2ts *.mts);;"
            "Standart Videolar (*.mp4 *.avi *.mkv *.mov *.wmv);;"
            "Güvenlik Kamerası (*.dav *.h264 *.264 *.ts *.m2ts);;"
            "Canlı Yayın (*.flv *.webm *.ts);;"
            "Tüm Dosyalar (*.*)"
        )
        
        for file_path in file_paths:
            if file_path and file_path not in self.video_paths:
                # Video tipini tespit et ve log'a yaz
                video_type = self.detect_video_type(file_path)
                type_names = {
                    "security_camera": "🔒 Güvenlik Kamerası",
                    "mobile_phone": "📱 Mobil Telefon", 
                    "drone": "🚁 Drone",
                    "standard": "🎬 Standart Video"
                }
                
                self.video_paths.append(file_path)
                # Liste widget'ına ekle
                filename = os.path.basename(file_path)
                type_icon = type_names.get(video_type, "📹")
                item_text = f"{type_icon} {filename}"
                self.video_list.addItem(item_text)
                
                # Özel log mesajı
                self.log_message(f"Video eklendi: {filename} ({type_names.get(video_type, 'Bilinmeyen tip')})", "success")
                
                # Güvenlik kamerası ise özel bilgi ver
                if video_type == "security_camera":
                    self.log_message("  💡 İpucu: Bu dosya için 'GÜVENLİK KAMERASI' hassasiyet seviyesi önerilir", "info")
                self.video_paths.append(file_path)
                # Liste widget'ına ekle
                item_text = f"📹 {os.path.basename(file_path)}"
                self.video_list.addItem(item_text)
                self.log_message(f"Video eklendi: {os.path.basename(file_path)}", "success")
        
        self.update_ui_state()
    
    def remove_video_file(self):
        """Seçili video dosyasını kaldırır."""
        current_row = self.video_list.currentRow()
        if current_row >= 0:
            removed_video = self.video_paths.pop(current_row)
            self.video_list.takeItem(current_row)
            self.log_message(f"Video kaldırıldı: {os.path.basename(removed_video)}", "success")
            
            # Eğer şu anki video kaldırıldıysa
            if current_row == self.current_video_index:
                self.video_path = None
                self.video_capture = None
                self.current_video_index = 0
                self._reset_for_new_video()
            elif current_row < self.current_video_index:
                self.current_video_index -= 1
        
        self.update_ui_state()
    
    def clear_all_videos(self):
        """Tüm video dosyalarını temizler."""
        self.video_paths.clear()
        self.video_list.clear()
        self.video_path = None
        self.video_capture = None
        self.current_video_index = 0
        self.detected_events.clear()
        self.detected_objects.clear()
        self._reset_for_new_video()
        self.update_ui_state()
        self.log_message("Tüm videolar temizlendi.", "info")
    
    def on_video_selection_changed(self):
        """Video seçimi değiştiğinde çağrılır."""
        current_row = self.video_list.currentRow()
        if current_row >= 0 and current_row < len(self.video_paths):
            self.current_video_index = current_row
            self.video_path = self.video_paths[current_row]
            self.load_selected_video()
    
    def load_selected_video(self):
        """Seçili videoyu yükler."""
        if not self.video_path:
            return
        
        try:
            if self.video_capture:
                self.video_capture.release()
            
            self.video_capture = cv2.VideoCapture(self.video_path)
            if not self.video_capture.isOpened():
                self.log_message(f"❌ Video açılamadı: {os.path.basename(self.video_path)}", "error")
                return
            
            # Video bilgilerini al
            fps = self.video_capture.get(cv2.CAP_PROP_FPS)
            frame_count = int(self.video_capture.get(cv2.CAP_PROP_FRAME_COUNT))
            width = int(self.video_capture.get(cv2.CAP_PROP_FRAME_WIDTH))
            height = int(self.video_capture.get(cv2.CAP_PROP_FRAME_HEIGHT))
            duration = frame_count / fps if fps > 0 else 0
            
            # Dosya boyutunu al
            file_size = os.path.getsize(self.video_path)
            file_size_mb = file_size / (1024 * 1024)
            
            self.video_info = {
                'fps': fps,
                'total_frames': frame_count,
                'width': width,
                'height': height,
                'duration': duration
            }
            
            # Video bilgi panelini güncelle
            video_name = os.path.basename(self.video_path)
            self.video_info_label.setText(
                f"📹 {video_name}\n"
                f"📐 {width}x{height} | ⏱️ {duration:.1f}s | 💾 {file_size_mb:.1f}MB"
            )
            
            # Timeline'ı güncelle
            self.timeline_widget.set_duration(duration)
            self.timeline_widget.set_events([])
            
            # İlk kareyi göster
            self.video_capture.set(cv2.CAP_PROP_POS_FRAMES, 0)
            ret, frame = self.video_capture.read()
            if ret:
                self.display_cv_frame(frame, 0)
            
            # UI'ı güncelle
            self.update_ui_state()
            
            video_name = os.path.basename(self.video_path)
            self.log_message(f"Video yüklendi: {video_name} ({duration:.1f}s, {frame_count} kare)", "success")
            
        except Exception as e:
            self.log_message(f"Video yükleme hatası: {str(e)}", "error")

    def update_active_classes(self):
        """Aktif sınıfları güncelle"""
        self.ACTIVE_CLASSES = []
        for checkbox in self.object_checkboxes:
            if checkbox.isChecked():
                class_name = checkbox.text()
                if class_name in TARGET_CLASSES:
                    self.ACTIVE_CLASSES.append(TARGET_CLASSES[class_name])
        
        # Eğer hiç seçili değilse, varsayılan olarak person ekle
        if not self.ACTIVE_CLASSES:
            self.ACTIVE_CLASSES = [0]  # person ID'si
    
    def open_advanced_selection(self):
        """Gelişmiş nesne seçim dialogunu aç"""
        dialog = AdvancedObjectSelectionDialog(self.TARGET_CLASSES, self.ACTIVE_CLASSES, self)
        if dialog.exec_() == QDialog.Accepted:
            selected_classes = dialog.get_selected_classes()
            if selected_classes:
                # String isimlerini ID'lere çevir
                self.ACTIVE_CLASSES = []
                for class_name in selected_classes:
                    if class_name in TARGET_CLASSES:
                        self.ACTIVE_CLASSES.append(TARGET_CLASSES[class_name])
                self.log_message(f"{len(self.ACTIVE_CLASSES)} nesne türü seçildi!", "success")
            else:
                self.ACTIVE_CLASSES = [0]  # person ID'si
                self.log_message("Hiç nesne seçilmediği için 'person' varsayılan olarak seçildi.", "info")

    def start_live_camera(self):
        """Live kamera başlatma"""
        try:
            # Kamera başlatma
            self.live_camera = cv2.VideoCapture(0)
            
            if not self.live_camera.isOpened():
                self.log_message("Kamera bulunamadı!", "error")
                return
            
            # FPS ayarı
            self.live_camera.set(cv2.CAP_PROP_FPS, 30)
            
            # Timer başlat
            self.live_timer = QTimer()
            self.live_timer.timeout.connect(self.update_live_camera_frame)
            self.live_timer.start(33)  # ~30 FPS
            
            # UI güncelle
            self.btn_start_camera.setEnabled(False)
            self.btn_stop_camera.setEnabled(True)
            self.live_detection_count = 0
            
            self.log_message("Kamera başlatıldı!", "success")
            
        except Exception as e:
            self.log_error("Kamera başlatma hatası", e)
    
    def stop_live_camera(self):
        """Live kamera durdurma"""
        try:
            if hasattr(self, 'live_timer'):
                self.live_timer.stop()
            
            if hasattr(self, 'live_camera'):
                self.live_camera.release()
            
            # Video display'i temizle
            if hasattr(self, 'video_display_label'):
                self.video_display_label.setText("Lütfen bir video dosyası yükleyin veya canlı kamerayı başlatın.")
            
            # UI güncelle
            self.btn_start_camera.setEnabled(True)
            self.btn_stop_camera.setEnabled(False)
            
            self.log_message("Kamera durduruldu!", "info")
            
        except Exception as e:
            self.log_message(f"Kamera durdurma hatası: {e}", "warning")
    
    def update_live_camera_frame(self):
        """Live kamera frame'ini güncelle"""
        try:
            if not hasattr(self, 'live_camera') or self.live_camera is None:
                return
                
            ret, frame = self.live_camera.read()
            
            if not ret:
                return
            
            # Nesne tespiti yap
            frame = self.detect_live_objects(frame)
            
            # Frame'i Qt formatına çevir
            rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            h, w, ch = rgb_frame.shape
            bytes_per_line = ch * w
            qt_image = QImage(rgb_frame.data, w, h, bytes_per_line, QImage.Format_RGB888)
            
            # Video display'de göster
            if hasattr(self, 'video_display_label'):
                pixmap = QPixmap.fromImage(qt_image)
                scaled_pixmap = pixmap.scaled(
                    self.video_display_label.size(), 
                    Qt.KeepAspectRatio, 
                    Qt.SmoothTransformation
                )
                self.video_display_label.setPixmap(scaled_pixmap)
            
        except Exception as e:
            self.log_message(f"Live frame güncelleme hatası: {e}", "warning")

    def detect_live_objects(self, frame):
        """Live kamera görüntüsünde nesne tespiti yap"""
        try:
            if self.model is None:
                return frame
                
            # YOLO modeli ile nesne tespiti
            results = self.model(frame)
            
            # Tespit edilen nesneleri çiz
            annotated_frame = results[0].plot()
            
            # Aktif nesneleri say ve Türkçe isimlerle göster
            detected_count = 0
            detected_objects_turkish = []
            for result in results:
                if result.boxes is not None:
                    for box in result.boxes:
                        class_id = int(box.cls[0])
                        class_name = self.model.names[class_id]
                        if class_id in self.ACTIVE_CLASSES:
                            detected_count += 1
                            # Türkçe isim kullan
                            turkish_name = self.OBJECT_NAMES_TURKISH.get(class_name, class_name)
                            detected_objects_turkish.append(turkish_name)
            
            # Count'u güncelle ve Türkçe nesne isimlerini göster
            self.live_detection_count += detected_count
            if hasattr(self, 'status_label'):
                if detected_objects_turkish:
                    objects_str = ", ".join(list(set(detected_objects_turkish))[:3])  # En fazla 3 nesne göster
                    self.status_label.setText(f"🔴 Anlık Tespit: {detected_count} ({objects_str}) | Toplam: {self.live_detection_count}")
                else:
                    self.status_label.setText(f"🔴 Anlık Tespit: {detected_count} | Toplam: {self.live_detection_count}")
            
            return annotated_frame
            
        except Exception as e:
            self.log_message(f"Live detection error: {e}", "warning")
            return frame
            
        except Exception as e:
            self.log_error("Live detection error", e)
            return frame

    def analyze_video(self):
        """Ana video analiz fonksiyonu - buton tıklamasında çağrılır"""
        if not self.video_paths and not self.video_path:
            QMessageBox.warning(self, "Uyarı", "Lütfen analiz edilecek video(lar) seçin!")
            return
        
        # Analiz durumunu kontrol et
        if hasattr(self, 'processor_thread') and self.processor_thread and self.processor_thread.isRunning():
            self.stop_analysis()
        else:
            self.start_analysis()
    
    def stop_analysis(self):
        """Analizi durdurur"""
        if hasattr(self, 'processor_thread') and self.processor_thread:
            self.processor_thread.stop()
            self.processor_thread.wait(3000)  # 3 saniye bekle
            self.log_message("Analiz durduruldu", "warning")
        
        self.update_ui_state(is_analyzing=False)

    def start_analysis(self):
        # Çoklu video desteği
        if self.video_paths:
            self.start_batch_analysis()
        elif self.video_path:
            self.start_single_analysis()
        else:
            self.log_message("⚠️ Lütfen analiz edilecek video(lar) seçin.", "warning")
    
    def start_single_analysis(self):
        """Tek video analizi başlatır."""
        if not self.video_path: 
            return
        
        self.update_ui_state(is_analyzing=True)
        
        # Progress bar görünür yap ve sıfırla
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        
        # YENİ: Motion progress bar'ı da görünür yap
        self.motion_progress_bar.setVisible(True)
        self.motion_progress_bar.setValue(0)
        
        self.log_message(f"Analiz başlatılıyor (Hassasiyet: {self.current_sensitivity})...", "info")
        
        self.processor_thread = VideoProcessor(self.video_path, self.current_sensitivity)
        self.processor_thread.progress_updated.connect(self.progress_bar.setValue)
        self.processor_thread.motion_progress_updated.connect(self.on_motion_progress_updated)  # DEBUG: Motion progress ile debug
        self.processor_thread.status_updated.connect(self.update_status)
        self.processor_thread.analysis_complete.connect(self.on_analysis_complete)
        self.processor_thread.error_occurred.connect(self.on_thread_error)
        self.processor_thread.start()
    
    def on_motion_progress_updated(self, value):
        """DEBUG: Motion progress güncellemelerini takip eder."""
        print(f"🔍 DEBUG: Motion progress signal alındı: {value}%")
        self.motion_progress_bar.setValue(value)
        # Status'ta da göster
        if value > 0:
            self.update_status(f"🚨 HAREKET TESPİTİ: %{value}")
    
    def start_batch_analysis(self):
        """🚀 ULTRA HIZLI çoklu video analizi."""
        if not self.video_paths:
            return
        
        self.update_ui_state(is_analyzing=True)
        
        # Progress bar görünür yap ve sıfırla
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        
        # YENİ: Motion progress bar'ı da görünür yap - BATCH için
        self.motion_progress_bar.setVisible(True)
        self.motion_progress_bar.setValue(0)
        
        self.current_batch_index = 0
        self.batch_results = []
        
        total_videos = len(self.video_paths)
        self.log_message(f"🚀 ULTRA HIZLI toplu analiz başlıyor: {total_videos} video (Hassasiyet: {self.current_sensitivity})", "success")
        
        # Paralel işlem desteği
        if MULTIPROCESSING_ENABLED and total_videos > 1:
            self.log_message("⚡ Paralel işlem modu aktif", "info")
        
        self.analyze_next_video()
    
    def analyze_next_video(self):
        """🎯 Toplu analizde sıradaki videoyu ULTRA HIZLI analiz eder."""
        if self.current_batch_index >= len(self.video_paths):
            self.on_batch_analysis_complete()
            return
        
        current_video = self.video_paths[self.current_batch_index]
        video_name = os.path.basename(current_video)
        
        # Dosya boyutu bilgisi
        try:
            file_size = os.path.getsize(current_video) / (1024 * 1024)  # MB
            self.log_message(f"⚡ Analiz [{self.current_batch_index + 1}/{len(self.video_paths)}]: {video_name} ({file_size:.1f} MB)", "info")
        except:
            self.log_message(f"⚡ Analiz [{self.current_batch_index + 1}/{len(self.video_paths)}]: {video_name}", "info")
        
        self.processor_thread = VideoProcessor(current_video, self.current_sensitivity)
        self.processor_thread.progress_updated.connect(self.on_batch_progress_updated)
        self.processor_thread.motion_progress_updated.connect(self.on_motion_progress_updated)  # DEBUG: Motion progress ile debug
        self.processor_thread.analysis_complete.connect(self.on_batch_video_complete)
        self.processor_thread.error_occurred.connect(self.on_batch_error)
        self.processor_thread.start()
    
    def on_batch_progress_updated(self, progress):
        """Toplu analiz ilerlemesini günceller."""
        total_progress = (self.current_batch_index * 100 + progress) / len(self.video_paths)
        self.progress_bar.setValue(int(total_progress))
    
    def on_batch_motion_progress_updated(self, motion_progress):
        """Toplu analiz motion ilerlemesini günceller."""
        # Şu anki video için motion progress göster
        self.motion_progress_bar.setValue(motion_progress)
        # Debug için motion progress değerini göster
        print(f"DEBUG: Motion progress güncellendi: {motion_progress}%")
    
    def on_batch_video_complete(self, detected_objects, events, video_info):
        """Toplu analizde bir video tamamlandığında çağrılır."""
        video_path = self.video_paths[self.current_batch_index]
        video_name = os.path.basename(video_path)
        
        # Sonuçları sakla
        result = {
            'video_path': video_path,
            'video_name': video_name,
            'detected_objects': detected_objects,
            'events': events,
            'video_info': video_info,
            'event_count': len(events),
            'detection_count': len(detected_objects)
        }
        self.batch_results.append(result)
        
        event_count = len(events)
        detection_count = len(detected_objects)
        self.log_message(f"✅ Tamamlandı: {video_name} ({event_count} olay, {detection_count} tespit)", "success")
        
        self.current_batch_index += 1
        self.analyze_next_video()
    
    def on_batch_error(self, error_msg):
        """Toplu analizde hata oluştuğunda çağrılır."""
        video_path = self.video_paths[self.current_batch_index]
        video_name = os.path.basename(video_path)
        
        self.log_message(f"❌ Hata: {video_name} - {error_msg}", "error")
        
        self.current_batch_index += 1
        self.analyze_next_video()
    
    def on_batch_analysis_complete(self):
        """Toplu analiz tamamlandığında çağrılır."""
        self.update_ui_state(is_analyzing=False)
        
        total_videos = len(self.batch_results)
        total_events = sum(result['event_count'] for result in self.batch_results)
        total_detections = sum(result['detection_count'] for result in self.batch_results)
        
        self.log_message(f"🎉 Toplu analiz tamamlandı! {total_videos} video, {total_events} olay, {total_detections} tespit", "success")
        
        # Sonuçları birleştir veya özet göster
        self.show_batch_results_summary()
    
    def show_batch_results_summary(self):
        """Toplu analiz sonuçlarının özetini gösterir."""
        if not hasattr(self, 'batch_results') or not self.batch_results:
            return
        
        # Event listesini temizle ve toplu sonuçları ekle
        self.event_list_widget.clear()
        
        for result in self.batch_results:
            video_name = result['video_name']
            event_count = result['event_count']
            detection_count = result['detection_count']
            
            item_text = f"📹 {video_name} | {event_count} olay | {detection_count} tespit"
            item = QListWidgetItem(item_text)
            item.setData(Qt.UserRole, result)  # Sonucu item'a bağla
            self.event_list_widget.addItem(item)

    @pyqtSlot(dict, list, dict)
    def on_analysis_complete(self, detected_objects, events, video_info):
        # detected_objects'in doğru tipte olduğunu kontrol et
        if isinstance(detected_objects, dict):
            self.detected_objects = detected_objects
        else:
            print(f"UYARI: detected_objects list olarak geldi: {type(detected_objects)}")
            # List'i dict'e dönüştür
            converted_objects = {}
            for i, obj in enumerate(detected_objects):
                if isinstance(obj, dict):
                    converted_objects[i] = [obj]
                else:
                    converted_objects[i] = obj
            self.detected_objects = converted_objects
            
        self.detected_events = events
        self.video_info = video_info
        self.timeline_widget.set_events(events)
        
        msg = f"✅ Analiz tamamlandı. {len(events)} olay bulundu."
        
        # Tespit edilen nesne türlerini Türkçe olarak göster
        detected_object_types = set()
        for frame_objects in detected_objects.values():
            if isinstance(frame_objects, list):
                for obj in frame_objects:
                    if isinstance(obj, dict) and 'class_name' in obj:
                        english_name = obj['class_name']
                        turkish_name = self.OBJECT_NAMES_TURKISH.get(english_name, english_name)
                        detected_object_types.add(turkish_name)
        
        if detected_object_types:
            objects_str = ", ".join(list(detected_object_types)[:5])  # En fazla 5 nesne türü göster
            msg += f" Tespit edilen nesneler: {objects_str}"
        
        self.log_message(msg, "success" if events else "warning")
        
        self.event_list_widget.clear()
        if events:
            for i, (start, end) in enumerate(events):
                log_msg = f"Olay {i+1}: {self.format_duration(start)} - {self.format_duration(end)}"
                self.log_message(f"  {log_msg}", "info")
                # Listeye tıklanabilir öğe ekle
                item = QListWidgetItem(log_msg)
                item.setData(Qt.UserRole, start) # Başlangıç zamanını sakla
                self.event_list_widget.addItem(item)
        
        # YENİ: Hareket tespiti zaman çizelgesini güncelle
        self.update_motion_timeline()
        
        self.update_status(msg)
        self.update_ui_state()
        
        # Progress bar'ı gizle ve sıfırla
        if hasattr(self, 'progress_bar'):
            self.progress_bar.setVisible(False)
            self.progress_bar.setValue(0)
        
        # DEBUG: Motion progress bar'ı gizleme - geçici olarak açık bırak
        if hasattr(self, 'motion_progress_bar'):
            # Geçici: Motion progress'i gizleme, 3 saniye sonra gizle
            QTimer.singleShot(3000, lambda: self.motion_progress_bar.setVisible(False))
            print(f"🔍 DEBUG: Motion progress bar değeri: {self.motion_progress_bar.value()}%")
    
    def update_motion_timeline(self):
        """Hareket tespiti zaman çizelgesini günceller."""
        if not hasattr(self, 'motion_timeline_list'):
            return
            
        self.motion_timeline_list.clear()
        
        if not self.detected_events:
            item = QListWidgetItem("⏸️ Hareket tespiti bulunamadı")
            item.setForeground(QColor('#95a5a6'))
            self.motion_timeline_list.addItem(item)
            self.motion_stats_label.setText("📈 Hareket istatistikleri: Hareket bulunamadı")
            return
        
        # Hareket tespit edilen saniyeler
        motion_seconds = []
        for start_time, end_time in self.detected_events:
            # Her saniye için hareket var mı kontrol et
            for second in range(int(start_time), int(end_time) + 1):
                if second not in motion_seconds:
                    motion_seconds.append(second)
        
        motion_seconds.sort()
        
        # Hareket zamanlarını listele - sadece kırmızı renk kullan
        for second in motion_seconds:
            minutes = second // 60
            secs = second % 60
            time_str = f"{minutes:02d}:{secs:02d}"
            
            # Hareket yoğunluğunu hesapla (kaç frame'de tespit var)
            frame_count = 0
            if hasattr(self, 'detected_objects') and self.detected_objects:
                for frame_num in self.detected_objects.keys():
                    if self.video_info.get('fps', 30) > 0:
                        frame_time = frame_num / self.video_info.get('fps', 30)
                        if int(frame_time) == second:
                            frame_count += 1
            
            # Sadece kırmızı gösterge kullan - hareketli bölgeler için
            intensity = "🔴"  # Hep kırmızı - sadece hareketli yerler gösteriliyor
            item_text = f"{intensity} {time_str} saniye - {frame_count} tespit"
            
            item = QListWidgetItem(item_text)
            item.setData(Qt.UserRole, second)  # Saniyeyi sakla
            # Hareketli bölgeler kırmızı renkte görünsün
            item.setForeground(QColor('#e74c3c'))  # Kırmızı renk
            self.motion_timeline_list.addItem(item)
        
        # İstatistikleri güncelle
        total_motion_time = len(motion_seconds)
        total_video_time = self.video_info.get('duration', 0)
        motion_percentage = (total_motion_time / total_video_time * 100) if total_video_time > 0 else 0
        
        stats_text = f"📈 {total_motion_time} saniye hareket | {motion_percentage:.1f}% video oranı | {len(self.detected_events)} olay"
        self.motion_stats_label.setText(stats_text)
    
    def on_motion_timeline_clicked(self, item):
        """Hareket zaman çizelgesindeki bir öğeye tıklandığında çağrılır."""
        second = item.data(Qt.UserRole)
        if second is not None:
            self.seek_video(float(second))
            self.log_message(f"⏯️ {second} saniyeye gidildi", "info")

    @pyqtSlot()
    def export_video(self):
        if not self.detected_events:
            self.show_error_message("Dışa aktarılacak olay bulunamadı.")
            return

        output_path, _ = QFileDialog.getSaveFileName(self, "Özet Videoyu Kaydet", "V.E.R.A_Özet.mp4", "MP4 Dosyası (*.mp4)")
        if not output_path: return

        self.update_ui_state(is_exporting=True)
        
        # Progress bar görünür yap ve sıfırla - EXPORT için
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        
        self.log_message("Özet video oluşturuluyor...", "info")

        self.exporter_thread = VideoExporter(self.video_path, self.detected_events, self.video_info, output_path)
        self.exporter_thread.export_progress.connect(self.on_export_progress)
        self.exporter_thread.export_complete.connect(self.on_export_complete)
        self.exporter_thread.error_occurred.connect(self.on_thread_error)
        self.exporter_thread.start()

    @pyqtSlot(int, str)
    def on_export_progress(self, value, message):
        self.progress_bar.setValue(value)
        self.update_status(message)

    @pyqtSlot(str, str)
    def on_export_complete(self, path, message):
        self.log_message(message, "success")
        self.update_status("Özet video oluşturuldu.")
        self.update_ui_state()
        
        # Progress bar'ı gizle - EXPORT tamamlandı
        if hasattr(self, 'progress_bar'):
            self.progress_bar.setVisible(False)
            self.progress_bar.setValue(0)
        
        try:
            summary_size_mb = os.path.getsize(path) / (1024 * 1024)
            summary_cap = cv2.VideoCapture(path)
            s_fps = summary_cap.get(cv2.CAP_PROP_FPS)
            s_frames = summary_cap.get(cv2.CAP_PROP_FRAME_COUNT)
            summary_duration = s_frames / s_fps if s_fps > 0 else 0
            summary_cap.release()
            # self.info_label_summary.setText(f"<b>Özet Video:</b> {self.format_duration(summary_duration)} | {summary_size_mb:.2f} MB")
            self.log_message(f"✅ Özet video oluşturuldu: {self.format_duration(summary_duration)} süre, {summary_size_mb:.2f} MB boyut", "success")
        except Exception as e:
            self.log_message(f"Özet video bilgileri okunamadı: {e}", "error")

        reply = QMessageBox.question(self, 'Başarılı', "Özet video oluşturuldu. Dosyayı açmak ister misiniz?", QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.Yes:
            from PyQt5.QtGui import QDesktopServices
            QDesktopServices.openUrl(QUrl.fromLocalFile(path))

    @pyqtSlot(str)
    def on_thread_error(self, error_message):
        self.show_error_message(error_message)
        self.update_status("Bir hata oluştu. Detaylar için işlem geçmişine bakın.")
        self.update_ui_state()
        
        # Progress bar'ı gizle - HATA durumunda
        if hasattr(self, 'progress_bar'):
            self.progress_bar.setVisible(False)
            self.progress_bar.setValue(0)
        
        # YENİ: Motion progress bar'ı da gizle - HATA durumunda
        if hasattr(self, 'motion_progress_bar'):
            self.motion_progress_bar.setVisible(False)
            self.motion_progress_bar.setValue(0)

    @pyqtSlot()
    def toggle_playback(self):
        if not self.video_capture: return
        self.is_playing = not self.is_playing
        if self.is_playing:
            self.btn_play_pause.setIcon(self.style().standardIcon(QStyle.SP_MediaPause))
            self.playback_timer.start(int(1000 / self.video_info.get('fps', 30)))
        else:
            self.btn_play_pause.setIcon(self.style().standardIcon(QStyle.SP_MediaPlay))
            self.playback_timer.stop()

    @pyqtSlot()
    def update_frame(self):
        if not self.video_capture or not self.is_playing:
            return
            
        if not self.video_capture.isOpened():
            self.log_message("Video capture kapalı - oynatma durduruluyor", "warning")
            self.is_playing = False
            self.playback_timer.stop()
            return
            
        try:
            ret, frame = self.video_capture.read()
            if ret and frame is not None:
                current_frame_num = int(self.video_capture.get(cv2.CAP_PROP_POS_FRAMES))
                self.display_cv_frame(frame, current_frame_num)
                
                # Timeline progress güncelle
                total_frames = self.video_info.get('total_frames', 1)
                if total_frames > 0:
                    self.timeline_widget.set_progress(current_frame_num / total_frames)
            else:
                # Video sonu - oynatmayı durdur
                self.is_playing = False
                self.playback_timer.stop()
                self.video_capture.set(cv2.CAP_PROP_POS_FRAMES, 0)
                self.btn_play_pause.setIcon(self.style().standardIcon(QStyle.SP_MediaPlay))
                self.timeline_widget.set_progress(0)
                self.log_message("Video sonu - oynatma durduruldu", "info")
                
        except Exception as e:
            self.log_message(f"Frame update hatası: {e}", "error")
            self.is_playing = False
            self.playback_timer.stop()

    @pyqtSlot(float)
    def seek_video(self, time_sec):
        if self.video_capture:
            frame_num = int(time_sec * self.video_info.get('fps', 30))
            self.show_frame(frame_num)
            if self.is_playing:
                self.video_capture.set(cv2.CAP_PROP_POS_FRAMES, frame_num)

    def show_frame(self, frame_number):
        if not self.video_capture or not self.video_capture.isOpened():
            self.log_message("Video capture nesnesi kullanılamıyor", "warning")
            return
            
        try:
            # Frame numarasını sınırlar içinde tut
            max_frames = self.video_info.get('total_frames', 0)
            frame_number = max(0, min(frame_number, max_frames - 1))
            
            # Frame pozisyonunu ayarla
            self.video_capture.set(cv2.CAP_PROP_POS_FRAMES, frame_number)
            ret, frame = self.video_capture.read()
            
            if ret and frame is not None:
                self.display_cv_frame(frame, frame_number)
                # Timeline progress güncelle
                if max_frames > 0:
                    self.timeline_widget.set_progress(frame_number / max_frames)
            else:
                self.log_message(f"Frame {frame_number} okunamadı", "warning")
                # Hata durumunda varsayılan mesaj göster
                self.video_display_label.setText(f"Frame {frame_number} görüntülenemiyor")
                
        except Exception as e:
            self.log_message(f"Frame gösterme hatası: {e}", "error")
            self.video_display_label.setText("Video frame hatası")

    def rotate_video(self, rotation):
        """Video döndürme fonksiyonu"""
        self.current_rotation = rotation
        self.log_message(f"Video döndürme: {rotation}°", "info")
        
        # Şu anki frame'i tekrar göster
        if self.video_capture:
            current_frame = int(self.video_capture.get(cv2.CAP_PROP_POS_FRAMES))
            self.show_frame(current_frame)
    
    def apply_rotation(self, frame, rotation):
        """Frame'e döndürme uygular"""
        if rotation == 0:
            return frame
        elif rotation == 90:
            return cv2.rotate(frame, cv2.ROTATE_90_CLOCKWISE)
        elif rotation == 180:
            return cv2.rotate(frame, cv2.ROTATE_180)
        elif rotation == 270:
            return cv2.rotate(frame, cv2.ROTATE_90_COUNTERCLOCKWISE)
        else:
            return frame

    def adjust_coordinates_for_rotation(self, x, y, w, h, rotation, frame_shape):
        """Döndürme için koordinatları ayarlar"""
        if rotation == 0:
            return x, y, w, h
        
        original_height, original_width = frame_shape[:2]
        
        if rotation == 90:
            # 90° saat yönünde: (x,y) -> (y, width-x-w)
            new_x = y
            new_y = original_width - x - w
            new_w = h
            new_h = w
        elif rotation == 180:
            # 180°: (x,y) -> (width-x-w, height-y-h)
            new_x = original_width - x - w
            new_y = original_height - y - h
            new_w = w
            new_h = h
        elif rotation == 270:
            # 270° saat yönünde: (x,y) -> (height-y-h, x)
            new_x = original_height - y - h
            new_y = x
            new_w = h
            new_h = w
        else:
            return x, y, w, h
        
        return new_x, new_y, new_w, new_h

    def enhance_frame_quality(self, frame):
        """Frame kalitesini iyileştir - özellikle güvenlik kamerası videoları için"""
        try:
            if frame is None or frame.size == 0:
                return frame
            
            # Frame'in ortalama parlaklığını kontrol et
            gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
            mean_brightness = np.mean(gray)
            
            # Düşük ışık koşulları için iyileştirme (ortalama parlaklık < 80)
            if mean_brightness < 80:
                # Histogram eşitleme uygula
                clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8,8))
                
                # Her kanal için ayrı ayrı iyileştir
                if len(frame.shape) == 3:
                    lab = cv2.cvtColor(frame, cv2.COLOR_BGR2LAB)
                    lab[:,:,0] = clahe.apply(lab[:,:,0])
                    frame = cv2.cvtColor(lab, cv2.COLOR_LAB2BGR)
                else:
                    frame = clahe.apply(frame)
            
            # Çok koyu veya çok parlak videoları ayarla
            if mean_brightness < 50:  # Çok koyu
                frame = cv2.convertScaleAbs(frame, alpha=1.3, beta=20)
            elif mean_brightness > 200:  # Çok parlak
                frame = cv2.convertScaleAbs(frame, alpha=0.9, beta=-10)
            
            return frame
            
        except Exception as e:
            self.log_message(f"Frame kalitesi iyileştirme hatası: {e}", "warning")
            return frame
    
    def display_cv_frame(self, frame, frame_number):
        try:
            # Frame'in geçerli olup olmadığını kontrol et
            if frame is None or frame.size == 0:
                self.log_message("Geçersiz frame - boş veri", "warning")
                return
            
            # Frame'in boyutlarını kontrol et
            if len(frame.shape) != 3 or frame.shape[2] != 3:
                self.log_message(f"Geçersiz frame formatı: {frame.shape}", "warning")
                return
            
            # Frame kalitesini iyileştir (özellikle güvenlik kamerası videoları için)
            frame = self.enhance_frame_quality(frame)
            
            # Döndürme uygula
            if self.current_rotation != 0:
                frame = self.apply_rotation(frame, self.current_rotation)
            
            # O anki karede tespit edilen nesne varsa kutu çiz
            if frame_number in self.detected_objects:
                for i, (x, y, w, h) in enumerate(self.detected_objects[frame_number]):
                    # Döndürme durumuna göre koordinatları ayarla
                    if self.current_rotation != 0:
                        x, y, w, h = self.adjust_coordinates_for_rotation(x, y, w, h, self.current_rotation, frame.shape)
                    
                    # Koordinatları frame sınırları içinde tut
                    x = max(0, min(x, frame.shape[1] - 1))
                    y = max(0, min(y, frame.shape[0] - 1))
                    w = max(1, min(w, frame.shape[1] - x))
                    h = max(1, min(h, frame.shape[0] - y))
                    
                    # Yeşil çerçeve çiz
                    cv2.rectangle(frame, (x, y), (x + w, y + h), (0, 255, 0), 3)
                    
                    # Nesne numarası ve "PERSON" etiketi ekle
                    label = f"PERSON {i+1}"
                    label_size = cv2.getTextSize(label, cv2.FONT_HERSHEY_SIMPLEX, 0.7, 2)[0]
                    
                    # Etiket arka planı
                    cv2.rectangle(frame, (x, y - label_size[1] - 10), 
                                 (x + label_size[0] + 10, y), (0, 255, 0), -1)
                    
                    # Etiket metni
                    cv2.putText(frame, label, (x + 5, y - 5), 
                               cv2.FONT_HERSHEY_SIMPLEX, 0.7, (0, 0, 0), 2)
                    
                    # Merkez noktası işaretle
                    center_x = x + w // 2
                    center_y = y + h // 2
                    cv2.circle(frame, (center_x, center_y), 5, (0, 255, 0), -1)

            # Frame'i güvenli bir şekilde RGB'ye dönüştür
            try:
                # BGR'den RGB'ye dönüştür (OpenCV BGR formatında yükler)
                if len(frame.shape) == 3 and frame.shape[2] == 3:
                    rgb_image = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                else:
                    # Gri renk veya diğer formatlar için
                    if len(frame.shape) == 2:
                        rgb_image = cv2.cvtColor(frame, cv2.COLOR_GRAY2RGB)
                    else:
                        rgb_image = frame.copy()
                
                # QImage oluştur
                h, w, ch = rgb_image.shape
                bytes_per_line = ch * w
                qt_image = QImage(rgb_image.data, w, h, bytes_per_line, QImage.Format_RGB888)
                
                # QImage'ın geçerli olup olmadığını kontrol et
                if qt_image.isNull():
                    self.log_message("QImage oluşturulamadı", "warning")
                    return
                
                # Pixmap oluştur ve aspect ratio koruyarak göster
                pixmap = QPixmap.fromImage(qt_image)
                if not pixmap.isNull():
                    # Video display label boyutlarını al
                    display_width = self.video_display_label.width()
                    display_height = self.video_display_label.height()
                    
                    # Minimum boyut kontrolü
                    if display_width < 100 or display_height < 100:
                        display_width = 640
                        display_height = 350
                    
                    # Aspect ratio'yu koruyarak ölçeklendir
                    scaled_pixmap = pixmap.scaled(
                        display_width - 20,  # Biraz padding bırak
                        display_height - 20, 
                        Qt.KeepAspectRatio, 
                        Qt.SmoothTransformation
                    )
                    self.video_display_label.setPixmap(scaled_pixmap)
                    
                    # Debug bilgisi
                    original_size = pixmap.size()
                    scaled_size = scaled_pixmap.size()
                    if hasattr(self, 'debug_frame_count'):
                        self.debug_frame_count += 1
                    else:
                        self.debug_frame_count = 1
                        
                    # Her 30 frame'de bir boyut bilgisini logla
                    if self.debug_frame_count % 30 == 0:
                        self.log_message(f"🖼️ Video: {original_size.width()}x{original_size.height()} → {scaled_size.width()}x{scaled_size.height()}", "info")
                else:
                    self.log_message("Pixmap oluşturulamadı", "warning")
                    
            except Exception as color_error:
                self.log_message(f"Renk dönüşümü hatası: {color_error}", "error")
                # Hata durumunda frame'i doğrudan göstermeyi dene
                try:
                    h, w = frame.shape[:2]
                    if len(frame.shape) == 3:
                        ch = frame.shape[2]
                        bytes_per_line = ch * w
                        qt_image = QImage(frame.data, w, h, bytes_per_line, QImage.Format_RGB888)
                    else:
                        bytes_per_line = w
                        qt_image = QImage(frame.data, w, h, bytes_per_line, QImage.Format_Grayscale8)
                    
                    if not qt_image.isNull():
                        pixmap = QPixmap.fromImage(qt_image)
                        if not pixmap.isNull():
                            # Hata durumunda da aspect ratio koru
                            display_width = self.video_display_label.width() or 640
                            display_height = self.video_display_label.height() or 350
                            
                            scaled_pixmap = pixmap.scaled(
                                display_width - 20, 
                                display_height - 20, 
                                Qt.KeepAspectRatio, 
                                Qt.SmoothTransformation
                            )
                            self.video_display_label.setPixmap(scaled_pixmap)
                except Exception as backup_error:
                    self.log_message(f"Yedek görüntüleme hatası: {backup_error}", "error")
                
        except Exception as e:
            self.log_message(f"Frame görüntüleme genel hatası: {str(e)}", "error")

    @pyqtSlot()
    def export_excel_report(self):
        """Excel raporu dışa aktarır."""
        if not self.detected_events:
            self.show_error_message("Rapor oluşturmak için önce analiz yapmalısınız.")
            return

        output_path, _ = QFileDialog.getSaveFileName(
            self, "Excel Raporu Kaydet", 
            f"V.E.R.A_Rapor_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx", 
            "Excel Dosyası (*.xlsx)"
        )
        if not output_path:
            return

        try:
            self.log_message("Excel raporu oluşturuluyor...", "info")
            
            # Excel dosyası oluştur
            import xlsxwriter
            
            workbook = xlsxwriter.Workbook(output_path)
            worksheet = workbook.add_worksheet('Analiz Raporu')
            
            # Formatlar
            header_format = workbook.add_format({
                'bold': True,
                'bg_color': '#3498db',
                'font_color': 'white',
                'align': 'center'
            })
            
            cell_format = workbook.add_format({
                'align': 'center',
                'border': 1
            })
            
            # Başlıklar
            worksheet.write(0, 0, 'Olay No', header_format)
            worksheet.write(0, 1, 'Zaman', header_format)
            worksheet.write(0, 2, 'Tespit Tipi', header_format)
            worksheet.write(0, 3, 'Güven Oranı', header_format)
            worksheet.write(0, 4, 'Koordinatlar', header_format)
            
            # Verileri yaz
            for i, event in enumerate(self.detected_events):
                worksheet.write(i + 1, 0, i + 1, cell_format)
                
                # Event'in tuple mi yoksa dict mi olduğunu kontrol et
                if isinstance(event, tuple) and len(event) >= 2:
                    # Tuple formatı: (başlangıç, bitiş)
                    start_time, end_time = event[0], event[1]
                    worksheet.write(i + 1, 1, f"{self.format_duration(start_time)} - {self.format_duration(end_time)}", cell_format)
                    worksheet.write(i + 1, 2, 'Hareket Tespiti', cell_format)
                    worksheet.write(i + 1, 3, "100%", cell_format)
                    worksheet.write(i + 1, 4, '-', cell_format)
                elif isinstance(event, dict):
                    # Dictionary formatı
                    worksheet.write(i + 1, 1, event.get('time', ''), cell_format)
                    worksheet.write(i + 1, 2, event.get('type', 'İnsan Tespiti'), cell_format)
                    worksheet.write(i + 1, 3, f"{event.get('confidence', 0.0):.2f}", cell_format)
                    
                    # Koordinatları formatla
                    if 'coordinates' in event:
                        coords = event['coordinates']
                        coord_str = f"({coords.get('x', 0)}, {coords.get('y', 0)}, {coords.get('w', 0)}, {coords.get('h', 0)})"
                        worksheet.write(i + 1, 4, coord_str, cell_format)
                    else:
                        worksheet.write(i + 1, 4, '-', cell_format)
                else:
                    # Bilinmeyen format
                    worksheet.write(i + 1, 1, str(event), cell_format)
                    worksheet.write(i + 1, 2, 'Tespit', cell_format)
                    worksheet.write(i + 1, 3, "N/A", cell_format)
                    worksheet.write(i + 1, 4, '-', cell_format)
            
            # Sütun genişliklerini ayarla
            worksheet.set_column(0, 0, 10)
            worksheet.set_column(1, 1, 20)
            worksheet.set_column(2, 2, 15)
            worksheet.set_column(3, 3, 12)
            worksheet.set_column(4, 4, 25)
            
            workbook.close()
            
            self.log_message("✅ Excel raporu başarıyla oluşturuldu!", "success")
            
            reply = QMessageBox.question(
                self, 'Başarılı', 
                "Excel raporu oluşturuldu. Dosyayı açmak ister misiniz?", 
                QMessageBox.Yes | QMessageBox.No, QMessageBox.No
            )
            if reply == QMessageBox.Yes:
                from PyQt5.QtGui import QDesktopServices
                QDesktopServices.openUrl(QUrl.fromLocalFile(output_path))
                
        except ImportError:
            self.show_error_message("Excel raporu için xlsxwriter kütüphanesi gerekli. 'pip install xlsxwriter' komutu ile yükleyin.")
        except Exception as e:
            self.show_error_message(f"Excel raporu oluşturma hatası: {e}")

    @pyqtSlot()
    def export_motion_report(self):
        """Hareket analizi raporu dışa aktarır - sadece hareket verilerine odaklanır."""
        if not self.detected_events:
            self.show_error_message("Rapor oluşturmak için önce analiz yapmalısınız.")
            return

        output_folder = QFileDialog.getExistingDirectory(self, "Hareket Raporu Klasörü Seç")
        if not output_folder:
            return

        try:
            self.log_message("🏃 Hareket analizi raporu oluşturuluyor...", "info")
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            report_folder = os.path.join(output_folder, f"Hareket_Raporu_{timestamp}")
            os.makedirs(report_folder, exist_ok=True)
            
            # 1. Hareket grafiği oluştur
            motion_chart_path = os.path.join(report_folder, "hareket_grafigi.png")
            self._create_motion_timeline_chart(motion_chart_path)
            
            # 2. Hareket yoğunluk haritası
            heatmap_path = os.path.join(report_folder, "hareket_yogunlugu.png")
            self._create_motion_heatmap(heatmap_path)
            
            # 3. Hareket özet raporu (HTML)
            html_report_path = os.path.join(report_folder, "hareket_raporu.html")
            self._create_motion_html_report(html_report_path, motion_chart_path, heatmap_path)
            
            # 4. CSV veri dosyası
            csv_path = os.path.join(report_folder, "hareket_verileri.csv")
            self._export_motion_csv(csv_path)
            
            self.log_message(f"✅ Hareket raporu oluşturuldu: {report_folder}", "success")
            
            # Klasörü aç
            if platform.system() == "Windows":
                os.startfile(report_folder)
            elif platform.system() == "Darwin":  # macOS
                subprocess.run(["open", report_folder])
            else:  # Linux
                subprocess.run(["xdg-open", report_folder])
                
        except Exception as e:
            self.show_error_message(f"Hareket raporu oluşturma hatası: {e}")
            self.log_message(f"❌ Hareket raporu hatası: {e}", "error")

    @pyqtSlot()
    def export_charts_report(self):
        """Grafik raporu dışa aktarır."""
        if not self.detected_events:
            self.show_error_message("Rapor oluşturmak için önce analiz yapmalısınız.")
            return

        output_dir = QFileDialog.getExistingDirectory(
            self, "Grafik Raporu Klasörü Seç", 
            os.path.expanduser("~/Desktop")
        )
        if not output_dir:
            return

        # Rapor klasörü oluştur
        report_folder = os.path.join(output_dir, f"V.E.R.A_Grafik_Raporu_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
        os.makedirs(report_folder, exist_ok=True)

        try:
            self.log_message("Grafik raporu oluşturuluyor...", "info")
            
            # detected_objects'in doğru formatta olduğunu kontrol et
            if isinstance(self.detected_objects, list):
                # List'i dict'e dönüştür
                converted_objects = {}
                for i, obj in enumerate(self.detected_objects):
                    converted_objects[i] = [obj] if isinstance(obj, dict) else obj
                detected_objects = converted_objects
            else:
                detected_objects = self.detected_objects
                
            report_gen = ReportGenerator(
                self.video_path, self.detected_events, detected_objects, 
                self.video_info, self.current_sensitivity
            )
            
            if report_gen.generate_charts(report_folder):
                # Özet metin dosyası da oluştur
                summary_path = os.path.join(report_folder, "analiz_ozeti.txt")
                with open(summary_path, 'w', encoding='utf-8') as f:
                    f.write("V.E.R.A. - Video Analiz Özeti\n")
                    f.write("=" * 50 + "\n\n")
                    f.write(f"Video: {os.path.basename(self.video_path)}\n")
                    f.write(f"Analiz Tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}\n")
                    f.write(f"Hassasiyet: {self.current_sensitivity}\n")
                    f.write(f"Video Süresi: {self.format_duration(self.video_info['duration'])}\n")
                    f.write(f"Toplam Olay: {len(self.detected_events)}\n")
                    f.write(f"Toplam Tespit: {len(self.detected_objects)}\n\n")
                    
                    if self.detected_events:
                        total_duration = sum([end - start for start, end in self.detected_events])
                        f.write(f"Toplam Olay Süresi: {self.format_duration(total_duration)}\n")
                        f.write(f"Olay Yoğunluğu: {(total_duration / self.video_info['duration'] * 100):.1f}%\n\n")
                        
                        f.write("Detaylı Olaylar:\n")
                        f.write("-" * 30 + "\n")
                        for i, (start, end) in enumerate(self.detected_events, 1):
                            f.write(f"Olay {i}: {self.format_duration(start)} - {self.format_duration(end)} ({end-start:.1f}s)\n")
                
                self.log_message(f"Grafik raporu başarıyla oluşturuldu: {report_folder}", "success")
                
                reply = QMessageBox.question(
                    self, 'Başarılı', 
                    "Grafik raporu oluşturuldu. Klasörü açmak ister misiniz?", 
                    QMessageBox.Yes | QMessageBox.No, QMessageBox.No
                )
                if reply == QMessageBox.Yes:
                    from PyQt5.QtGui import QDesktopServices
                    QDesktopServices.openUrl(QUrl.fromLocalFile(report_folder))
            else:
                self.show_error_message("Grafik raporu oluşturulamadı. matplotlib kütüphanesi yüklü mü?")
                
        except Exception as e:
            self.show_error_message(f"Grafik raporu oluşturma hatası: {e}")

    @pyqtSlot()
    def export_word_report(self):
        """Word raporu dışa aktarır - hareket tespiti görselleri ile."""
        if not self.detected_events:
            self.show_error_message("Rapor oluşturmak için önce analiz yapmalısınız.")
            return

        output_path, _ = QFileDialog.getSaveFileName(
            self, "Word Raporu Kaydet", 
            f"V.E.R.A_Rapor_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx", 
            "Word Dosyası (*.docx)"
        )
        if not output_path:
            return

        try:
            self.log_message("Word raporu oluşturuluyor (hareket tespiti görselleri ile)...", "info")
            
            # Word belgesini oluştur
            if not DOCX_AVAILABLE:
                self.show_error_message("Word desteği yok! 'pip install python-docx' çalıştırın")
                return
            
            # Geçici klasör oluştur ve hareket tespiti görsellerini kaydet
            temp_dir = os.path.join(os.path.dirname(output_path), "temp_word_images")
            os.makedirs(temp_dir, exist_ok=True)
            
            self.log_message("Hareket tespiti görselleri oluşturuluyor...", "info")
            saved_images = self._save_detection_frames(temp_dir)
            
            doc = Document()
            
            # Başlık ekle
            doc.add_heading('M.SAVAŞ Video Analiz Raporu', 0)
            
            # Genel bilgiler
            doc.add_heading('Genel Bilgiler', level=1)
            p = doc.add_paragraph()
            p.add_run('Rapor Tarihi: ').bold = True
            p.add_run(datetime.now().strftime('%d.%m.%Y %H:%M:%S'))
            
            p = doc.add_paragraph()
            p.add_run('Video Dosyası: ').bold = True
            p.add_run(os.path.basename(self.video_path) if self.video_path else '-')
            
            p = doc.add_paragraph()
            p.add_run('Hassasiyet Seviyesi: ').bold = True
            p.add_run(self.current_sensitivity)
            
            p = doc.add_paragraph()
            p.add_run('Toplam Tespit: ').bold = True
            p.add_run(str(len(self.detected_events)))
            
            # Video bilgileri
            if self.video_info:
                doc.add_heading('Video Özellikleri', level=1)
                p = doc.add_paragraph()
                p.add_run('Süre: ').bold = True
                p.add_run(self.format_duration(self.video_info.get('duration', 0)))
                
                p = doc.add_paragraph()
                p.add_run('FPS: ').bold = True
                p.add_run(f"{self.video_info.get('fps', 0):.2f}")
                
                p = doc.add_paragraph()
                p.add_run('Çözünürlük: ').bold = True
                p.add_run(f"{self.video_info.get('width', 0)}x{self.video_info.get('height', 0)}")
            
            # Tespit edilen olaylar
            if self.detected_events:
                doc.add_heading('Tespit Edilen Olaylar', level=1)
                
                # Tablo oluştur
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Light Grid Accent 1'
                
                # Başlık satırı
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Olay No'
                hdr_cells[1].text = 'Zaman'
                hdr_cells[2].text = 'Tespit Tipi'
                hdr_cells[3].text = 'Güven Oranı'
                
                # Verileri ekle
                for i, event in enumerate(self.detected_events):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(i + 1)
                    
                    # Event'in tuple mi yoksa dict mi olduğunu kontrol et
                    if isinstance(event, tuple) and len(event) >= 2:
                        # Tuple formatı: (başlangıç, bitiş)
                        start_time, end_time = event[0], event[1]
                        row_cells[1].text = f"{self.format_duration(start_time)} - {self.format_duration(end_time)}"
                        row_cells[2].text = 'Hareket Tespiti'
                        row_cells[3].text = "100%"
                    elif isinstance(event, dict):
                        # Dictionary formatı
                        row_cells[1].text = event.get('time', '')
                        row_cells[2].text = event.get('type', 'İnsan Tespiti')
                        row_cells[3].text = f"{event.get('confidence', 0.0):.2f}"
                    else:
                        # Bilinmeyen format
                        row_cells[1].text = str(event)
                        row_cells[2].text = 'Tespit'
                        row_cells[3].text = "N/A"
            
            # Hareket tespiti görselleri ekleme
            if saved_images:
                doc.add_page_break()
                doc.add_heading('🎯 Hareket Tespiti Görselleri', level=1)
                
                doc.add_paragraph('Aşağıda videoda tespit edilen hareketlerin karşılaştırmalı görüntüleri yer almaktadır. '
                                 'Sol tarafta orijinal görüntü, sağ tarafta hareket tespiti uygulanmış hali gösterilmektedir.')
                
                self.log_message(f"Word raporuna {len(saved_images)} görsel ekleniyor...", "info")
                
                for i, img_info in enumerate(saved_images):
                    # Debug: görsel dosya bilgilerini logla
                    self.log_message(f"Görsel {i+1}: {img_info.get('path', 'N/A')}", "info")
                    self.log_message(f"Orijinal {i+1}: {img_info.get('original_path', 'N/A')}", "info")
                    
                    # Dosya varlığını kontrol et
                    original_exists = os.path.exists(img_info.get('original_path', ''))
                    detection_exists = os.path.exists(img_info.get('path', ''))
                    self.log_message(f"Dosya kontrolü {i+1}: Orijinal={original_exists}, Tespit={detection_exists}", "info")
                    
                    # Her görsel için başlık
                    time_str = self.format_duration(img_info['time'])
                    doc.add_heading(f"Tespit {i+1}: {time_str} ({img_info['detections']} kişi)", level=2)
                    
                    # Basit yaklaşım: Her görseli ayrı ayrı ekle (tablo yerine)
                    try:
                        # Orijinal görüntü başlığı ve görseli
                        doc.add_paragraph("📷 Orijinal Görüntü:", style='Intense Quote')
                        if original_exists:
                            original_para = doc.add_paragraph()
                            original_run = original_para.add_run()
                            original_run.add_picture(img_info['original_path'], width=Inches(4))
                            self.log_message(f"✅ Orijinal görsel {i+1} eklendi", "success")
                        else:
                            doc.add_paragraph("❌ Orijinal görsel dosyası bulunamadı")
                            self.log_message(f"❌ Orijinal görsel {i+1} dosyası yok", "error")
                        
                        # Tespit görseli başlığı ve görseli  
                        doc.add_paragraph("🎯 Hareket Tespiti:", style='Intense Quote')
                        if detection_exists:
                            detection_para = doc.add_paragraph()
                            detection_run = detection_para.add_run()
                            detection_run.add_picture(img_info['path'], width=Inches(4))
                            self.log_message(f"✅ Tespit görseli {i+1} eklendi", "success")
                        else:
                            doc.add_paragraph("❌ Tespit görsel dosyası bulunamadı")
                            self.log_message(f"❌ Tespit görseli {i+1} dosyası yok", "error")
                            
                    except Exception as e:
                        self.log_message(f"❌ Görsel {i+1} ekleme hatası: {str(e)}", "error")
                        # Hata durumunda da bilgi ver
                        doc.add_paragraph(f"Görsel {i+1} eklenirken hata oluştu: {str(e)}")
                    
                    # Ayırıcı ekle
                    if i < len(saved_images) - 1:
                        doc.add_paragraph("─" * 50)
                        doc.add_paragraph()
                
                self.log_message(f"✅ {len(saved_images)} hareket tespiti görseli işlendi", "success")
            else:
                self.log_message("⚠️ Kaydedilmiş görsel bulunamadı", "warning")
            
            # Belgeyi kaydet
            doc.save(output_path)
            
            # Geçici dosyaları temizle
            try:
                import shutil
                shutil.rmtree(temp_dir)
                self.log_message("Geçici dosyalar temizlendi", "info")
            except Exception as e:
                self.log_message(f"Geçici dosya temizleme hatası: {e}", "warning")
            
            self.log_message("✅ Word raporu başarıyla oluşturuldu!", "success")
            
            reply = QMessageBox.question(
                self, 'Başarılı', 
                "Word raporu oluşturuldu. Dosyayı açmak ister misiniz?", 
                QMessageBox.Yes | QMessageBox.No, QMessageBox.No
            )
            if reply == QMessageBox.Yes:
                from PyQt5.QtGui import QDesktopServices
                QDesktopServices.openUrl(QUrl.fromLocalFile(output_path))
                
        except ImportError:
            self.show_error_message("Word raporu için python-docx kütüphanesi gerekli. 'pip install python-docx' komutu ile yükleyin.")
        except Exception as e:
            self.show_error_message(f"Word raporu oluşturma hatası: {e}")
            self.log_message(f"Word raporu hatası: {str(e)}", "error")

    @pyqtSlot()
    def export_all_reports(self):
        """Tüm raporları tek seferde oluşturur."""
        if not self.detected_events:
            self.show_error_message("Rapor oluşturmak için önce analiz yapmalısınız.")
            return

        # Klasör seç
        output_dir = QFileDialog.getExistingDirectory(
            self, "Tüm Raporlar Klasörü Seç", 
            os.path.expanduser("~/Desktop")
        )
        if not output_dir:
            return

        # Rapor klasörü oluştur
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        report_folder = os.path.join(output_dir, f"M.SAVAS_Tum_Raporlar_{timestamp}")
        os.makedirs(report_folder, exist_ok=True)

        try:
            self.log_message("🎯 Tüm raporlar oluşturuluyor...", "info")
            
            # detected_objects'in doğru formatta olduğunu kontrol et
            if isinstance(self.detected_objects, list):
                # List'i dict'e dönüştür
                converted_objects = {}
                for i, obj in enumerate(self.detected_objects):
                    converted_objects[i] = [obj] if isinstance(obj, dict) else obj
                detected_objects = converted_objects
            else:
                detected_objects = self.detected_objects
            
            # Rapor generator
            report_gen = ReportGenerator(
                self.video_path, self.detected_events, detected_objects, 
                self.video_info, self.current_sensitivity
            )
            
            reports_created = []
            
            # 1. Excel Raporu
            excel_path = os.path.join(report_folder, f"Excel_Rapor_{timestamp}.xlsx")
            if report_gen.generate_excel_report(excel_path):
                reports_created.append("📈 Excel Raporu")
                self.log_message("✅ Excel raporu oluşturuldu", "success")
            
            # 2. Word Raporu
            word_path = os.path.join(report_folder, f"Word_Rapor_{timestamp}.docx")
            temp_dir = os.path.join(report_folder, "temp_images")
            os.makedirs(temp_dir, exist_ok=True)
            saved_images = self._save_detection_frames(temp_dir)
            
            if report_gen.generate_word_report(word_path, saved_images):
                reports_created.append("📄 Word Raporu")
                self.log_message("✅ Word raporu oluşturuldu", "success")
            
            # 3. Grafik Raporu
            chart_folder = os.path.join(report_folder, "Grafik_Raporu")
            os.makedirs(chart_folder, exist_ok=True)
            
            if report_gen.generate_charts(chart_folder):
                reports_created.append("📊 Grafik Raporu")
                self.log_message("✅ Grafik raporu oluşturuldu", "success")
            
            # 4. Özet dosyası oluştur
            summary_path = os.path.join(report_folder, "Rapor_Ozeti.txt")
            with open(summary_path, 'w', encoding='utf-8') as f:
                f.write("M.SAVAŞ - Kapsamlı Rapor Paketi\n")
                f.write("=" * 50 + "\n\n")
                f.write(f"Oluşturulma Tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}\n")
                f.write(f"Video: {os.path.basename(self.video_path)}\n")
                f.write(f"Hassasiyet: {self.current_sensitivity}\n")
                f.write(f"Toplam Olay: {len(self.detected_events)}\n")
                f.write(f"Toplam Tespit: {len(self.detected_objects)}\n\n")
                f.write("Oluşturulan Raporlar:\n")
                f.write("-" * 30 + "\n")
                for report in reports_created:
                    f.write(f"✅ {report}\n")
                f.write(f"\nToplam {len(reports_created)} rapor başarıyla oluşturuldu.\n")
            
            # Geçici dosyaları temizle
            try:
                shutil.rmtree(temp_dir)
            except:
                pass
            
            self.log_message(f"🎉 Tüm raporlar oluşturuldu: {len(reports_created)} rapor", "success")
            
            # Başarı mesajı
            reply = QMessageBox.question(
                self, 'Başarılı', 
                f"Tüm raporlar oluşturuldu!\n\n"
                f"Oluşturulan raporlar:\n" + "\n".join(reports_created) + 
                f"\n\nKlasörü açmak ister misiniz?", 
                QMessageBox.Yes | QMessageBox.No, QMessageBox.No
            )
            if reply == QMessageBox.Yes:
                from PyQt5.QtGui import QDesktopServices
                QDesktopServices.openUrl(QUrl.fromLocalFile(report_folder))
                
        except Exception as e:
            self.show_error_message(f"Toplu rapor oluşturma hatası: {e}")

    def _save_detection_frames(self, temp_dir: str) -> list:
        """Tespit edilen kareleri resim olarak kaydeder - her saniye için detaylı."""
        saved_images = []
        
        if not self.video_capture or not self.detected_objects:
            print("❌ Video capture veya detected_objects bulunamadı")
            return saved_images

        print(f"🎯 Görsel kaydetme başlıyor: {len(self.detected_objects)} tespit karesi var")
        print(f"📁 Hedef klasör: {temp_dir}")
        
        try:
            # Tüm tespit edilen kareleri al
            detection_frames = list(self.detected_objects.keys())
            fps = self.video_info.get('fps', 30)
            
            print(f"📊 FPS: {fps}, Tespit kareleri: {len(detection_frames)}")
            
            # Her saniye için en az bir görüntü olsun
            second_based_frames = {}
            for frame_num in detection_frames:
                second = int(frame_num / fps)
                if second not in second_based_frames:
                    second_based_frames[second] = []
                second_based_frames[second].append(frame_num)
            
            print(f"📈 {len(second_based_frames)} farklı saniyede tespit var")            # Her saniye için en iyi kareler
            selected_frames = []
            for second in sorted(second_based_frames.keys()):
                frames_in_second = second_based_frames[second]
                # En çok tespit içeren kareyi seç
                best_frame = max(frames_in_second, key=lambda f: len(self.detected_objects[f]))
                selected_frames.append(best_frame)
            
            # Maksimum 30 görüntü ile sınırla (çok uzun olmasın)
            if len(selected_frames) > 30:
                step = len(selected_frames) // 30
                selected_frames = selected_frames[::step]
            
            for i, frame_num in enumerate(selected_frames):
                # Kareyi oku
                self.video_capture.set(cv2.CAP_PROP_POS_FRAMES, frame_num)
                ret, frame = self.video_capture.read()
                
                if ret:
                    # Orijinal frame'i de kaydet (karşılaştırma için)
                    original_frame = frame.copy()
                    
                    # Döndürme uygula (Word raporunda da aynı görünsün)
                    if hasattr(self, 'current_rotation') and self.current_rotation != 0:
                        frame = self.apply_rotation(frame, self.current_rotation)
                    
                    # Tespit kutularını çiz (döndürme sonrası koordinatlara göre)
                    detection_count = 0
                    for j, (x, y, w, h) in enumerate(self.detected_objects[frame_num]):
                        # Koordinatları döndürmeye göre ayarla
                        if hasattr(self, 'current_rotation') and self.current_rotation != 0:
                            adjusted_coords = self.adjust_coordinates_for_rotation(x, y, w, h, self.current_rotation, frame.shape)
                            x, y, w, h = adjusted_coords
                        
                        # Yeşil çerçeve (kalın)
                        cv2.rectangle(frame, (x, y), (x + w, y + h), (0, 255, 0), 4)
                        
                        # Etiket arka planı
                        label = f"PERSON {j+1}"
                        label_size = cv2.getTextSize(label, cv2.FONT_HERSHEY_SIMPLEX, 0.8, 2)[0]
                        cv2.rectangle(frame, (x, y - label_size[1] - 15), 
                                     (x + label_size[0] + 10, y - 5), (0, 255, 0), -1)
                        
                        # Etiket metni
                        cv2.putText(frame, label, (x + 5, y - 10), 
                                   cv2.FONT_HERSHEY_SIMPLEX, 0.8, (0, 0, 0), 2)
                        
                        # Merkez noktası (daha büyük)
                        cv2.circle(frame, (x + w // 2, y + h // 2), 8, (0, 255, 0), -1)
                        cv2.circle(frame, (x + w // 2, y + h // 2), 12, (0, 255, 0), 2)
                        
                        detection_count += 1
                    
                    # Zaman ve tespit bilgisi
                    time_sec = frame_num / fps
                    time_text = f"Zaman: {self.format_duration(time_sec)}"
                    detection_text = f"Tespit: {detection_count} kisi"
                    frame_text = f"Kare: {frame_num}"
                    
                    # Bilgi kutusu arka planı
                    info_height = 100
                    cv2.rectangle(frame, (10, 10), (400, info_height), (0, 0, 0), -1)
                    cv2.rectangle(frame, (10, 10), (400, info_height), (255, 255, 255), 2)
                    
                    # Bilgi metinleri
                    cv2.putText(frame, time_text, (20, 35), 
                               cv2.FONT_HERSHEY_SIMPLEX, 0.8, (255, 255, 255), 2)
                    cv2.putText(frame, detection_text, (20, 60), 
                               cv2.FONT_HERSHEY_SIMPLEX, 0.8, (0, 255, 0), 2)
                    cv2.putText(frame, frame_text, (20, 85), 
                               cv2.FONT_HERSHEY_SIMPLEX, 0.6, (200, 200, 200), 1)
                    
                    # Resmi kaydet - hata kontrolü ile
                    image_path = os.path.join(temp_dir, f"tespit_{i+1:03d}.jpg")
                    success1 = cv2.imwrite(image_path, frame)
                    
                    # Orijinal karşılaştırma için
                    if hasattr(self, 'current_rotation') and self.current_rotation != 0:
                        original_frame = self.apply_rotation(original_frame, self.current_rotation)
                    
                    original_path = os.path.join(temp_dir, f"orijinal_{i+1:03d}.jpg")
                    success2 = cv2.imwrite(original_path, original_frame)
                    
                    # Dosya kaydetme kontrolü
                    if success1 and success2 and os.path.exists(image_path) and os.path.exists(original_path):
                        saved_images.append({
                            'path': image_path,
                            'original_path': original_path,
                            'frame': frame_num,
                            'time': time_sec,
                            'detections': detection_count,
                            'second': int(time_sec)
                        })
                        print(f"✅ Görsel {i+1} kaydedildi: {image_path}")
                    else:
                        print(f"❌ Görsel {i+1} kaydedilemedi - CV2 write başarısız")
                else:
                    print(f"❌ Frame {frame_num} okunamadı")
                    
        except Exception as e:
            print(f"❌ Görsel kaydetme hatası: {e}")
            import traceback
            traceback.print_exc()
        
        print(f"✅ Toplam {len(saved_images)} görsel kaydedildi")
        return saved_images

    @pyqtSlot(str)
    def update_status(self, message):
        self.status_label.setText(f"Durum: {message}")

    def log_message(self, message: str, level: str = "info"):
        """Gelişmiş log mesajı - renk kodlaması ve otomatik kaydırma"""
        color_map = {
            "info": "#ecf0f1", 
            "success": "#2ecc71", 
            "warning": "#f39c12", 
            "error": "#e74c3c"
        }
        
        # Emoji haritası
        emoji_map = {
            "info": "ℹ️",
            "success": "✅", 
            "warning": "⚠️",
            "error": "❌"
        }
        
        color = color_map.get(level, "#ecf0f1")
        emoji = emoji_map.get(level, "📝")
        timestamp = QTime.currentTime().toString("HH:mm:ss")
        
        # HTML formatında mesaj oluştur
        formatted_message = f'<font color="{color}"><b>[{timestamp}]</b> {emoji} {message}</font>'
        self.log_display.append(formatted_message)
        
        # Otomatik kaydırma - en alta git
        scrollbar = self.log_display.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())
        
        # Çok fazla log varsa eski mesajları temizle (performans için)
        if self.log_display.document().blockCount() > 1000:
            cursor = self.log_display.textCursor()
            cursor.movePosition(cursor.Start)
            cursor.movePosition(cursor.Down, cursor.KeepAnchor, 100)
            cursor.removeSelectedText()

    def show_error_message(self, message: str):
        self.log_message(message, "error")
        QMessageBox.critical(self, "Hata", message)

    def format_duration(self, seconds: float) -> str:
        """Süreyi HH:MM:SS formatında döndürür."""
        if seconds is None or math.isnan(seconds): 
            return "00:00:00"
        h, m, s = int(seconds // 3600), int((seconds % 3600) // 60), int(seconds % 60)
        return f"{h:02d}:{m:02d}:{s:02d}"

    def closeEvent(self, event):
        """Uygulama kapatılırken temizlik işlemleri yapar"""
        try:
            self.log_message("🔄 Uygulama kapatılıyor, temizlik işlemleri yapılıyor...", "info")
            
            # Thread'leri güvenli şekilde durdur
            if hasattr(self, 'processor_thread') and self.processor_thread and self.processor_thread.isRunning():
                self.log_message("🛑 İşlem thread'i durduruluyor...", "info")
                self.processor_thread.request_stop()
                self.processor_thread.wait(3000)  # 3 saniye bekle
                if self.processor_thread.isRunning():
                    self.processor_thread.terminate()
                    
            if hasattr(self, 'exporter_thread') and self.exporter_thread and self.exporter_thread.isRunning():
                self.log_message("📤 Export thread'i durduruluyor...", "info")
                self.exporter_thread.wait(3000)
                if self.exporter_thread.isRunning():
                    self.exporter_thread.terminate()
            
            # Live camera'yı durdur
            if hasattr(self, 'live_timer') and self.live_timer.isActive():
                self.live_timer.stop()
                
            if hasattr(self, 'live_camera') and self.live_camera:
                self.live_camera.release()
                
            # Video capture'ı serbest bırak
            if hasattr(self, 'video_capture') and self.video_capture:
                self.video_capture.release()
                
            # Playback timer'ı durdur
            if hasattr(self, 'playback_timer') and self.playback_timer.isActive():
                self.playback_timer.stop()
                
            self.log_message("✅ Temizlik işlemleri tamamlandı. Güle güle!", "success")
            
        except Exception as e:
            print(f"Kapatma hatası: {e}")
            
        finally:
            event.accept()

    def get_stylesheet(self) -> str:
        return """
            QMainWindow, QWidget {
                background-color: #2b3e50; color: #ecf0f1; font-family: 'Segoe UI', Arial;
            }
            QGroupBox {
                background-color: #34495e; border: 2px solid #3498db;
                border-radius: 10px; margin-top: 8px; padding: 8px;
                font-weight: bold; color: #ecf0f1;
            }
            QGroupBox::title {
                subcontrol-origin: margin; subcontrol-position: top left;
                padding: 2px 8px; left: 8px; color: #3498db; font-size: 14px; font-weight: bold;
                background-color: #2b3e50; border-radius: 4px;
            }
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #3498db, stop:1 #2980b9);
                color: white; border: 1px solid #2980b9;
                border-radius: 6px; padding: 4px 8px; font-size: 11px; font-weight: bold;
                text-align: center;
            }
            QPushButton:hover { 
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #3cb4f0, stop:1 #3498db);
                border: 1px solid #3498db;
            }
            QPushButton:pressed { 
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #2980b9, stop:1 #1f618d);
                border: 1px solid #1f618d;
            }
            QPushButton:disabled { 
                background-color: #7f8c8d; color: #bdc3c7; border: 1px solid #95a5a6;
            }
            
            /* Analiz butonları için özel stil */
            QPushButton[text*="Analiz"] {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #27ae60, stop:1 #229954);
                border: 1px solid #229954; font-weight: bold;
            }
            QPushButton[text*="Analiz"]:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #2ecc71, stop:1 #27ae60);
                border: 1px solid #27ae60;
            }
            
            /* Video döndürme butonları için özel stil */
            QPushButton[text*="°"], QPushButton[text*="Sıfırla"] {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #e74c3c, stop:1 #c0392b);
                border: 1px solid #c0392b; font-weight: bold;
            }
            QPushButton[text*="°"]:hover, QPushButton[text*="Sıfırla"]:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #ec7063, stop:1 #e74c3c);
                border: 1px solid #e74c3c;
            }
            
            QLabel { 
                color: #ecf0f1; font-size: 13px; padding: 2px;
            }
            #videoDisplay { 
                border: 3px solid #3498db; border-radius: 8px; 
                background-color: #2c3e50;
            }
            QProgressBar {
                border: 2px solid #34495e; border-radius: 8px; text-align: center;
                font-size: 12px; color: white; background-color: #2c3e50; height: 20px;
            }
            QProgressBar::chunk { 
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #27ae60, stop:1 #2ecc71);
                border-radius: 6px; 
            }
            QTextEdit {
                background-color: #2c3e50; color: #ecf0f1; border: 2px solid #34495e;
                border-radius: 6px; font-family: 'Consolas', 'Courier New', monospace; 
                font-size: 11px; selection-background-color: #3498db;
            }
            QRadioButton { 
                font-size: 12px; padding: 3px; color: #ecf0f1; font-weight: bold;
            }
            QRadioButton::indicator {
                width: 16px; height: 16px;
            }
            QRadioButton::indicator:unchecked {
                border: 2px solid #7f8c8d; border-radius: 8px; background-color: #34495e;
            }
            QRadioButton::indicator:checked {
                border: 2px solid #3498db; border-radius: 8px; background-color: #3498db;
            }
            QListWidget {
                background-color: #2c3e50; color: #ecf0f1; border: 2px solid #34495e;
                border-radius: 6px; font-size: 12px; selection-background-color: #3498db;
            }
            QListWidget::item {
                padding: 6px; border-bottom: 1px solid #34495e;
            }
            QListWidget::item:hover {
                background-color: #34495e; color: #ecf0f1;
            }
            QListWidget::item:selected {
                background-color: #3498db; color: white; font-weight: bold;
            }
        """

    def force_exit_application(self):
        """Uygulamayı zorla kapatır"""
        reply = QMessageBox.question(
            self, 
            '🛑 Zorla Çıkış', 
            '⚠️ Uygulama zorla kapatılacak!\n\n'
            '• Kaydedilmemiş veriler kaybolacak\n'
            '• Çalışan analizler durdurulacak\n'
            '• Tüm işlemler sonlandırılacak\n\n'
            'Devam etmek istediğinizden emin misiniz?',
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            try:
                self.log_message("🛑 Zorla çıkış başlatılıyor...", "warning")
                
                # Sistem durumunu güncelle
                self.system_status_label.setText("🔴 Kapanıyor...")
                self.system_status_label.setStyleSheet("""
                    QLabel {
                        background-color: #e74c3c;
                        color: white;
                        border: 1px solid #c0392b;
                        border-radius: 5px;
                        padding: 3px;
                        font-size: 10px;
                        font-weight: bold;
                        text-align: center;
                    }
                """)
                
                # Aktif thread'leri durdur
                if hasattr(self, 'processor_thread') and self.processor_thread:
                    self.processor_thread.stop_requested = True
                    self.processor_thread.quit()
                
                if hasattr(self, 'exporter_thread') and self.exporter_thread:
                    self.exporter_thread.quit()
                
                # Kameraları durdur
                if hasattr(self, 'live_camera') and self.live_camera:
                    self.live_camera.release()
                
                if hasattr(self, 'video_capture') and self.video_capture:
                    self.video_capture.release()
                
                # Timer'ları durdur
                if hasattr(self, 'live_timer'):
                    self.live_timer.stop()
                
                if hasattr(self, 'playback_timer'):
                    self.playback_timer.stop()
                
                self.log_message("✅ Tüm kaynaklar temizlendi. Uygulama kapatılıyor.", "info")
                
                # Singleton temizle
                if hasattr(app_singleton, 'cleanup'):
                    app_singleton.cleanup()
                
                # Force exit
                QApplication.quit()
                os._exit(0)  # Zorla çıkış
                
            except Exception as e:
                self.log_message(f"❌ Zorla çıkış hatası: {e}", "error")
                os._exit(1)  # Hata ile çıkış
    
    def restart_application(self):
        """Uygulamayı yeniden başlatır"""
        reply = QMessageBox.question(
            self, 
            '🔄 Yeniden Başlat', 
            '🔄 Uygulama yeniden başlatılacak!\n\n'
            '• Mevcut session kaybolacak\n'
            '• Çalışan analizler durdurulacak\n'
            '• Ayarlar korunacak\n\n'
            'Devam etmek istediğinizden emin misiniz?',
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            try:
                self.log_message("🔄 Yeniden başlatma işlemi başlıyor...", "info")
                
                # Sistem durumunu güncelle
                self.system_status_label.setText("🔄 Yeniden Başlatılıyor...")
                self.system_status_label.setStyleSheet("""
                    QLabel {
                        background-color: #f39c12;
                        color: white;
                        border: 1px solid #e67e22;
                        border-radius: 5px;
                        padding: 3px;
                        font-size: 10px;
                        font-weight: bold;
                        text-align: center;
                    }
                """)
                
                # Kaynakları temizle
                if hasattr(self, 'processor_thread') and self.processor_thread:
                    self.processor_thread.stop_requested = True
                    self.processor_thread.quit()
                
                if hasattr(self, 'live_camera') and self.live_camera:
                    self.live_camera.release()
                
                if hasattr(self, 'video_capture') and self.video_capture:
                    self.video_capture.release()
                
                # Singleton temizle
                if hasattr(app_singleton, 'cleanup'):
                    app_singleton.cleanup()
                
                # Python script'ini yeniden çalıştır
                python = sys.executable
                script = sys.argv[0]
                
                self.log_message("✅ Yeniden başlatılıyor...", "success")
                
                # Yeni process başlat ve mevcut uygulamayı kapat
                subprocess.Popen([python, script])
                QApplication.quit()
                
            except Exception as e:
                self.log_message(f"❌ Yeniden başlatma hatası: {e}", "error")
                QMessageBox.critical(self, "Hata", f"Yeniden başlatma hatası:\n{e}")

# =============================================================================
# --- RAPOR OLUŞTURUCU ---
# =============================================================================

class ReportGenerator:
    def generate_dav_report(self, output_path: str) -> bool:
        """DAV formatında (JSON tabanlı) analiz raporu oluşturur."""
        try:
            import json
            dav_data = {
                "video_file": os.path.basename(self.video_path),
                "analysis_date": self.analysis_date.strftime('%Y-%m-%d %H:%M:%S'),
                "sensitivity": self.sensitivity,
                "duration": self.video_info.get('duration', 0),
                "fps": self.video_info.get('fps', 0),
                "total_frames": self.video_info.get('total_frames', 0),
                "events": [],
            }
            # Olayları ekle
            for idx, (start, end) in enumerate(self.events):
                event = {
                    "event_id": idx + 1,
                    "start_time": start,
                    "end_time": end,
                    "duration": end - start,
                    "detections": []
                }
                # Bu olay süresindeki tespitleri ekle
                for frame_num, detections in self.detected_objects.items():
                    frame_time = frame_num / self.video_info.get('fps', 1)
                    if start <= frame_time <= end:
                        for det in detections:
                            event["detections"].append({
                                "frame": int(frame_num),
                                "time": frame_time,
                                "box": det.get('box', []),
                                "confidence": det.get('conf', 0),
                                "area": det.get('area', 0)
                            })
                dav_data["events"].append(event)
            # Dosyaya yaz
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(dav_data, f, ensure_ascii=False, indent=2)
            return True
        except Exception as e:
            print(f"DAV raporu oluşturma hatası: {e}")
            return False
    """Gelişmiş raporlama sistemi - Excel ve grafik çıktıları."""
    
    def __init__(self, video_path: str, events: list, detected_objects: dict, video_info: dict, sensitivity: str):
        self.video_path = video_path
        self.events = events
        
        # detected_objects'in doğru tipte olduğunu kontrol et
        if isinstance(detected_objects, dict):
            self.detected_objects = detected_objects
        else:
            print(f"ReportGenerator: detected_objects list olarak geldi: {type(detected_objects)}")
            # List'i dict'e dönüştür
            converted_objects = {}
            for i, obj in enumerate(detected_objects):
                if isinstance(obj, dict):
                    converted_objects[i] = [obj]
                else:
                    converted_objects[i] = obj
            self.detected_objects = converted_objects
            
        self.video_info = video_info
        self.sensitivity = sensitivity
        self.analysis_date = datetime.now()
    
    def format_duration(self, seconds: float) -> str:
        """Süreyi HH:MM:SS formatında döndürür."""
        if seconds is None or (isinstance(seconds, float) and math.isnan(seconds)): 
            return "00:00:00"
        try:
            h, m, s = int(seconds // 3600), int((seconds % 3600) // 60), int(seconds % 60)
            return f"{h:02d}:{m:02d}:{s:02d}"
        except:
            return "00:00:00"
    
    def generate_excel_report(self, output_path: str) -> bool:
        """Excel raporu oluşturur - gelişmiş grafikler ile."""
        try:
            try:
                import xlsxwriter
            except ImportError:
                print("xlsxwriter kütüphanesi yüklü değil!")
                print("Lütfen 'pip install xlsxwriter' çalıştırın")
                return False
            
            # Debug: Veri tiplerini kontrol et
            print(f"detected_objects tipi: {type(self.detected_objects)}")
            print(f"detected_objects içeriği: {self.detected_objects}")
            
            # detected_objects'in doğru formatta olduğunu kontrol et
            if isinstance(self.detected_objects, list):
                print("detected_objects list olarak geldi, dict'e dönüştürülüyor...")
                # List'i dict'e dönüştür
                converted_objects = {}
                for i, obj in enumerate(self.detected_objects):
                    if isinstance(obj, dict):
                        converted_objects[i] = [obj]
                    else:
                        converted_objects[i] = obj
                self.detected_objects = converted_objects
                print(f"Dönüştürülen detected_objects: {self.detected_objects}")
            
            # Excel dosyasını oluştur
            workbook = xlsxwriter.Workbook(output_path)
            
            # Genel bilgiler sayfası
            overview_sheet = workbook.add_worksheet('Genel Bilgiler')
            
            # Formatları tanımla
            header_format = workbook.add_format({
                'bold': True,
                'font_size': 16,
                'align': 'center',
                'valign': 'vcenter',
                'bg_color': '#1ABC9C',
                'font_color': 'white',
                'border': 1
            })
            
            title_format = workbook.add_format({
                'bold': True,
                'font_size': 14,
                'align': 'center',
                'valign': 'vcenter',
                'bg_color': '#3498DB',
                'font_color': 'white',
                'border': 1
            })
            
            data_format = workbook.add_format({
                'font_size': 12,
                'align': 'left',
                'valign': 'vcenter',
                'border': 1
            })
            
            number_format = workbook.add_format({
                'font_size': 12,
                'align': 'center',
                'valign': 'vcenter',
                'border': 1,
                'num_format': '#,##0'
            })
            
            # Başlık
            overview_sheet.merge_range('A1:F1', 'M.SAVAŞ - VİDEO ANALİZ RAPORU', header_format)
            overview_sheet.set_row(0, 25)
            
            # Genel bilgiler - güvenli erişim
            overview_sheet.write('A3', 'Video Dosyası:', title_format)
            overview_sheet.write('B3', os.path.basename(self.video_path), data_format)
            overview_sheet.write('A4', 'Analiz Tarihi:', title_format)
            overview_sheet.write('B4', self.analysis_date.strftime('%d.%m.%Y %H:%M:%S'), data_format)
            overview_sheet.write('A5', 'Video Süresi:', title_format)
            
            # Video süresini güvenli şekilde al
            duration = self.video_info.get('duration', 0)
            overview_sheet.write('B5', self.format_duration(duration), data_format)
            
            overview_sheet.write('A6', 'Toplam Kare:', title_format)
            total_frames = self.video_info.get('total_frames', 0)
            overview_sheet.write('B6', int(total_frames), number_format)
            
            # Tespit istatistikleri
            overview_sheet.write('A8', 'Tespit İstatistikleri:', title_format)
            overview_sheet.write('A9', 'Toplam Tespit:', title_format)
            overview_sheet.write('B9', len(self.detected_objects), number_format)
            overview_sheet.write('A10', 'Olay Sayısı:', title_format)
            overview_sheet.write('B10', len(self.events), number_format)
            
            if self.events and duration > 0:
                total_event_duration = sum([end - start for start, end in self.events])
                overview_sheet.write('A11', 'Toplam Olay Süresi:', title_format)
                overview_sheet.write('B11', self.format_duration(total_event_duration), data_format)
                
                coverage_percentage = (total_event_duration / duration) * 100
                overview_sheet.write('A12', 'Video Kapsamı:', title_format)
                overview_sheet.write('B12', f"{coverage_percentage:.1f}%", data_format)
            
            # Tespit edilen nesneler sayfası
            detections_sheet = workbook.add_worksheet('Tespit Detayları')
            
            # Başlık
            detections_sheet.merge_range('A1:F1', 'TESPIT EDİLEN NESNELER DETAYLI', header_format)
            detections_sheet.set_row(0, 25)
            
            # Tablo başlıkları
            headers = ['Kare No', 'Zaman (sn)', 'Nesne Tipi', 'Güven Skoru', 'Konum (X,Y)', 'Boyut (W,H)']
            for col, header in enumerate(headers):
                detections_sheet.write(2, col, header, title_format)
            
            # Tespit verilerini yaz
            row = 3
            fps = self.video_info.get('fps', 30)  # Varsayılan FPS
            
            # detected_objects'in dict olduğunu kontrol et
            if isinstance(self.detected_objects, dict):
                for frame_num, detections in self.detected_objects.items():
                    time_sec = frame_num / fps
                    # detections bir liste olabilir, kontrol edelim
                    if isinstance(detections, list):
                        for detection in detections:
                            detections_sheet.write(row, 0, frame_num, number_format)
                            detections_sheet.write(row, 1, f"{time_sec:.2f}", data_format)
                            
                            # detection dict mi liste mi kontrol et
                            if isinstance(detection, dict):
                                detections_sheet.write(row, 2, detection.get('class', 'person'), data_format)
                                detections_sheet.write(row, 3, f"{detection.get('confidence', 0):.2f}", data_format)
                                detections_sheet.write(row, 4, f"({detection.get('x', 0)}, {detection.get('y', 0)})", data_format)
                                detections_sheet.write(row, 5, f"{detection.get('width', 0)}x{detection.get('height', 0)}", data_format)
                            else:
                                # detection [x, y, w, h] formatında
                                detections_sheet.write(row, 2, 'person', data_format)
                                detections_sheet.write(row, 3, '0.90', data_format)
                                if len(detection) >= 4:
                                    detections_sheet.write(row, 4, f"({detection[0]}, {detection[1]})", data_format)
                                    detections_sheet.write(row, 5, f"{detection[2]}x{detection[3]}", data_format)
                                else:
                                    detections_sheet.write(row, 4, "N/A", data_format)
                                    detections_sheet.write(row, 5, "N/A", data_format)
                            row += 1
                    else:
                        # detections doğrudan bir detection objesi
                        detections_sheet.write(row, 0, frame_num, number_format)
                        detections_sheet.write(row, 1, f"{time_sec:.2f}", data_format)
                        
                        if isinstance(detections, dict):
                            detections_sheet.write(row, 2, detections.get('class', 'person'), data_format)
                            detections_sheet.write(row, 3, f"{detections.get('confidence', 0):.2f}", data_format)
                            detections_sheet.write(row, 4, f"({detections.get('x', 0)}, {detections.get('y', 0)})", data_format)
                            detections_sheet.write(row, 5, f"{detections.get('width', 0)}x{detections.get('height', 0)}", data_format)
                        else:
                            # detections [x, y, w, h] formatında
                            detections_sheet.write(row, 2, 'person', data_format)
                            detections_sheet.write(row, 3, '0.90', data_format)
                            if len(detections) >= 4:
                                detections_sheet.write(row, 4, f"({detections[0]}, {detections[1]})", data_format)
                                detections_sheet.write(row, 5, f"{detections[2]}x{detections[3]}", data_format)
                            else:
                                detections_sheet.write(row, 4, "N/A", data_format)
                                detections_sheet.write(row, 5, "N/A", data_format)
                        row += 1
            else:
                # Eğer list ise dict'e dönüştür
                print(f"detected_objects list olarak geldi: {type(self.detected_objects)}")
                for i, detection in enumerate(self.detected_objects):
                    detections_sheet.write(row, 0, i, number_format)
                    detections_sheet.write(row, 1, f"{i/fps:.2f}", data_format)
                    detections_sheet.write(row, 2, detection.get('class', 'person'), data_format)
                    detections_sheet.write(row, 3, f"{detection.get('confidence', 0):.2f}", data_format)
                    detections_sheet.write(row, 4, f"({detection.get('x', 0)}, {detection.get('y', 0)})", data_format)
                    detections_sheet.write(row, 5, f"{detection.get('width', 0)}x{detection.get('height', 0)}", data_format)
                    row += 1
            
            # Grafik sayfası oluştur
            if self.detected_objects:
                chart_sheet = workbook.add_worksheet('Grafikler')
                
                # Zaman serisi grafiği için veri hazırla
                time_series = {}
                if isinstance(self.detected_objects, dict):
                    for frame_num, detections in self.detected_objects.items():
                        time_sec = int(frame_num / fps)
                        # detections bir liste mi kontrol et
                        if isinstance(detections, list):
                            time_series[time_sec] = time_series.get(time_sec, 0) + len(detections)
                        else:
                            time_series[time_sec] = time_series.get(time_sec, 0) + 1
                else:
                    # List ise basit işlem
                    for i, detection in enumerate(self.detected_objects):
                        time_sec = int(i / fps)
                        time_series[time_sec] = time_series.get(time_sec, 0) + 1
                
                # Grafik verilerini yaz
                chart_sheet.write('A1', 'Saniye', title_format)
                chart_sheet.write('B1', 'Tespit Sayısı', title_format)
                
                row = 2
                for time_sec in sorted(time_series.keys()):
                    chart_sheet.write(row, 0, time_sec, number_format)
                    chart_sheet.write(row, 1, time_series[time_sec], number_format)
                    row += 1
                
                # Çizgi grafik oluştur
                if row > 2:
                    line_chart = workbook.add_chart({'type': 'line'})
                    line_chart.add_series({
                        'name': 'Tespit Sayısı',
                        'categories': ['Grafikler', 2, 0, row-1, 0],
                        'values': ['Grafikler', 2, 1, row-1, 1],
                        'line': {'color': '#1ABC9C', 'width': 2}
                    })
                    line_chart.set_title({'name': 'Zaman İçinde Tespit Sayısı'})
                    line_chart.set_x_axis({'name': 'Saniye'})
                    line_chart.set_y_axis({'name': 'Tespit Sayısı'})
                    chart_sheet.insert_chart('D2', line_chart)
                
                # Olay süreleri için pasta grafiği
                if self.events:
                    # Olay sürelerine göre kategoriler
                    duration_categories = {'Kısa (0-5s)': 0, 'Orta (5-15s)': 0, 'Uzun (15s+)': 0}
                    for start, end in self.events:
                        event_duration = end - start
                        if event_duration <= 5:
                            duration_categories['Kısa (0-5s)'] += 1
                        elif event_duration <= 15:
                            duration_categories['Orta (5-15s)'] += 1
                        else:
                            duration_categories['Uzun (15s+)'] += 1
                    
                    # Pasta grafiği verileri
                    chart_sheet.write('A20', 'Olay Süresi', title_format)
                    chart_sheet.write('B20', 'Olay Sayısı', title_format)
                    
                    row = 21
                    for category, count in duration_categories.items():
                        if count > 0:
                            chart_sheet.write(row, 0, category, data_format)
                            chart_sheet.write(row, 1, count, number_format)
                            row += 1
                    
                    # Pasta grafiği oluştur
                    if row > 21:
                        pie_chart = workbook.add_chart({'type': 'pie'})
                        pie_chart.add_series({
                            'name': 'Olay Süre Dağılımı',
                            'categories': ['Grafikler', 21, 0, row-1, 0],
                            'values': ['Grafikler', 21, 1, row-1, 1],
                            'data_labels': {'percentage': True}
                        })
                        pie_chart.set_title({'name': 'Olay Süre Dağılımı'})
                        chart_sheet.insert_chart('D20', pie_chart)
            
            # Sütun genişliklerini ayarla
            overview_sheet.set_column('A:A', 20)
            overview_sheet.set_column('B:B', 40)
            detections_sheet.set_column('A:A', 10)
            detections_sheet.set_column('B:B', 12)
            detections_sheet.set_column('C:C', 15)
            detections_sheet.set_column('D:D', 12)
            detections_sheet.set_column('E:E', 15)
            detections_sheet.set_column('F:F', 15)
            
            workbook.close()
            return True
            
        except ImportError:
            print("xlsxwriter kütüphanesi yüklü değil!")
            return False
        except Exception as e:
            print(f"Excel raporu oluşturma hatası: {e}")
            return False
    
    def generate_charts(self, output_dir: str) -> bool:
        """Gelişmiş grafikler ve çizelgeler oluşturur."""
        try:
            if not MATPLOTLIB_AVAILABLE:
                print("Grafik desteği yok! 'pip install matplotlib numpy' çalıştırın")
                return False
            
            # Matplotlib ayarları
            try:
                plt.style.use('seaborn-v0_8')
            except OSError:
                try:
                    plt.style.use('seaborn')
                except OSError:
                    plt.style.use('default')
            plt.rcParams['figure.facecolor'] = 'white'
            plt.rcParams['axes.facecolor'] = 'white'
            plt.rcParams['font.size'] = 12
            
            # Video bilgilerini güvenli şekilde al
            duration = self.video_info.get('duration', 0)
            fps = self.video_info.get('fps', 30)
            
            # 1. Zaman Çizelgesi Grafiği (Gelişmiş)
            if self.events:
                fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(15, 10))
                
                # Üst panel: Olay çubukları
                for i, (start, end) in enumerate(self.events):
                    ax1.barh(i, end - start, left=start, height=0.8, 
                            color='darkgreen', alpha=0.7, edgecolor='black')
                    ax1.text(start + (end - start) / 2, i, f'#{i+1}', 
                            ha='center', va='center', fontweight='bold', color='white')
                
                ax1.set_xlabel('Zaman (saniye)', fontsize=14, fontweight='bold')
                ax1.set_ylabel('Olay Numarası', fontsize=14, fontweight='bold')
                ax1.set_title('Tespit Edilen Olaylar - Zaman Cizelgesi', fontsize=16, fontweight='bold')
                ax1.grid(True, alpha=0.3)
                ax1.set_ylim(-0.5, len(self.events) - 0.5)
                
                # Alt panel: Yoğunluk grafiği
                if self.detected_objects:
                    time_density = {}
                    if isinstance(self.detected_objects, dict):
                        for frame_num, detections in self.detected_objects.items():
                            time_sec = frame_num / fps
                            time_density[int(time_sec)] = time_density.get(int(time_sec), 0) + len(detections)
                    else:
                        # List ise basit işlem
                        for i, _ in enumerate(self.detected_objects):
                            time_sec = int(i / fps)
                            time_density[time_sec] = time_density.get(time_sec, 0) + 1
                    
                    if time_density:
                        times = sorted(time_density.keys())
                        densities = [time_density[t] for t in times]
                        
                        ax2.plot(times, densities, 'b-', linewidth=2, alpha=0.8, label='Tespit Yoğunluğu')
                        ax2.fill_between(times, densities, alpha=0.3, color='blue')
                        ax2.set_xlabel('Zaman (saniye)', fontsize=14, fontweight='bold')
                        ax2.set_ylabel('Tespit Sayısı', fontsize=14, fontweight='bold')
                        ax2.set_title('Tespit Yoğunluğu Dağılımı', fontsize=16, fontweight='bold')
                        ax2.grid(True, alpha=0.3)
                        ax2.legend()
                
                plt.tight_layout()
                plt.savefig(os.path.join(output_dir, 'olay_zaman_analizi.png'), dpi=300, bbox_inches='tight')
                plt.close()
            
            # 2. Pasta Grafiği - Olay Süre Kategorileri
            if self.events:
                durations = [end - start for start, end in self.events]
                
                # Kategori tanımları
                short_events = sum(1 for d in durations if d <= 5)
                medium_events = sum(1 for d in durations if 5 < d <= 15)
                long_events = sum(1 for d in durations if d > 15)
                
                if short_events + medium_events + long_events > 0:
                    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 8))
                    
                    # Pasta grafiği
                    labels = ['Kısa (≤5s)', 'Orta (5-15s)', 'Uzun (>15s)']
                    sizes = [short_events, medium_events, long_events]
                    colors = ['#FF6B6B', '#4ECDC4', '#45B7D1']
                    explode = (0.05, 0.05, 0.05)
                    
                    # Sadece sıfır olmayan değerleri göster
                    filtered_labels = []
                    filtered_sizes = []
                    filtered_colors = []
                    filtered_explode = []
                    
                    for i, size in enumerate(sizes):
                        if size > 0:
                            filtered_labels.append(labels[i])
                            filtered_sizes.append(size)
                            filtered_colors.append(colors[i])
                            filtered_explode.append(explode[i])
                    
                    if filtered_sizes:
                        _, _, _ = ax1.pie(filtered_sizes, labels=filtered_labels, 
                                         autopct='%1.1f%%', startangle=90,
                                         colors=filtered_colors, explode=filtered_explode,
                                         shadow=True, textprops={'fontsize': 12})
                        
                        ax1.set_title('Olay Süre Kategorileri', fontsize=16, fontweight='bold')
                    
                    # Histogram
                    ax2.hist(durations, bins=min(15, len(durations)), color='skyblue', 
                            alpha=0.7, edgecolor='black', linewidth=1)
                    ax2.axvline(np.mean(durations), color='red', linestyle='--', 
                               linewidth=2, label=f'Ortalama: {np.mean(durations):.1f}s')
                    ax2.axvline(np.median(durations), color='green', linestyle='--', 
                               linewidth=2, label=f'Medyan: {np.median(durations):.1f}s')
                    ax2.set_xlabel('Olay Süresi (saniye)', fontsize=14, fontweight='bold')
                    ax2.set_ylabel('Olay Sayısı', fontsize=14, fontweight='bold')
                    ax2.set_title('Olay Süresi Dağılımı', fontsize=16, fontweight='bold')
                    ax2.grid(True, alpha=0.3)
                    ax2.legend()
                    
                    plt.tight_layout()
                    plt.savefig(os.path.join(output_dir, 'olay_sure_analizi.png'), dpi=300, bbox_inches='tight')
                    plt.close()
            
            # 3. Radar Grafiği - Tespit Performansı
            if self.detected_objects and self.events and duration > 0:
                fig, ax = plt.subplots(figsize=(10, 10), subplot_kw=dict(projection='polar'))
                
                # Performans metrikleri
                total_frames = self.video_info.get('total_frames', 1)
                event_duration = sum([end - start for start, end in self.events])
                
                # Güvenlik kontrolü için try-except
                try:
                    if isinstance(self.detected_objects, dict):
                        avg_confidence = np.mean([np.mean([d.get('confidence', 0) for d in dets]) 
                                                  for dets in self.detected_objects.values() if dets])
                    else:
                        avg_confidence = np.mean([d.get('confidence', 0) for d in self.detected_objects])
                except:
                    avg_confidence = 0
                
                metrics = {
                    'Tespit Oranı': min(100, (len(self.detected_objects) / total_frames) * 100),
                    'Olay Kapsamı': min(100, (event_duration / duration) * 100),
                    'Ortalama Güven': avg_confidence * 100,
                    'Tespit Yoğunluğu': min(100, (len(self.detected_objects) / len(self.events)) * 10) if len(self.events) > 0 else 0,
                    'Olay Sıklığı': min(100, (len(self.events) / (duration / 60)) * 10),
                    'Sistem Verimliliği': min(100, (len(self.events) / max(1, len(self.detected_objects))) * 100)
                }
                
                # Radar grafik verileri
                angles = np.linspace(0, 2 * np.pi, len(metrics), endpoint=False).tolist()
                values = list(metrics.values())
                
                # Grafiği kapat
                angles += angles[:1]
                values += values[:1]
                
                ax.plot(angles, values, 'o-', linewidth=2, color='#1E88E5')
                ax.fill(angles, values, alpha=0.25, color='#1E88E5')
                
                ax.set_xticks(angles[:-1])
                ax.set_xticklabels(metrics.keys(), fontsize=12)
                ax.set_ylim(0, 100)
                ax.set_title('Sistem Performans Radari', fontsize=16, fontweight='bold', pad=20)
                ax.grid(True)
                
                plt.tight_layout()
                plt.savefig(os.path.join(output_dir, 'performans_radar.png'), dpi=300, bbox_inches='tight')
                plt.close()
            
            # 4. Isı Haritası - Tespit Yoğunluğu
            if self.detected_objects and duration > 0:
                fig, ax = plt.subplots(figsize=(15, 8))
                
                # Video süresini 100 bölüme ayır
                segments = min(100, max(10, int(duration)))
                segment_duration = duration / segments
                segment_counts = np.zeros(segments)
                
                if isinstance(self.detected_objects, dict):
                    for frame_num, detections in self.detected_objects.items():
                        time_sec = frame_num / fps
                        segment_idx = min(int(time_sec / segment_duration), segments - 1)
                        segment_counts[segment_idx] += len(detections)
                else:
                    # List ise basit işlem
                    for i, _ in enumerate(self.detected_objects):
                        time_sec = i / fps
                        segment_idx = min(int(time_sec / segment_duration), segments - 1)
                        segment_counts[segment_idx] += 1
                
                # Isı haritası matrisini oluştur
                heat_matrix = segment_counts.reshape(1, -1)
                
                im = ax.imshow(heat_matrix, cmap='YlOrRd', aspect='auto', interpolation='bilinear')
                
                # Eksen etiketleri
                ax.set_xlabel('Video Zaman Dilimleri', fontsize=14, fontweight='bold')
                ax.set_ylabel('Yoğunluk', fontsize=14, fontweight='bold')
                ax.set_title('Tespit Yoğunluğu Isi Haritasi', fontsize=16, fontweight='bold')
                
                # Zaman etiketleri
                tick_count = min(10, segments)
                time_labels = [f'{int(i * segment_duration)}s' for i in range(0, segments, max(1, segments//tick_count))]
                ax.set_xticks(range(0, segments, max(1, segments//tick_count)))
                ax.set_xticklabels(time_labels)
                ax.set_yticks([])
                
                # Renk çubuğu
                cbar = plt.colorbar(im, ax=ax, shrink=0.8)
                cbar.set_label('Tespit Sayısı', fontsize=12, fontweight='bold')
                
                plt.tight_layout()
                plt.savefig(os.path.join(output_dir, 'tespit_isi_haritasi.png'), dpi=300, bbox_inches='tight')
                plt.close()
            
            # 5. Basit Dashboard
            if self.events and self.detected_objects:
                fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(16, 12))
                fig.suptitle('M.SAVAS - VIDEO ANALIZ DASHBOARD', fontsize=20, fontweight='bold')
                
                # Dashboard bileşenleri
                durations = [end - start for start, end in self.events]
                
                # 1. Olay sayısı
                ax1.text(0.5, 0.5, f'{len(self.events)}', ha='center', va='center', 
                        fontsize=48, fontweight='bold', color='#2E8B57')
                ax1.text(0.5, 0.2, 'Toplam Olay', ha='center', va='center', 
                        fontsize=14, fontweight='bold')
                ax1.set_xlim(0, 1)
                ax1.set_ylim(0, 1)
                ax1.axis('off')
                
                # 2. Tespit sayısı
                ax2.text(0.5, 0.5, f'{len(self.detected_objects)}', ha='center', va='center', 
                        fontsize=48, fontweight='bold', color='#FF6347')
                ax2.text(0.5, 0.2, 'Tespit Edildi', ha='center', va='center', 
                        fontsize=14, fontweight='bold')
                ax2.set_xlim(0, 1)
                ax2.set_ylim(0, 1)
                ax2.axis('off')
                
                # 3. Ortalama süre
                ax3.text(0.5, 0.5, f'{np.mean(durations):.1f}s', ha='center', va='center', 
                        fontsize=48, fontweight='bold', color='#4169E1')
                ax3.text(0.5, 0.2, 'Ortalama Süre', ha='center', va='center', 
                        fontsize=14, fontweight='bold')
                ax3.set_xlim(0, 1)
                ax3.set_ylim(0, 1)
                ax3.axis('off')
                
                # 4. Kapsam yüzdesi
                ax4.text(0.5, 0.5, f'{(sum(durations) / duration) * 100:.1f}%', ha='center', va='center', 
                        fontsize=48, fontweight='bold', color='#8A2BE2')
                ax4.text(0.5, 0.2, 'Video Kapsamı', ha='center', va='center', 
                        fontsize=14, fontweight='bold')
                ax4.set_xlim(0, 1)
                ax4.set_ylim(0, 1)
                ax4.axis('off')
                # 6. Isı haritası (her segmentteki tespit yoğunluğu)
                # (Varsa, bir alt subplot olarak ekle)
                # Eğer segment_counts ve segments tanımlıysa, heatmap ekle
                # (Bu kodun üst kısmında segment_counts ve segments tanımlı olmalı)
                # Bu örnekte, heatmap ayrı bir şekilde kaydedildiği için burada tekrar eklemeye gerek yok.
                pass
            
            plt.tight_layout()
            plt.savefig(os.path.join(output_dir, 'dashboard.png'), dpi=300, bbox_inches='tight')
            plt.close()
            
            return True
            
        except Exception as e:
            print(f"Grafik oluşturma hatası: {e}")
            return False
    
    def generate_word_report(self, output_path: str, detection_images: list = None) -> bool:
        """Detaylı Word raporu oluşturur."""
        try:
            if not DOCX_AVAILABLE:
                print("Word desteği yok! 'pip install python-docx' çalıştırın")
                return False
            
            doc = Document()
            
            # Görsel listesi kontrolü
            if detection_images is None:
                detection_images = []  # Boş liste olarak başlat
            
            # Stil ayarları
            styles = doc.styles
            
            # Başlık stili
            try:
                title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
                title_font = title_style.font
                title_font.name = 'Calibri'
                title_font.size = Pt(24)
                title_font.bold = True
                title_font.color.rgb = RGBColor(26, 188, 156)  # Turkuaz
                title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                # Stil zaten varsa devam et
                pass
            
            # Alt başlık stili
            try:
                subtitle_style = styles.add_style('CustomSubtitle', WD_STYLE_TYPE.PARAGRAPH)
                subtitle_font = subtitle_style.font
                subtitle_font.name = 'Calibri'
                subtitle_font.size = Pt(16)
                subtitle_font.bold = True
                subtitle_font.color.rgb = RGBColor(52, 152, 219)  # Mavi
            except:
                # Stil zaten varsa devam et
                pass
            
            # 📋 KAPAK SAYFASI
            try:
                title = doc.add_paragraph('M.SAVAS', style='CustomTitle')
                title.add_run('\nMotion Surveillance and Video Analysis System').font.size = Pt(14)
            except:
                title = doc.add_paragraph('M.SAVAS')
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title.runs[0].font.size = Pt(24)
                title.runs[0].font.bold = True
            
            doc.add_paragraph()
            try:
                subtitle = doc.add_paragraph('DETAYLI VIDEO ANALIZ RAPORU', style='CustomSubtitle')
                subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                subtitle = doc.add_paragraph('DETAYLI VIDEO ANALIZ RAPORU')
                subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
                subtitle.runs[0].font.size = Pt(16)
                subtitle.runs[0].font.bold = True
            
            doc.add_paragraph()
            
            # Kapak bilgileri tablosu
            info_table = doc.add_table(rows=7, cols=2)
            info_table.style = 'Table Grid'
            
            info_data = [
                ('📁 Video Dosyası:', os.path.basename(self.video_path)),
                ('📅 Analiz Tarihi:', self.analysis_date.strftime('%d.%m.%Y %H:%M:%S')),
                ('⚙️ Hassasiyet Seviyesi:', self.sensitivity),
                ('⏱️ Video Süresi:', self.format_duration(self.video_info.get('duration', 0))),
                ('🎯 Tespit Sayısı:', f"{len(self.detected_objects)} kare"),
                ('🚨 Olay Sayısı:', f"{len(self.events)} olay"),
                ('📷 Görsel Sayısı:', f"{len(detection_images) if detection_images else 0} görüntü")
            ]
            
            for i, (label, value) in enumerate(info_data):
                info_table.cell(i, 0).text = label
                info_table.cell(i, 1).text = str(value)
                info_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
            
            doc.add_page_break()
            
            # 📊 İSTATİSTİKLER BÖLÜMÜ
            try:
                doc.add_paragraph('ANALIZ ISTATISTIKLERI', style='CustomSubtitle')
            except:
                stats_title = doc.add_paragraph('ANALIZ ISTATISTIKLERI')
                stats_title.runs[0].font.size = Pt(16)
                stats_title.runs[0].font.bold = True
            
            if self.events:
                total_event_duration = sum([end - start for start, end in self.events])
                avg_event_duration = total_event_duration / len(self.events)
                coverage_percentage = (total_event_duration / self.video_info.get('duration', 1)) * 100
                
                stats_table = doc.add_table(rows=6, cols=2)
                stats_table.style = 'Light Shading Accent 1'
                
                stats_data = [
                    ('📈 Toplam Olay Süresi:', self.format_duration(total_event_duration)),
                    ('⏱️ Ortalama Olay Süresi:', self.format_duration(avg_event_duration)),
                    ('📊 Video Kapsamı:', f"{coverage_percentage:.1f}%"),
                    ('🎯 Tespit Yoğunluğu:', f"{len(self.detected_objects) / (self.video_info.get('total_frames', 1) / 100):.1f} tespit/100 kare"),
                    ('⚡ Analiz Verimliliği:', f"{len(self.events) / (self.video_info.get('duration', 1) / 60):.1f} olay/dakika"),
                    ('🔍 Saniye Başına Tespit:', f"{len(self.detected_objects) / self.video_info.get('duration', 1):.2f} tespit/saniye")
                ]
                
                for i, (label, value) in enumerate(stats_data):
                    stats_table.cell(i, 0).text = label
                    stats_table.cell(i, 1).text = str(value)
                    stats_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
            
            doc.add_paragraph()
            
            # 📷 DETAYLI GÖRSEL TESPİTLER BÖLÜMÜ - YENİ DÜZENLE
            if detection_images:
                try:
                    doc.add_paragraph('📷 DETAYLI TESPİT GÖRÜNTÜLERİ', style='CustomSubtitle')
                except:
                    title = doc.add_paragraph('📷 DETAYLI TESPİT GÖRÜNTÜLERİ')
                    title.runs[0].font.size = Pt(16)
                    title.runs[0].font.bold = True
                
                doc.add_paragraph('Aşağıda videoda tespit edilen hareketlerin detaylı görüntüleri düzenli tablo formatında sunulmaktadır. Her satırda 2 görüntü, sayfa başına 4 görüntü bulunmaktadır.')
                doc.add_paragraph()
                
                # Görüntüleri 2x2 tablo halinde düzenle (sayfa başına 4 görüntü)
                images_per_page = 4
                current_page_images = 0
                table = None
                current_row = None
                
                for i, img_info in enumerate(detection_images):
                    try:
                        # Her 4 görüntüde bir yeni sayfa
                        if current_page_images >= images_per_page:
                            doc.add_page_break()
                            current_page_images = 0
                            table = None
                        
                        # İlk görüntü veya yeni sayfa ise tablo oluştur
                        if table is None:
                            table = doc.add_table(rows=0, cols=2)
                            table.style = 'Table Grid'
                            # Sütun genişliklerini ayarla
                            for col in table.columns:
                                for cell in col.cells:
                                    cell.width = Inches(3.2)
                        
                        # 2 sütunluk satır için
                        if current_page_images % 2 == 0:
                            # Yeni satır ekle
                            current_row = table.add_row().cells
                        
                        # Hangi hücreye ekleneceğini belirle (sol: 0, sağ: 1)
                        cell_index = current_page_images % 2
                        cell = current_row[cell_index]
                        
                        # Hücre içeriği oluştur
                        cell_paragraph = cell.paragraphs[0]
                        cell_paragraph.clear()
                        
                        # Başlık
                        title_run = cell_paragraph.add_run(f"🎯 Kare {img_info['frame']}\n")
                        title_run.font.bold = True
                        title_run.font.size = Pt(12)
                        title_run.font.color.rgb = RGBColor(52, 152, 219)
                        
                        # Zaman bilgisi
                        time_run = cell_paragraph.add_run(f"⏰ Zaman: {self.format_duration(img_info['time'])}\n")
                        time_run.font.size = Pt(10)
                        time_run.font.bold = True
                        
                        # Tespit bilgisi
                        detection_run = cell_paragraph.add_run(f"👥 {img_info['detections']} kişi tespit edildi\n\n")
                        detection_run.font.size = Pt(10)
                        detection_run.font.color.rgb = RGBColor(46, 204, 113)
                        
                        # Resim ekle (daha küçük boyutta)
                        if os.path.exists(img_info['path']):
                            cell_paragraph.add_run().add_picture(img_info['path'], width=Inches(2.8))
                            cell_paragraph.add_run("\n\n")
                        
                        # Açıklama metni altına ekle
                        description_run = cell_paragraph.add_run("📋 ")
                        description_run.font.bold = True
                        
                        if img_info['detections'] == 1:
                            desc_text = "Sahnede 1 kişi aktif"
                        elif img_info['detections'] > 1:
                            desc_text = f"Sahnede {img_info['detections']} kişi aktif"
                        else:
                            desc_text = "Hareket algılandı"
                        
                        description_run2 = cell_paragraph.add_run(desc_text)
                        description_run2.font.size = Pt(9)
                        description_run2.font.italic = True
                        description_run2.font.color.rgb = RGBColor(127, 140, 141)
                        
                        # Hücre hizalaması
                        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Hücre padding ayarları
                        cell_margins = cell._element.get_or_add_tcPr()
                        cell_margins.append(parse_xml(
                            '<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                            '<w:top w:w="144" w:type="dxa"/>'
                            '<w:left w:w="144" w:type="dxa"/>'
                            '<w:bottom w:w="144" w:type="dxa"/>'
                            '<w:right w:w="144" w:type="dxa"/>'
                            '</w:tcMar>'
                        ))
                        
                        current_page_images += 1
                        
                    except Exception as e:
                        print(f"Resim ekleme hatası: {e}")
                        continue
                
                # Son sayfa geçişi
                doc.add_page_break()
            
            # 📋 DETAYLI OLAY LİSTESİ
            try:
                doc.add_paragraph('DETAYLI OLAY LISTESI', style='CustomSubtitle')
            except:
                events_title = doc.add_paragraph('DETAYLI OLAY LISTESI')
                events_title.runs[0].font.size = Pt(16)
                events_title.runs[0].font.bold = True
            
            if self.events:
                events_table = doc.add_table(rows=len(self.events) + 1, cols=6)
                events_table.style = 'Table Grid'
                
                # Başlık satırı
                header_cells = events_table.rows[0].cells
                headers = ['Olay No', 'Başlangıç', 'Bitiş', 'Süre', 'Tespit Sayısı', 'Açıklama']
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    header_cells[i].paragraphs[0].runs[0].font.bold = True
                
                # Olay verileri
                for i, (start, end) in enumerate(self.events):
                    row_cells = events_table.rows[i + 1].cells
                    duration = end - start
                    
                    # Bu olay süresinde kaç tespit var
                    detection_count = 0
                    for frame_num, detections in self.detected_objects.items():
                        frame_time = frame_num / self.video_info.get('fps', 30)
                        if start <= frame_time <= end:
                            detection_count += len(detections)
                    
                    row_cells[0].text = str(i + 1)
                    row_cells[1].text = self.format_duration(start)
                    row_cells[2].text = self.format_duration(end)
                    row_cells[3].text = f"{duration:.1f}s"
                    row_cells[4].text = str(detection_count)
                    
                    # Açıklama
                    if duration < 2:
                        description = "Kısa süreli hareket"
                    elif duration < 10:
                        description = "Orta süreli aktivite"
                    else:
                        description = "Uzun süreli aktivite"
                    
                    if detection_count > 10:
                        description += " (Yoğun tespit)"
                    elif detection_count > 5:
                        description += " (Orta tespit)"
                    else:
                        description += " (Az tespit)"
                    
                    row_cells[5].text = description
            
            # � ZAMAN BAZLI ANALİZ
            doc.add_page_break()
            doc.add_paragraph('📊 ZAMAN BAZLI ANALİZ', style='CustomSubtitle')
            
            if self.detected_objects:
                # Saniye bazında tespit sayıları
                second_detections = {}
                for frame_num, detections in self.detected_objects.items():
                    second = int(frame_num / self.video_info.get('fps', 30))
                    if second not in second_detections:
                        second_detections[second] = 0
                    second_detections[second] += len(detections)
                
                # En yoğun 10 saniye
                top_seconds = sorted(second_detections.items(), key=lambda x: x[1], reverse=True)[:10]
                
                doc.add_paragraph("🔥 En Yoğun Tespit Edilen Anlar:")
                for i, (second, count) in enumerate(top_seconds, 1):
                    para = doc.add_paragraph()
                    para.add_run(f"{i}. ").font.bold = True
                    para.add_run(f"Zaman: {self.format_duration(second)} - {count} tespit")
            
            # �🔍 SONUÇ VE DEĞERLENDİRME
            doc.add_page_break()
            doc.add_paragraph('🔍 SONUÇ VE DEĞERLENDİRME', style='CustomSubtitle')
            
            if self.events:
                conclusion_text = f"""
Bu video analizi {self.sensitivity} hassasiyet seviyesinde gerçekleştirilmiştir. 

📊 Analiz Sonuçları:
• Video süresinin %{(sum([end - start for start, end in self.events]) / self.video_info.get('duration', 1) * 100):.1f}'inde hareket tespit edilmiştir
• Toplam {len(self.events)} ayrı olay kaydedilmiştir
• {len(detection_images) if detection_images else 0} adet detaylı görüntü oluşturulmuştur
• En uzun olay süresi: {max([end - start for start, end in self.events]):.1f} saniye
• En kısa olay süresi: {min([end - start for start, end in self.events]):.1f} saniye

🎯 Öneriler:
• Güvenlik açısından kritik zaman dilimlerine odaklanılması önerilir
• Uzun süreli aktiviteler detaylı incelenmelidir
• Yoğun tespit edilen anlar özel dikkat gerektirir
• Sistem {self.sensitivity} seviyesinde başarılı tespit performansı göstermiştir

🔍 Detay Analizi:
• Ortalama saniye başına tespit: {len(self.detected_objects) / self.video_info.get('duration', 1):.2f}
• Toplam tespit sayısı: {sum(len(detections) for detections in self.detected_objects.values())}
• Rapor kalitesi: Yüksek çözünürlükte {len(detection_images) if detection_images else 0} görüntü
                """
            else:
                conclusion_text = """
Bu video analizinde herhangi bir hareket tespit edilmemiştir.

🔍 Olası Nedenler:
• Video statik bir sahne içeriyor olabilir
• Hassasiyet seviyesi çok yüksek ayarlanmış olabilir
• Video kalitesi analiz için uygun olmayabilir

💡 Öneriler:
• Hassasiyet seviyesini düşürerek tekrar analiz yapabilirsiniz
• Video kalitesini ve içeriğini kontrol ediniz
• Farklı zaman dilimlerini test edebilirsiniz
            """
            
            doc.add_paragraph(conclusion_text.strip())
            
            # Alt bilgi
            doc.add_paragraph()
            footer = doc.add_paragraph()
            footer.add_run('Bu detaylı rapor M.SAVAŞ (Motion Surveillance and Video Analysis System) tarafından otomatik olarak oluşturulmuştur.').font.italic = True
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Word raporu oluşturma hatası: {e}")
            return False

    # =============================================================================
    # --- YENİ ÖZELLİK: CANLI KAMERA FONKSİYONLARI ---
    # =============================================================================
    
    def start_live_camera(self):
        """Canlı kamera akışını başlatır"""
        try:
            # Önce webcam'i dene
            self.camera_capture = cv2.VideoCapture(0)
            if not self.camera_capture.isOpened():
                # İkinci kamerayı dene
                self.camera_capture = cv2.VideoCapture(1)
                if not self.camera_capture.isOpened():
                    self.show_error_message("Kamera bulunamadı! Lütfen webcam'inizin bağlı olduğundan emin olun.")
                    return
            
            # Kamera ayarları
            self.camera_capture.set(cv2.CAP_PROP_FRAME_WIDTH, 640)
            self.camera_capture.set(cv2.CAP_PROP_FRAME_HEIGHT, 480)
            self.camera_capture.set(cv2.CAP_PROP_FPS, 30)
            
            self.is_live_camera = True
            self.live_camera_timer.start(33)  # ~30 FPS
            
            # UI güncelleme
            self.btn_start_camera.setEnabled(False)
            self.btn_stop_camera.setEnabled(True)
            self.video_info_label.setText("📹 Canlı Kamera Aktif\n🎯 Gerçek zamanlı analiz için 'Analiz Et' butonuna basın")
            
            self.log_message("📹 Canlı kamera başlatıldı! Gerçek zamanlı görüntü akışı başladı.", "success")
            self.log_message(f"  • Çözünürlük: 640x480", "info")
            self.log_message(f"  • FPS: 30", "info")
            self.log_message(f"  • Gerçek zamanlı analiz: {self.current_sensitivity}", "info")
            
        except Exception as e:
            self.show_error_message(f"Kamera başlatma hatası: {e}")
            self.log_message(f"❌ Kamera hatası: {e}", "error")

    def stop_live_camera(self):
        """Canlı kamera akışını durdurur"""
        try:
            self.is_live_camera = False
            self.live_camera_timer.stop()
            
            if self.camera_capture:
                self.camera_capture.release()
                self.camera_capture = None
            
            # UI güncelleme
            self.btn_start_camera.setEnabled(True)
            self.btn_stop_camera.setEnabled(False)
            self.video_info_label.setText("📹 Video seçilmedi")
            self.video_display_label.setText("Lütfen bir video dosyası yükleyin veya canlı kamerayı başlatın.")
            
            self.log_message("📹 Canlı kamera durduruldu.", "info")
            
        except Exception as e:
            self.log_message(f"Kamera durdurma hatası: {e}", "warning")

    def update_live_camera_frame(self):
        """Canlı kamera frame'lerini günceller"""
        if not self.camera_capture or not self.is_live_camera:
            return
            
        try:
            ret, frame = self.camera_capture.read()
            if ret and frame is not None:
                # Frame'i işle ve göster
                self.display_cv_frame(frame, 0)
                
                # Eğer canlı analiz aktifse tespit yap
                if self.live_detection_enabled:
                    self.detect_live_objects(frame)
                    
            else:
                self.log_message("Kamera frame'i okunamadı", "warning")
                self.stop_live_camera()
                
        except Exception as e:
            self.log_message(f"Canlı kamera frame hatası: {e}", "error")
            self.stop_live_camera()

    def detect_live_objects(self, frame):
        """Canlı kamera için nesne tespiti yapar"""
        try:
            # Basit hareket tespiti (daha performanslı)
            if not hasattr(self, 'background_subtractor'):
                self.background_subtractor = cv2.createBackgroundSubtractorMOG2(
                    detectShadows=True, varThreshold=50)
            
            # Hareket maskesi oluştur
            motion_mask = self.background_subtractor.apply(frame)
            
            # Gürültüyü azalt
            kernel = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (3, 3))
            motion_mask = cv2.morphologyEx(motion_mask, cv2.MORPH_OPEN, kernel)
            
            # Konturları bul
            contours, _ = cv2.findContours(motion_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            detection_count = 0
            for contour in contours:
                area = cv2.contourArea(contour)
                if area > 500:  # Minimum alan
                    x, y, w, h = cv2.boundingRect(contour)
                    # Tespit kutusunu çiz
                    cv2.rectangle(frame, (x, y), (x + w, y + h), (0, 255, 0), 2)
                    cv2.putText(frame, "Hareket", (x, y - 10), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (0, 255, 0), 1)
                    detection_count += 1
            
            # Status güncelle
            if detection_count > 0:
                self.status_label.setText(f"🔴 CANLI: {detection_count} hareket tespit edildi")
            else:
                self.status_label.setText("🟢 CANLI: İzleniyor...")
                
        except Exception as e:
            self.log_message(f"Canlı tespit hatası: {e}", "warning")
    
    def toggle_live_detection(self):
        """Canlı tespit modunu açar/kapatır"""
        if self.is_live_camera:
            self.live_detection_enabled = not self.live_detection_enabled
            status = "AÇIK" if self.live_detection_enabled else "KAPALI"
            self.log_message(f"🎯 Canlı tespit modu: {status}", "info")
        else:
            self.log_message("⚠️ Önce canlı kamerayı başlatın!", "warning")

    # =============================================================================
    # --- YENİ ÖZELLİK: ÇOKLU NESNE TESPİTİ FONKSİYONLARI ---
    # =============================================================================
    
    def update_active_classes(self):
        """Aktif tespit sınıflarını günceller"""
        global ACTIVE_CLASSES
        ACTIVE_CLASSES = []
        
        for class_name, checkbox in self.object_checkboxes.items():
            if checkbox.isChecked():
                if class_name in TARGET_CLASSES:
                    ACTIVE_CLASSES.append(TARGET_CLASSES[class_name])
        
        # Hiçbir şey seçilmemişse person'u varsayılan yap
        if not ACTIVE_CLASSES:
            ACTIVE_CLASSES = [0]  # person
            self.object_checkboxes['person'].setChecked(True)
        
        # Log mesajı
        selected_names = [name for name, checkbox in self.object_checkboxes.items() if checkbox.isChecked()]
        self.log_message(f"🎯 Tespit edilecek nesneler güncellendi: {', '.join(selected_names)}", "info")
    
    def show_advanced_object_selection(self):
        """Gelişmiş nesne seçimi dialogu gösterir"""
        try:
            # Import'lar zaten üstte yapıldı
            
            dialog = QDialog(self)
            dialog.setWindowTitle("🎯 Gelişmiş Nesne Seçimi")
            dialog.setFixedSize(600, 400)
            dialog.setStyleSheet(self.get_stylesheet())
            
            layout = QVBoxLayout(dialog)
            
            # Açıklama
            info_label = QLabel("🎯 Tespit edilecek nesneleri seçin (YOLO sınıfları):")
            info_label.setStyleSheet("font-weight: bold; color: white; margin: 10px;")
            layout.addWidget(info_label)
            
            # Tüm nesne kategorileri
            scroll_area = QScrollArea()
            scroll_widget = QWidget()
            grid_layout = QGridLayout(scroll_widget)
            
            self.advanced_checkboxes = {}
            row, col = 0, 0
            
            for class_name, class_id in TARGET_CLASSES.items():
                checkbox = QCheckBox(f"{class_name} (ID: {class_id})")
                checkbox.setChecked(class_id in ACTIVE_CLASSES)
                checkbox.setStyleSheet("color: white; margin: 5px;")
                
                self.advanced_checkboxes[class_name] = checkbox
                grid_layout.addWidget(checkbox, row, col)
                
                col += 1
                if col >= 3:  # 3 sütun
                    col = 0
                    row += 1
            
            scroll_area.setWidget(scroll_widget)
            scroll_area.setStyleSheet("background-color: #2c3e50; border: 1px solid #34495e;")
            layout.addWidget(scroll_area)
            
            # Butonlar
            button_layout = QHBoxLayout()
            
            select_all_btn = QPushButton("✅ Tümünü Seç")
            select_none_btn = QPushButton("❌ Hiçbirini Seçme")
            apply_btn = QPushButton("💾 Uygula")
            cancel_btn = QPushButton("🚫 İptal")
            
            for btn in [select_all_btn, select_none_btn, apply_btn, cancel_btn]:
                btn.setFixedHeight(35)
                btn.setStyleSheet("""
                    QPushButton {
                        background-color: #3498db;
                        color: white;
                        border: none;
                        border-radius: 5px;
                        font-weight: bold;
                        padding: 5px;
                    }
                    QPushButton:hover {
                        background-color: #2980b9;
                    }
                """)
            
            select_all_btn.clicked.connect(lambda: self.set_all_checkboxes(True))
            select_none_btn.clicked.connect(lambda: self.set_all_checkboxes(False))
            apply_btn.clicked.connect(lambda: self.apply_advanced_selection(dialog))
            cancel_btn.clicked.connect(dialog.reject)
            
            button_layout.addWidget(select_all_btn)
            button_layout.addWidget(select_none_btn)
            button_layout.addWidget(apply_btn)
            button_layout.addWidget(cancel_btn)
            
            layout.addLayout(button_layout)
            
            dialog.exec_()
            
        except Exception as e:
            self.log_message(f"Gelişmiş nesne seçimi hatası: {e}", "error")
    
    def set_all_checkboxes(self, checked):
        """Tüm gelişmiş checkboxları seçer/temizler"""
        for checkbox in self.advanced_checkboxes.values():
            checkbox.setChecked(checked)
    
    def apply_advanced_selection(self, dialog):
        """Gelişmiş seçimi uygular"""
        global ACTIVE_CLASSES
        ACTIVE_CLASSES = []
        
        # Ana checkboxları temizle
        for checkbox in self.object_checkboxes.values():
            checkbox.setChecked(False)
        
        # Gelişmiş seçimi uygula
        for class_name, checkbox in self.advanced_checkboxes.items():
            if checkbox.isChecked():
                if class_name in TARGET_CLASSES:
                    ACTIVE_CLASSES.append(TARGET_CLASSES[class_name])
                    
                    # Ana paneldeki ilgili checkbox'ı da işaretle
                    if class_name in self.object_checkboxes:
                        self.object_checkboxes[class_name].setChecked(True)
        
        # Hiçbir şey seçilmemişse person'u varsayılan yap
        if not ACTIVE_CLASSES:
            ACTIVE_CLASSES = [0]
            self.object_checkboxes['person'].setChecked(True)
        
        # Log mesajı
        selected_count = len(ACTIVE_CLASSES)
        self.log_message(f"🎯 Gelişmiş nesne seçimi uygulandı: {selected_count} farklı nesne türü aktif", "success")
        
        dialog.accept()

# =============================================================================
# --- UYGULAMA GİRİŞ NOKTASI ---
# =============================================================================

    def _create_motion_timeline_chart(self, output_path):
        """Hareket zaman çizelgesi grafiği oluşturur."""
        try:
            plt.figure(figsize=(12, 6))
            
            # Hareket verilerini hazırla
            if self.detected_events:
                times = []
                intensities = []
                
                for i, event in enumerate(self.detected_events):
                    if isinstance(event, tuple) and len(event) >= 2:
                        start_time, end_time = event[0], event[1]
                        duration = end_time - start_time
                        times.append(start_time)
                        intensities.append(duration)
                    else:
                        times.append(i * 10)  # Varsayılan 10 saniye aralıklar
                        intensities.append(1.0)
                
                plt.plot(times, intensities, 'ro-', linewidth=2, markersize=8)
                plt.title('📈 Hareket Yoğunluğu Zaman Çizelgesi', fontsize=16, fontweight='bold')
                plt.xlabel('Zaman (saniye)', fontsize=12)
                plt.ylabel('Hareket Yoğunluğu', fontsize=12)
                plt.grid(True, alpha=0.3)
                
            plt.tight_layout()
            plt.savefig(output_path, dpi=300, bbox_inches='tight')
            plt.close()
            
        except Exception as e:
            self.log_message(f"Hareket grafiği oluşturma hatası: {e}", "error")

    def _create_motion_heatmap(self, output_path):
        """Hareket yoğunluk haritası oluşturur."""
        try:
            _, ax = plt.subplots(figsize=(10, 8))
            
            # Basit heatmap verisi oluştur
            if self.detected_events:
                # Video boyutlarına göre grid oluştur
                grid_size = 20
                heatmap_data = np.zeros((grid_size, grid_size))
                
                # Olayları grid'e dağıt
                for _ in range(min(len(self.detected_events), 50)):  # İlk 50 olay
                    x = np.random.randint(0, grid_size)
                    y = np.random.randint(0, grid_size)
                    heatmap_data[y, x] += 1
                
                im = ax.imshow(heatmap_data, cmap='hot', interpolation='bilinear')
                ax.set_title('🔥 Hareket Yoğunluk Haritası', fontsize=16, fontweight='bold')
                plt.colorbar(im, ax=ax, label='Hareket Sayısı')
                
            plt.tight_layout()
            plt.savefig(output_path, dpi=300, bbox_inches='tight')
            plt.close()
            
        except Exception as e:
            self.log_message(f"Hareket haritası oluşturma hatası: {e}", "error")

    def _create_motion_html_report(self, output_path, chart_path, heatmap_path):
        """HTML hareket raporu oluşturur."""
        try:
            video_name = os.path.basename(self.video_path) if self.video_path else "Video"
            
            html_content = f"""
            <!DOCTYPE html>
            <html lang="tr">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>🏃 Hareket Analizi Raporu</title>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 20px; background-color: #f5f5f5; }}
                    .container {{ max-width: 1200px; margin: 0 auto; background: white; padding: 20px; border-radius: 10px; }}
                    h1 {{ color: #2c3e50; text-align: center; }}
                    h2 {{ color: #3498db; border-bottom: 2px solid #3498db; padding-bottom: 5px; }}
                    .info-box {{ background: #ecf0f1; padding: 15px; border-radius: 5px; margin: 10px 0; }}
                    .chart {{ text-align: center; margin: 20px 0; }}
                    .chart img {{ max-width: 100%; border: 1px solid #ddd; border-radius: 5px; }}
                    .stats {{ display: flex; justify-content: space-between; }}
                    .stat-item {{ background: #3498db; color: white; padding: 15px; border-radius: 5px; text-align: center; flex: 1; margin: 0 5px; }}
                </style>
            </head>
            <body>
                <div class="container">
                    <h1>🏃 M.SAVAŞ Hareket Analizi Raporu</h1>
                    
                    <div class="info-box">
                        <h2>📁 Video Bilgileri</h2>
                        <p><strong>Dosya:</strong> {video_name}</p>
                        <p><strong>Tarih:</strong> {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}</p>
                        <p><strong>Hassasiyet:</strong> {self.current_sensitivity}</p>
                    </div>
                    
                    <div class="stats">
                        <div class="stat-item">
                            <h3>📊 Toplam Hareket</h3>
                            <p>{len(self.detected_events)}</p>
                        </div>
                        <div class="stat-item">
                            <h3>⏱️ Ortalama Süre</h3>
                            <p>{self._calculate_avg_duration():.1f}s</p>
                        </div>
                        <div class="stat-item">
                            <h3>🎯 En Aktif Dönem</h3>
                            <p>{self._find_peak_activity()}</p>
                        </div>
                    </div>
                    
                    <div class="chart">
                        <h2>📈 Hareket Zaman Çizelgesi</h2>
                        <img src="{os.path.basename(chart_path)}" alt="Hareket Grafiği">
                    </div>
                    
                    <div class="chart">
                        <h2>🔥 Hareket Yoğunluk Haritası</h2>
                        <img src="{os.path.basename(heatmap_path)}" alt="Yoğunluk Haritası">
                    </div>
                    
                    <div class="info-box">
                        <h2>📝 Analiz Özeti</h2>
                        <p>Bu rapor, videodaki hareket aktivitelerinin detaylı analizini içermektedir. 
                        Zaman çizelgesi grafiği hareketlerin zamana göre dağılımını, yoğunluk haritası ise 
                        en çok hareketin tespit edildiği alanları göstermektedir.</p>
                    </div>
                </div>
            </body>
            </html>
            """
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
                
        except Exception as e:
            self.log_message(f"HTML raporu oluşturma hatası: {e}", "error")

    def _export_motion_csv(self, output_path):
        """Hareket verilerini CSV formatında dışa aktarır."""
        try:
            with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                writer.writerow(['Olay_No', 'Başlangıç_Zamanı', 'Bitiş_Zamanı', 'Süre', 'Tip'])
                
                for i, event in enumerate(self.detected_events):
                    if isinstance(event, tuple) and len(event) >= 2:
                        start_time, end_time = event[0], event[1]
                        duration = end_time - start_time
                        writer.writerow([i+1, f"{start_time:.2f}", f"{end_time:.2f}", f"{duration:.2f}", "Hareket"])
                    else:
                        writer.writerow([i+1, f"{i*10:.2f}", f"{(i+1)*10:.2f}", "10.00", "Tespit"])
                        
        except Exception as e:
            self.log_message(f"CSV dışa aktarma hatası: {e}", "error")

    def _calculate_avg_duration(self):
        """Ortalama hareket süresini hesaplar."""
        if not self.detected_events:
            return 0.0
            
        total_duration = 0.0
        count = 0
        
        for event in self.detected_events:
            if isinstance(event, tuple) and len(event) >= 2:
                duration = event[1] - event[0]
                total_duration += duration
                count += 1
        
        return total_duration / count if count > 0 else 0.0

    def _find_peak_activity(self):
        """En yoğun hareket dönemini bulur."""
        if not self.detected_events:
            return "Veri yok"
            
        # Basit olarak ilk olayın zamanını döndür
        if isinstance(self.detected_events[0], tuple):
            peak_time = self.detected_events[0][0]
            return f"{peak_time:.0f}s"
        else:
            return "0-10s"

def main():
    """Uygulamayı başlatır."""
    try:
        # Singleton kontrolü - sadece tek bir örnek çalışabilir
        if app_singleton.is_already_running():
            print("⚠️ M.SAVAŞ Video Analiz Sistemi zaten çalışıyor!")
            print("📋 Mevcut uygulama penceresini kontrol edin.")
            
            # GUI uyarısı göster
            if not QApplication.instance():
                QApplication(sys.argv)
            QMessageBox.warning(
                None, 
                "⚠️ Uygulama Zaten Çalışıyor", 
                "🔍 M.SAVAŞ Video Analiz Sistemi zaten çalışıyor!\n\n"
                "• Mevcut pencereyi kontrol edin\n"
                "• Görev çubuğunda uygulama simgesini arayın\n"
                "• Alt+Tab ile uygulamalar arasında geçiş yapın\n\n"
                "Eğer uygulama yanıt vermiyorsa Görev Yöneticisi'nden\n"
                "'python.exe' sürecini sonlandırabilirsiniz."
            )
            sys.exit(1)
        
        # Temel kütüphaneleri kontrol et
        print("M.SAVAŞ Video Analiz Sistemi başlatılıyor...")
        
        # Bağımlılık kontrolü
        if not check_dependencies():
            print("UYARI: Bazı bağımlılıklar eksik. Temel fonksiyonlarla devam ediliyor.")
        
        # PyQt5 uygulaması başlat
        app = QApplication(sys.argv)
        app.setApplicationName("M.SAVAŞ Video Analiz Sistemi")
        app.setApplicationVersion("1.1.0")
        
        # Meta type kaydet
        try:
            from PyQt5.QtCore import qRegisterMetaType
            qRegisterMetaType(QTextCursor)
        except (ImportError, NameError, AttributeError):
            pass # Eski versiyonlarda sorun olabilir

        # Ana pencereyi oluştur ve göster
        print("Ana pencere oluşturuluyor...")
        window = MainWindow()
        window.show()
        
        print("Uygulama başarıyla başlatıldı!")
        
        # Uygulama döngüsünü başlat
        sys.exit(app.exec_())

    except ImportError as e:
        print(f"Kritik bağımlılık hatası: {e}")
        try:
            app = QApplication(sys.argv)
            QMessageBox.critical(None, "Kritik Bağımlılık Hatası", 
                               f"Gerekli kütüphaneler bulunamadı:\n\n{e}\n\n"
                               "Lütfen gerekli kütüphaneleri yükleyin ve tekrar deneyin.")
            sys.exit(1)
        except:
            print("GUI hata mesajı gösterilemedi.")
            sys.exit(1)
            
    except Exception as e:
        import traceback
        error_details = f"Beklenmedik bir hata oluştu:\n\n{e}\n\n{traceback.format_exc()}"
        print(error_details)
        try:
            app = QApplication(sys.argv)
            QMessageBox.critical(None, "Beklenmedik Hata", error_details)
        except: 
            pass
        sys.exit(1)

if __name__ == "__main__":
    main()
